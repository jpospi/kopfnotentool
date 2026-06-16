print("Starting...")
import tkinter as tk
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext, simpledialog
import sqlite3
import threading
import queue
import logging
import sys
import os
import json
import shutil
import tempfile
import re
import io
import statistics
import pandas as pd
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Tuple, Optional, Any, Callable
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from app_paths import load_app_paths

APP_PATHS = load_app_paths()
DEFAULT_SCHOOL_YEAR = "2024/2025"
DEFAULT_TERM = 1

# Ensure temp directory exists with proper permissions
temp_dir = APP_PATHS.temp_dir
temp_dir.mkdir(exist_ok=True, parents=True)
if os.name != "nt" and os.path.exists(temp_dir):
    os.chmod(temp_dir, 0o755) # Set proper permissions

# Ensure logs directory exists
log_dir = APP_PATHS.logs_dir
log_dir.mkdir(exist_ok=True, parents=True)

# Logging-Konfiguration
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler(str(APP_PATHS.logs_dir / "kopfnoten_gui.log"), encoding="utf-8"),
        logging.StreamHandler(sys.stdout),
    ],
)

# Globale Konstanten
TEMPLATE_MANAGER_ENABLED = False  # UI aus; Code bleibt für spätere Wiederaktivierung
GITHUB_REPO_URL = "https://github.com/jpospi/kopfnotentool"
MAX_CLASSES_PER_JAHRGANG = 9  # Autoimport bis 9 Züge (05a … 05i)
CLASS_SUFFIX_LETTERS = "abcdefghi"

FAECHER_MAPPING = {
    # Deutsch
    "De": "Deutsch",
    "Deu": "Deutsch",
    "Deutsch": "Deutsch",
    # Mathematik
    "Ma": "Mathematik",
    "Mat": "Mathematik",
    "Mathematik": "Mathematik",
    # Englisch
    "En": "Englisch",
    "Eng": "Englisch",
    "Englisch": "Englisch",
    # Ethik
    "Et": "Ethik",
    "Eth": "Ethik",
    "Ethik": "Ethik",
    "Ethik (Et)": "Ethik",
    # Kunst
    "Ku": "Kunst",
    "Kun": "Kunst",
    "Kunst": "Kunst",
    # Musik
    "Mu": "Musik",
    "Mus": "Musik",
    "Musik": "Musik",
    # Naturwissenschaften
    "Na": "Naturwissenschaften",
    "Nat": "Naturwissenschaften",
    "Naturwissenschaften": "Naturwissenschaften",
    # Sport
    "Sp": "Sport",
    "Spo": "Sport",
    "Sport": "Sport",
    # Informatik
    "In": "Informatik",
    "Inf": "Informatik",
    "Informatik": "Informatik",
    # Physik
    "Ph": "Physik",
    "Phy": "Physik",
    "Physik": "Physik",
    # Chemie
    "Ch": "Chemie",
    "Che": "Chemie",
    "Chemie": "Chemie",
    # Biologie
    "Bi": "Biologie",
    "Bio": "Biologie",
    "Biologie": "Biologie",
    # Geschichte -> Gesellschaftslehre
    "Ge": "Gesellschaftslehre",
    "Ges": "Gesellschaftslehre",
    "Geschichte": "Gesellschaftslehre",
    # Gesellschaftskunde -> Gesellschaftslehre
    "Gl": "Gesellschaftslehre",
    "Gesell": "Gesellschaftslehre",
    "Gesellschaftskunde": "Gesellschaftslehre",
    # Französisch
    "Fr": "Französisch",
    "Fra": "Französisch",
    "Französisch": "Französisch",
    # Spanisch
    "Es": "Spanisch",
    "Spa": "Spanisch",
    "Spanisch": "Spanisch",
    # Arbeitslehre
    "Al": "Arbeitslehre",
    "Arbeitslehre": "Arbeitslehre",
    # Religion
    "Re": "Religion",
    "Rel": "Religion",
    "Rel.1": "Religion",
    "Religion": "Religion",
    "Religion (ev)": "Religion (ev)",
    "Religion (kath)": "Religion (kath)",
    "Rev": "Religion (ev)",
    "Rrk": "Religion (kath)",
    # Neue Fächer
    "Holz": "Holzbearbeitung",
    "Holzbearbeitung": "Holzbearbeitung",
    "Präs": "Präsentationstechnik",
    "Präsentationstechnik": "Präsentationstechnik",
    "Ernährung": "Ernährung",
    "Ernährung und Gesundheit": "Ernährung und Gesundheit",
    "Kraft": "Kraftsport",
    "Kraftsport": "Kraftsport",
    "Kreativ": "Kreativ gestalten",
    "Kreativ gestalten": "Kreativ gestalten",
    "Textil": "Textiles Gestalten",
    "Textiles Gestalten": "Textiles Gestalten",
}

# Fächer, die eine Gruppe bilden (gegenseitiger Ausschluss)
TRIAD_RELIGION_ETHIK = [
    ("Ethik", None),
    ("Religion", "evangelisch"),
    ("Religion", "katholisch")
]

# Definition der Status-Typen und Mapping
# Standard: "Nebenfach"
# WPU-Flag wird dynamisch erkannt
SUBJECT_STATUS_CONFIG = {
    # Hauptfächer (explizit)
    "Deutsch": "Hauptfach",
    "Mathematik": "Hauptfach",
    "Englisch": "Hauptfach",
    "Gesellschaftslehre": "Hauptfach",
    
    # Nebenfächer (explizit zur Sicherheit, aber Fallback ist Nebenfach)
    "Französisch": "Nebenfach (WPU)",
    "Spanisch": "Nebenfach (WPU)",
    "Ethik": "Nebenfach",
    "Religion": "Nebenfach",
    "Religion (ev)": "Nebenfach",
    "Religion (kath)": "Nebenfach",
    "Kunst": "Nebenfach",
    "Musik": "Nebenfach",
    "Naturwissenschaften": "Nebenfach",
    "Sport": "Nebenfach",
    "Informatik": "Nebenfach (WPU)",
    "Physik": "Nebenfach",
    "Chemie": "Nebenfach",
    "Biologie": "Nebenfach",
    "Arbeitslehre": "Nebenfach",
    "Praxistag": "Nebenfach",
    "Holzbearbeitung": "Nebenfach (WPU)",
    "Präsentationstechnik": "Nebenfach (WPU)",
    "Ernährung": "Nebenfach (WPU)",
    "Ernährung und Gesundheit": "Nebenfach (WPU)",
    "Kraftsport": "Nebenfach (WPU)",
    "Kreativ gestalten": "Nebenfach (WPU)",
    "Textiles Gestalten": "Nebenfach (WPU)",
}

# WPU Fächer Muster (für automatische Erkennung)
WPU_PATTERNS = ["WPU", "WPU1", "WPU2", "WP", "WP1", "WP2", "(W)"]

class LinuxPathManager:
    """Linux-spezifische Pfad-Verwaltung"""
    @staticmethod
    def ensure_directory(path: Path) -> Path:
        """Erstellt Verzeichnis mit korrekten Linux-Berechtigungen"""
        path = Path(path)
        try:
            path.mkdir(parents=True, exist_ok=True)
            os.chmod(path, 0o755)
            return path
        except PermissionError:
            logging.error(f"Berechtigung verweigert für: {path}")
            raise
        except Exception as e:
            logging.error(f"Fehler beim Erstellen von {path}: {e}")
            raise

    @staticmethod
    def check_file_permissions(file_path: Path) -> bool:
        """Prüft Dateiberechtigungen unter Linux"""
        file_path = Path(file_path)
        if not file_path.exists():
            return False
        try:
            with open(file_path, "rb") as f:
                f.read(1)
            return True
        except PermissionError:
            logging.warning(f"Keine Berechtigung für: {file_path}")
            return False
        except Exception:
            return False

class SimpleTemplateDesigner:
    """Vereinfachter Template-Designer ohne komplexe Validierung"""
    def __init__(self, parent):
        self.parent = parent
        self.logger = logging.getLogger("template_designer")

    def create_template_designer_window(self):
        """Öffnet Template-Designer-Fenster"""
        designer_window = tk.Toplevel(self.parent)
        designer_window.title("Template-Designer")
        designer_window.geometry("800x600")

        # Header
        header_frame = ttk.Frame(designer_window)
        header_frame.pack(fill=tk.X, padx=10, pady=5)
        ttk.Label(
            header_frame, text="Template-Designer", font=("Arial", 16, "bold")
        ).pack(anchor=tk.W)
        ttk.Label(
            header_frame,
            text="Erstellen Sie einfache Templates für den Export",
            font=("Arial", 10),
        ).pack(anchor=tk.W)

        # Template-Optionen
        options_frame = ttk.LabelFrame(designer_window, text="Template-Einstellungen")
        options_frame.pack(fill=tk.X, padx=10, pady=5)

        # Template-Typ Auswahl
        type_frame = ttk.Frame(options_frame)
        type_frame.pack(fill=tk.X, padx=5, pady=5)
        ttk.Label(type_frame, text="Template-Typ:").pack(side=tk.LEFT)
        template_type = tk.StringVar(value="horizontal")
        ttk.Radiobutton(
            type_frame,
            text="Horizontal (3 Zeilen)",
            variable=template_type,
            value="horizontal",
        ).pack(side=tk.LEFT, padx=10)

        # Spaltenanzahl
        cols_frame = ttk.Frame(options_frame)
        cols_frame.pack(fill=tk.X, padx=5, pady=5)
        ttk.Label(cols_frame, text="Max. Spalten:").pack(side=tk.LEFT)
        max_cols = tk.StringVar(value="15")
        ttk.Spinbox(cols_frame, from_=10, to=20, textvariable=max_cols, width=5).pack(
            side=tk.LEFT, padx=5
        )

        # Vorschau
        preview_frame = ttk.LabelFrame(designer_window, text="Vorschau")
        preview_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        preview_text = scrolledtext.ScrolledText(
            preview_frame, height=20, font=("Courier", 10)
        )
        preview_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        def update_preview():
            """Aktualisiert die Template-Vorschau"""
            template_content = self.generate_template_content(
                template_type.get(), int(max_cols.get())
            )
            preview_text.delete(1.0, tk.END)
            preview_text.insert(1.0, template_content)

        # Buttons
        button_frame = ttk.Frame(designer_window)
        button_frame.pack(fill=tk.X, padx=10, pady=5)
        ttk.Button(
            button_frame, text="Vorschau aktualisieren", command=update_preview
        ).pack(side=tk.LEFT, padx=5)
        ttk.Button(
            button_frame,
            text="Template erstellen",
            command=lambda: self.create_template_file(
                template_type.get(), int(max_cols.get()), designer_window
            ),
        ).pack(side=tk.LEFT, padx=5)
        ttk.Button(
            button_frame, text="Schließen", command=designer_window.destroy
        ).pack(side=tk.RIGHT, padx=5)

        # Initiale Vorschau
        update_preview()

    def generate_template_content(self, template_type: str, max_cols: int) -> str:
        """Generiert Template-Inhalt als Text-Vorschau"""
        return (
            """KOPFNOTEN - KLASSE {{klasse}}
Export-Datum: {{export_datum}}

Horizontale Tabelle (3 Zeilen × dynamische Spaltenanzahl):

{% for schueler in schueler_liste %}
Schüler: {{schueler.name}}
Fach: {% for fach in schueler.faecher_spalten %}{{fach}}{% if not loop.last %} | {% endif %}{% endfor %}
AV-Note: {% for note in schueler.av_noten %}{{note}}{% if not loop.last %} | {% endif %}{% endfor %}
SV-Note: {% for note in schueler.sv_noten %}{{note}}{% if not loop.last %} | {% endif %}{% endfor %}
{% if not schueler.ist_letzter %}---{% endif %}
{% endfor %}
"""
        )

    @staticmethod
    def create_working_horizontal_template(filename: str, max_cols: int = 15):
        """Creates a working horizontal template with dynamic columns per student"""
        try:
            from docxtpl import DocxTemplate
            from docx import Document
            from docx.shared import Inches
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from tkinter import messagebox
            from pathlib import Path
            from docx.oxml.parser import parse_xml
            import logging

            # Create new document
            doc = Document()

            # Header section (outside any loops)
            header_para = doc.add_paragraph()
            header_para.add_run("KOPFNOTEN - KLASSE {{ klasse }}")
            
            date_para = doc.add_paragraph()
            date_para.add_run("Export-Datum: {{ export_datum }}")
            
            doc.add_paragraph() # Empty line for spacing

            # FOR loop opening - in its own paragraph
            for_para = doc.add_paragraph()
            for_para.add_run("{% for schueler in schueler_liste %}")

            # Student name heading - in its own paragraph
            name_para = doc.add_paragraph()
            name_para.add_run("{{ schueler.name }}")
            
            doc.add_paragraph() # Spacing before table

            # Create a placeholder table with just the first column
            # The real dynamic table will be created during template rendering
            table = doc.add_table(rows=3, cols=1)
            table.style = "Table Grid"
            table.autofit = False
            
            # Headers for first column
            headers = ["Fach", "AV", "SV"]
            for i, header in enumerate(headers):
                cell = table.cell(i, 0)
                run = cell.paragraphs[0].add_run(header)
                run.bold = True
            
            # Add Jinja code for dynamically adding the subject columns
            # This will insert the necessary columns during template rendering
            dynamic_code_para = doc.add_paragraph()
            dynamic_code_para.add_run("{{! ! ! DYNAMIC_TABLE_PLACEHOLDER ! ! !}}")
            
            # Add spacing after table
            doc.add_paragraph()

            # Page break control - in separate paragraphs
            # First the if statement
            if_para = doc.add_paragraph()
            if_para.add_run("{% if not schueler.ist_letzter %}")
            
            # Then the page break in its own paragraph
            break_para = doc.add_paragraph()
            break_run = break_para.add_run()
            break_run._r.append(parse_xml(r'<w:br w:type="page"/>'))
            
            # Close the if block
            endif_para = doc.add_paragraph()
            endif_para.add_run("{% endif %}")
            
            # Add spacing before endfor
            doc.add_paragraph()
            
            # Close the for loop - in its own paragraph
            endfor_para = doc.add_paragraph()
            endfor_para.add_run("{% endfor %}")
            
            # Save template
            doc.save(filename)
            
            # Success message
            messagebox.showinfo(
                "Template erstellt",
                f"Template erfolgreich erstellt:\n"
                f"{Path(filename).name}\n\n"
                f"Typ: horizontal mit dynamischer Spaltenanzahl\n"
                f"Optimiert für variable Fächeranzahl pro Schüler\n"
                f"Beschriftungsbreite: 1.0''\n"
                f"Tabellenlänge: 8.0'' (fest)"
            )
            
        except Exception as e:
            logging.error(f"Fehler beim Erstellen des Templates: {e}")
            messagebox.showerror("Template-Fehler", f"Fehler beim Erstellen: {e}")

    def create_template_file(self, template_type: str, max_cols: int, parent_window):
        """Erstellt Template-Datei über Dialog"""
        try:
            # Get filename
            filename = filedialog.asksaveasfilename(
                title="Template speichern",
                defaultextension=".docx",
                filetypes=[("Word-Dokument", "*.docx")],
                initialdir=str(APP_PATHS.templates_dir),
            )
            if not filename:
                return

            # Create template
            if template_type == "horizontal":
                # Create temporary directory for template creation
                with tempfile.TemporaryDirectory() as temp_dir:
                    temp_path = Path(temp_dir) / Path(filename).name
                    # Create template
                    self.create_working_horizontal_template(str(temp_path), max_cols)
                    # If template was created successfully, copy it to final location
                    if temp_path.exists():
                        shutil.copy2(temp_path, filename)
                parent_window.destroy()
            else:
                messagebox.showinfo(
                    "Info", "Vertikale Templates werden nicht unterstützt."
                )
        except Exception as e:
            self.logger.error(f"Fehler beim Erstellen der Template-Datei: {e}")
            messagebox.showerror("Template-Fehler", f"Fehler: {e}")

class KopfnotenImporter:
    """Import-Klasse für Excel-Dateien"""
    def __init__(self, db_path: str, school_year: str = DEFAULT_SCHOOL_YEAR, term: int = DEFAULT_TERM):
        self.db_path = Path(db_path)
        self.conn = None
        self.faecher_cache = {}
        self.logger = logging.getLogger("importer")
        self.school_year = school_year
        self.term = int(term)

    def __enter__(self):
        self.conn = self._create_database()
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        if self.conn:
            self.conn.close()

    def _create_database(self) -> sqlite3.Connection:
        """Erstellt die Datenbank mit normalisierten Tabellen"""
        self.db_path.parent.mkdir(parents=True, exist_ok=True)
        conn = sqlite3.connect(self.db_path)
        conn.execute("PRAGMA foreign_keys = ON")
        conn.executescript(
            """
            CREATE TABLE IF NOT EXISTS schueler (
                schueler_id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL,
                klasse TEXT NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                target_subjects INTEGER,
                is_active BOOLEAN DEFAULT 1,
                UNIQUE(name, klasse)
            );

            CREATE TABLE IF NOT EXISTS faecher (
                fach_id INTEGER PRIMARY KEY AUTOINCREMENT,
                fach_kurz TEXT NOT NULL,
                fach_lang TEXT NOT NULL,
                fach_typ TEXT,
                ist_wahlpflicht BOOLEAN DEFAULT 0,
                wahlpflicht_gruppe TEXT,
                UNIQUE(fach_kurz, fach_typ, wahlpflicht_gruppe)
            );

            CREATE TABLE IF NOT EXISTS noten (
                noten_id INTEGER PRIMARY KEY AUTOINCREMENT,
                schueler_id INTEGER NOT NULL,
                fach_id INTEGER NOT NULL,
                note_av INTEGER CHECK(note_av BETWEEN 1 AND 6),
                note_sv INTEGER CHECK(note_sv BETWEEN 1 AND 6),
                note_av_special TEXT,
                note_sv_special TEXT,
                manual_av_lock BOOLEAN DEFAULT 0,
                manual_sv_lock BOOLEAN DEFAULT 0,
                ist_wahlpflicht_belegung BOOLEAN DEFAULT 0,
                lehrer_kuerzel TEXT,
                schuljahr TEXT DEFAULT '2024/2025',
                halbjahr INTEGER DEFAULT 1,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (schueler_id) REFERENCES schueler(schueler_id),
                FOREIGN KEY (fach_id) REFERENCES faecher(fach_id),
                UNIQUE(schueler_id, fach_id, schuljahr, halbjahr)
            );

            CREATE INDEX IF NOT EXISTS idx_schueler_klasse ON schueler(klasse);
            CREATE INDEX IF NOT EXISTS idx_noten_schueler ON noten(schueler_id);
            """
        )
        conn.commit()
        
        # MIGRATION: Check if target_subjects exists (for existing DBs)
        try:
            cursor = conn.execute("PRAGMA table_info(schueler)")
            columns = [info[1] for info in cursor.fetchall()]
            if "target_subjects" not in columns:
                conn.execute("ALTER TABLE schueler ADD COLUMN target_subjects INTEGER")
            if "is_active" not in columns:
                conn.execute("ALTER TABLE schueler ADD COLUMN is_active BOOLEAN DEFAULT 1")
                conn.commit()
        except Exception as e:
            logging.error(f"Migration error (target_subjects): {e}")

        # Migration: manuelle Noten-Priorität (Locks je AV/SV)
        try:
            cursor = conn.execute("PRAGMA table_info(noten)")
            note_columns = [info[1] for info in cursor.fetchall()]
            if "manual_av_lock" not in note_columns:
                conn.execute("ALTER TABLE noten ADD COLUMN manual_av_lock BOOLEAN DEFAULT 0")
            if "manual_sv_lock" not in note_columns:
                conn.execute("ALTER TABLE noten ADD COLUMN manual_sv_lock BOOLEAN DEFAULT 0")
            conn.commit()
        except Exception as e:
            logging.error(f"Migration error (manual lock columns): {e}")

        # Migration: Sonderkürzel für Noten (z.B. GB/NF)
        try:
            cursor = conn.execute("PRAGMA table_info(noten)")
            note_columns = [info[1] for info in cursor.fetchall()]
            if "note_av_special" not in note_columns:
                conn.execute("ALTER TABLE noten ADD COLUMN note_av_special TEXT")
            if "note_sv_special" not in note_columns:
                conn.execute("ALTER TABLE noten ADD COLUMN note_sv_special TEXT")
            conn.commit()
        except Exception as e:
            logging.error(f"Migration error (special note columns): {e}")

        # Migration: Geschichte/Gesellschaftskunde -> Gesellschaftslehre
        conn.execute(
            "UPDATE faecher SET fach_lang = 'Gesellschaftslehre' WHERE fach_lang IN ('Geschichte', 'Gesellschaftskunde')"
        )
        conn.commit()
        return conn

    def _parse_note_mit_wahlpflicht(self, note_str: str) -> Tuple[Optional[int], Optional[str], bool, Optional[str]]:
        """Extrahiert Note/Sonderkürzel, Wahlpflicht-Flag und Lehrer-Kürzel aus einem Notenwert
        
        Format: Note kann alleine oder mit Lehrer-Kürzel vorkommen:
        - "3" -> (3, False, None)
        - "3\nGEO,RET" -> (3, False, "GEO,RET")
        - "2 (W)" -> (2, True, None)
        - "2 (W)\nMÜL" -> (2, True, "MÜL")
        """
        if pd.isna(note_str) or note_str == "":
            return None, None, False, None
        
        note_str = str(note_str).strip()
        
        # Split by newline to separate note from teacher
        parts = note_str.split('\n')
        note_part = parts[0].strip()
        lehrer_kuerzel = parts[1].strip() if len(parts) > 1 and parts[1].strip() else None
        
        # WP-Kennungen: W, WP, WP1, WP2, WPU, WPU1, WPU2, Ergänzung Praxistag
        wp_patterns = [
            r"\(W\)", r"\(WP\)", r"\(WPU\)", r"\(WP1\)", r"\(WP2\)", r"\(WPU1\)", r"\(WPU2\)",
            r"\bW\b", r"\bWP\b", r"\bWPU\b", r"\bWP1\b", r"\bWP2\b", r"\bWPU1\b", r"\bWPU2\b",
            r"Praxistag"
        ]
        
        # Parse the note part for WP
        ist_wahlpflicht = False
        for pattern in wp_patterns:
            if re.search(pattern, note_part, re.IGNORECASE):
                ist_wahlpflicht = True
                break
        
        # Check teacher initials for WP markers and clean them
        if lehrer_kuerzel:
            for pattern in wp_patterns:
                if re.search(pattern, lehrer_kuerzel, re.IGNORECASE):
                    ist_wahlpflicht = True
                    lehrer_kuerzel = re.sub(pattern, "", lehrer_kuerzel, flags=re.IGNORECASE).strip()
            
            # Remove trailing commas or spaces that might remain
            lehrer_kuerzel = re.sub(r"[,;\s]+$", "", lehrer_kuerzel).strip()
            lehrer_kuerzel = re.sub(r"^[,;\s]+", "", lehrer_kuerzel).strip()
            
            if not lehrer_kuerzel:
                lehrer_kuerzel = None

        if note_part.startswith("-"):
            return None, None, ist_wahlpflicht, lehrer_kuerzel

        special_match = re.search(r"\b(GB|NF)\b", note_part, re.IGNORECASE)
        if special_match:
            return None, special_match.group(1).upper(), ist_wahlpflicht, lehrer_kuerzel
            
        decimal_match = re.search(r"(\d+\.?\d*)", note_part)
        if decimal_match:
            try:
                note_float = float(decimal_match.group(1))
                note = int(round(note_float))
                if 1 <= note <= 6:
                    return note, None, ist_wahlpflicht, lehrer_kuerzel
            except ValueError:
                pass
                
        digit_match = re.search(r"(\d)", note_part)
        if digit_match:
            try:
                note = int(digit_match.group(1))
                if 1 <= note <= 6:
                    return note, None, ist_wahlpflicht, lehrer_kuerzel
            except ValueError:
                pass
                
        return None, None, ist_wahlpflicht, lehrer_kuerzel

    def _extract_wahlpflicht_gruppe(self, fach_name: str) -> Tuple[str, Optional[str]]:
        """Extrahiert Wahlpflichtgruppe aus Fachnamen"""
        patterns = [
            (r"\(WPU1\)", "WPU1"),
            (r"\(WPU2\)", "WPU2"),
            (r"\(WPU\s*1\)", "WPU1"),
            (r"\(WPU\s*2\)", "WPU2"),
            (r"\(WP1\)", "WP1"),
            (r"\(WP2\)", "WP2"),
            (r"\(WP\)", "WP"),
            (r"\(W\)", "WP"),
            (r"^WP1\b", "WP1"),
            (r"^WP2\b", "WP2"),
            (r"^WPU1\b", "WPU1"),
            (r"^WPU2\b", "WPU2"),
            (r"^WP\b", "WP"),
            (r"^WPU\b", "WPU"),
            (r"^W\b", "WP"),
            (r"Praxistag", "Praxistag"),
            (r"\bWP1\b", "WP1"),
            (r"\bWP2\b", "WP2"),
            (r"\bWPU1\b", "WPU1"),
            (r"\bWPU2\b", "WPU2"),
        ]
        for pattern, gruppe in patterns:
            if re.search(pattern, fach_name):
                fach_clean = re.sub(pattern, "", fach_name).strip()
                return fach_clean, gruppe
        return fach_name, None

    def _get_or_create_fach(
        self,
        fach_kurz: str,
        fach_typ: str = None,
        ist_wahlpflicht: bool = False,
        wahlpflicht_gruppe: str = None,
    ) -> int:
        """Holt oder erstellt ein Fach und gibt die ID zurück"""
        cache_key = f"{fach_kurz}_{fach_typ}_{wahlpflicht_gruppe}"
        if cache_key in self.faecher_cache:
            return self.faecher_cache[cache_key]

        fach_clean, wp_gruppe = self._extract_wahlpflicht_gruppe(fach_kurz)
        if wp_gruppe:
            wahlpflicht_gruppe = wp_gruppe
            ist_wahlpflicht = True
        fach_kurz = fach_clean

        cursor = self.conn.execute(
            """SELECT fach_id FROM faecher
               WHERE fach_kurz = ?
               AND (fach_typ = ? OR (fach_typ IS NULL AND ? IS NULL))
               AND (wahlpflicht_gruppe = ? OR (wahlpflicht_gruppe IS NULL AND ? IS NULL))""",
            (fach_kurz, fach_typ, fach_typ, wahlpflicht_gruppe, wahlpflicht_gruppe),
        )
        result = cursor.fetchone()
        if result:
            fach_id = result[0]
        else:
            fach_lang = FAECHER_MAPPING.get(fach_kurz, fach_kurz)
            cursor = self.conn.execute(
                """INSERT INTO faecher (fach_kurz, fach_lang, fach_typ, ist_wahlpflicht, wahlpflicht_gruppe)
                   VALUES (?, ?, ?, ?, ?)""",
                (fach_kurz, fach_lang, fach_typ, ist_wahlpflicht, wahlpflicht_gruppe),
            )
            fach_id = cursor.lastrowid

        self.faecher_cache[cache_key] = fach_id
        return fach_id

    def _get_or_create_schueler(self, name: str, klasse: str) -> int:
        """Holt oder erstellt einen Schüler und gibt die ID zurück"""
        cursor = self.conn.execute(
            "SELECT schueler_id FROM schueler WHERE name = ? AND klasse = ?",
            (name, klasse),
        )
        result = cursor.fetchone()
        if result:
            return result[0]
        else:
            cursor = self.conn.execute(
                "INSERT INTO schueler (name, klasse) VALUES (?, ?)", (name, klasse)
            )
            return cursor.lastrowid

    def import_excel_file(self, file_path: str):
        """Importiert eine Excel-Datei"""
        file_path = Path(file_path)
        klasse = (
            file_path.stem.split("_")[-1] if "_" in file_path.stem else file_path.stem
        )
        # Jahrgang extrahieren (z.B. "05a" -> 5)
        jahrgang = None
        jahr_match = re.search(r"(\d+)", klasse)
        if jahr_match:
            jahrgang = int(jahr_match.group(1))

        self.logger.info(f"Importiere Datei: {file_path.name} (Klasse: {klasse}, Jahrgang: {jahrgang})")
        try:
            df = pd.read_excel(file_path, engine="openpyxl")
            if "Name" not in df.columns or "Art" not in df.columns:
                raise ValueError("Spalten 'Name' oder 'Art' nicht gefunden")

            meta_columns = ["Name", "Art", "KN", "Abstg."]
            fach_info = []
            for idx, col in enumerate(df.columns):
                if col not in meta_columns and pd.notna(col):
                    fach_info.append((idx, col))

            fach_columns_clean = []
            rel_count = 0
            
            
            for idx, col_name in fach_info:
                # Column names only contain subject, teacher is in the cell value
                fach_clean = str(col_name).strip()

                # SPH-Exportartefakt ignorieren:
                # In Jg. 9 enthalten manche Dateien eine Sammelspalte wie
                # "Al~Bio~Che~...~WPU", die KEIN echtes Fach ist.
                if fach_clean.count("~") >= 3:
                    self.logger.info(f"Ignoriere Sammelspalte ohne Fachbezug: {fach_clean}")
                    continue
                
                # BEREINIGUNG: Entferne Zusätze wie (U1), (U2), (U 1) etc.
                # Regex: Sucht nach (U, gefolgt von Leerzeichen/Zahlen, Klammer zu) am Ende oder mittendrin
                fach_clean = re.sub(r'\s*\(\s*U\s*\d+\s*\)', '', fach_clean, flags=re.IGNORECASE).strip()
                
                # TUT in 5/6 ignorieren
                if jahrgang in [5, 6] and fach_clean.upper() == "TUT":
                    continue

                # Typ ermitteln
                typ = None
                if fach_clean == "Re":
                    rel_count += 1
                    if rel_count == 1:
                        typ = "evangelisch"
                    else:
                        typ = "katholisch"
                
                fach_columns_clean.append((idx, fach_clean, typ))

            schueler_noten = {}
            for row_idx, row in df.iterrows():
                name = row.iloc[df.columns.get_loc("Name")]
                art = row.iloc[df.columns.get_loc("Art")]
                if pd.isna(name) or pd.isna(art):
                    continue

                if name not in schueler_noten:
                    schueler_noten[name] = {"AV": {}, "SV": {}}

                for col_idx, fach_kurz, fach_typ in fach_columns_clean:
                    note_raw = row.iloc[col_idx]
                    note, special_note, ist_wahlpflicht, lehrer_kuerzel = self._parse_note_mit_wahlpflicht(note_raw)
                    if note is not None or special_note is not None or ist_wahlpflicht:
                        schueler_noten[name][art][(fach_kurz, fach_typ)] = {
                            "note": note,
                            "special_note": special_note,
                            "ist_wahlpflicht": ist_wahlpflicht,
                            "lehrer_kuerzel": lehrer_kuerzel
                        }

            # --- Universelle Platzhalter-Logik ---
            # 1. Nicht-Klassenverbands-Fächer nach Jahrgang definieren
            non_class_band = {"Re", "Et", "Eth", "Ethik"}
            if jahrgang and jahrgang >= 7:
                non_class_band.update({"De", "Deu", "Deutsch", "Ma", "Mat", "Mathematik", "En", "Eng", "Englisch", "WPU1", "WP1"})
            if jahrgang and jahrgang >= 9:
                non_class_band.update({"WPU2", "WP2"})

            # 2. Alle Fächer und Lehrer-Kürzel sammeln
            klassen_faecher = {} # (kurz, typ) -> {kuerzel: count}
            for schueler_data in schueler_noten.values():
                for art in ["AV", "SV"]:
                    for (f_k, f_t), data in schueler_data[art].items():
                        if (f_k, f_t) not in klassen_faecher:
                            klassen_faecher[(f_k, f_t)] = {}
                        kuerzel = data.get("lehrer_kuerzel")
                        if kuerzel:
                            klassen_faecher[(f_k, f_t)][kuerzel] = klassen_faecher[(f_k, f_t)].get(kuerzel, 0) + 1
            
            # 3. Häufigstes Kürzel bestimmen (für Klassenverband-Zuweisung)
            default_lehrer = {}
            for f_key, counts in klassen_faecher.items():
                if counts:
                    default_lehrer[f_key] = max(counts, key=counts.get)
                else:
                    default_lehrer[f_key] = None
            
            # 3.b Aggressive Deduplizierung von Religion und Ethik BEVOR Platzhalter eingefügt werden
            for name, data in schueler_noten.items():
                # Wir prüfen AV und SV getrennt, aber meist sind keys identisch
                for art in ["AV", "SV"]:
                    entries = data[art]
                    # Gruppiere Keys nach Fach-Typ (Religion oder Ethik) -> NUTZE MAPPING für Aliase (Re, Et, etc.)
                    rel_keys = [k for k in entries.keys() if FAECHER_MAPPING.get(k[0], k[0]).startswith("Religion")]
                    eth_keys = [k for k in entries.keys() if FAECHER_MAPPING.get(k[0], k[0]).startswith("Ethik")]
                    all_rel_eth = rel_keys + eth_keys
                    
                    # Hilfsfunktion zum Bestimmen des besten Eintrags
                    def get_best_key(keys, entries_dict):
                        if not keys: return None
                        # 1. Bevorzuge Eintrag mit Note
                        for k in keys:
                            if entries_dict[k]["note"] is not None or entries_dict[k].get("special_note") is not None:
                                return k
                        # 2. Bevorzuge Eintrag mit Lehrerkürzel
                        for k in keys:
                            if entries_dict[k]["lehrer_kuerzel"]:
                                return k
                        # 3. Nimm den ersten
                        return keys[0]

                    # Check if ANY Rel/Eth has a grade
                    graded_keys = [k for k in all_rel_eth if entries[k]["note"] is not None or entries[k].get("special_note") is not None]
                    
                    if graded_keys:
                        # Wenn mindestens eine Note existiert, behalte NUR den besten benoteten Eintrag
                        # Alle anderen (auch vom anderen Typ) werden gelöscht
                        winner = get_best_key(graded_keys, entries)
                        for k in all_rel_eth:
                            if k != winner:
                                del entries[k]
                    else:
                        # Keine Note vorhanden -> Wir wollen MAXIMAL 1x Religion und 1x Ethik als Platzhalter
                        # Religion bereinigen
                        if len(rel_keys) > 1:
                            best_rel = get_best_key(rel_keys, entries)
                            for k in rel_keys:
                                if k != best_rel:
                                    del entries[k]
                        # Ethik bereinigen
                        if len(eth_keys) > 1:
                            best_eth = get_best_key(eth_keys, entries)
                            for k in eth_keys:
                                if k != best_eth:
                                    del entries[k]

                # 4. Fehlende Fächer ergänzen (Platzhalter)
                # ACHTUNG: Wir müssen sicherstellen, dass wir nicht doppelt hinzufügen
                for name_place, data_place in schueler_noten.items():
                    # Checke aktuelle Fächer basierend auf den BEREINIGTEN Listen
                    current_faecher_keys = set(data_place["AV"].keys()) | set(data_place["SV"].keys())
                    # Normalisiere auch hier für den Check
                    current_faecher_kurz = {FAECHER_MAPPING.get(k[0], k[0]) for k in current_faecher_keys}
                    
                    # Wenn weder Religion noch Ethik vorhanden sind, Platzhalter einfügen
                    # Checke auf startswith("Religion") um auch Varianten zu fangen
                    has_religion = any(f.startswith("Religion") for f in current_faecher_kurz)
                    has_ethik = any(f.startswith("Ethik") for f in current_faecher_kurz)

                    if not has_religion and not has_ethik:
                        # Platzhalter für Religion
                        rel_key = ("Religion", "evangelisch") 
                        if rel_key not in data_place["AV"]:
                            data_place["AV"][rel_key] = {"note": None, "special_note": None, "ist_wahlpflicht": False, "lehrer_kuerzel": None}
                            data_place["SV"][rel_key] = {"note": None, "special_note": None, "ist_wahlpflicht": False, "lehrer_kuerzel": None}
                        # Platzhalter für Ethik
                        eth_key = ("Ethik", None)
                        if eth_key not in data_place["AV"]:
                            data_place["AV"][eth_key] = {"note": None, "special_note": None, "ist_wahlpflicht": False, "lehrer_kuerzel": None}
                            data_place["SV"][eth_key] = {"note": None, "special_note": None, "ist_wahlpflicht": False, "lehrer_kuerzel": None}
                        
                        # Update local sets to prevent re-adding below
                        current_faecher_kurz.add("Religion")
                        current_faecher_kurz.add("Ethik")
                        current_faecher_keys.add(rel_key)
                        current_faecher_keys.add(eth_key)

                    for f_key in default_lehrer:
                        if f_key not in current_faecher_keys:
                            # Religion/Ethik Sondermodus: Nicht doppelt einfügen
                            norm_key = FAECHER_MAPPING.get(f_key[0], f_key[0])
                            if norm_key.startswith("Religion") or norm_key.startswith("Ethik"):
                                if has_religion or has_ethik or "Religion" in current_faecher_kurz or "Ethik" in current_faecher_kurz:
                                    continue
                            
                            lehrer = default_lehrer[f_key]
                            data_place["AV"][f_key] = {"note": None, "special_note": None, "ist_wahlpflicht": False, "lehrer_kuerzel": lehrer}
                            data_place["SV"][f_key] = {"note": None, "special_note": None, "ist_wahlpflicht": False, "lehrer_kuerzel": lehrer}
                            
                            # Update local set
                            current_faecher_keys.add(f_key)
                            current_faecher_kurz.add(norm_key)

                # Lehrer für bereits vorhandene Fächer ergänzen, falls dort fehlend
                for art in ["AV", "SV"]:
                    for f_key, f_data in data[art].items():
                        if not f_data.get("lehrer_kuerzel") and f_key in default_lehrer:
                            f_data["lehrer_kuerzel"] = default_lehrer[f_key]

            schueler_count = 0
            noten_count = 0
            for name, noten_data in schueler_noten.items():
                schueler_id = self._get_or_create_schueler(name, klasse)
                schueler_count += 1

                faecher_gesamt = set()
                faecher_gesamt.update(noten_data.get("AV", {}).keys())
                faecher_gesamt.update(noten_data.get("SV", {}).keys())

                for fach_kurz, fach_typ in faecher_gesamt:
                    fach_id = self._get_or_create_fach(fach_kurz, fach_typ)
                    
                    # Prüfen, ob das Fach global ein WP-Fach ist
                    cursor = self.conn.execute("SELECT ist_wahlpflicht FROM faecher WHERE fach_id = ?", (fach_id,))
                    fach_is_wp = bool(cursor.fetchone()[0])

                    av_data = noten_data.get("AV", {}).get((fach_kurz, fach_typ), {})
                    sv_data = noten_data.get("SV", {}).get((fach_kurz, fach_typ), {})
                    note_av = av_data.get("note")
                    note_sv = sv_data.get("note")
                    note_av_special = av_data.get("special_note")
                    note_sv_special = sv_data.get("special_note")
                    
                    # WP belegung ist wahr wenn entweder in der Zelle markiert ODER global am Fach
                    ist_wahlpflicht_belegung = (
                        av_data.get("ist_wahlpflicht", False) or 
                        sv_data.get("ist_wahlpflicht", False) or
                        fach_is_wp
                    )
                    # Teacher should be the same for both AV and SV, prefer AV if both exist
                    lehrer_kuerzel = av_data.get("lehrer_kuerzel") or sv_data.get("lehrer_kuerzel")


                    if (
                        note_av is not None
                        or note_sv is not None
                        or note_av_special is not None
                        or note_sv_special is not None
                        or ist_wahlpflicht_belegung
                        or True # Immer Platzhalter erlauben
                    ):
                        cursor = self.conn.execute(
                            """SELECT
                                   noten_id,
                                   note_av,
                                   note_sv,
                                   note_av_special,
                                   note_sv_special,
                                   COALESCE(manual_av_lock, 0),
                                   COALESCE(manual_sv_lock, 0)
                               FROM noten
                               WHERE schueler_id = ? AND fach_id = ?
                               AND schuljahr = ? AND halbjahr = ?""",
                            (schueler_id, fach_id, self.school_year, self.term),
                        )
                        existing = cursor.fetchone()
                        if existing:
                            existing_note_av = existing[1]
                            existing_note_sv = existing[2]
                            existing_note_av_special = existing[3]
                            existing_note_sv_special = existing[4]
                            manual_av_lock = bool(existing[5])
                            manual_sv_lock = bool(existing[6])
                            final_note_av = existing_note_av if manual_av_lock else note_av
                            final_note_sv = existing_note_sv if manual_sv_lock else note_sv
                            final_note_av_special = existing_note_av_special if manual_av_lock else note_av_special
                            final_note_sv_special = existing_note_sv_special if manual_sv_lock else note_sv_special
                            self.conn.execute(
                                """UPDATE noten
                                   SET note_av = ?, note_sv = ?,
                                       note_av_special = ?, note_sv_special = ?,
                                       ist_wahlpflicht_belegung = ?, lehrer_kuerzel = ?
                                   WHERE noten_id = ?""",
                                (
                                    final_note_av,
                                    final_note_sv,
                                    final_note_av_special,
                                    final_note_sv_special,
                                    ist_wahlpflicht_belegung,
                                    lehrer_kuerzel,
                                    existing[0],
                                ),
                            )
                        else:
                            self.conn.execute(
                                """INSERT INTO noten
                                   (schueler_id, fach_id, note_av, note_sv, note_av_special, note_sv_special, ist_wahlpflicht_belegung, lehrer_kuerzel, schuljahr, halbjahr)
                                   VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                                (
                                    schueler_id,
                                    fach_id,
                                    note_av,
                                    note_sv,
                                    note_av_special,
                                    note_sv_special,
                                    ist_wahlpflicht_belegung,
                                    lehrer_kuerzel,
                                    self.school_year,
                                    self.term,
                                ),
                            )
                        noten_count += 1

            self.conn.commit()
            self.logger.info(
                f"Verarbeitet: {schueler_count} Schüler mit {noten_count} Noteneinträgen"
            )
        except Exception as e:
            self.logger.error(f"Fehler beim Import von {file_path.name}: {str(e)}")
            self.conn.rollback()
            raise

    def _clean_existing_subjects(self):
        """Bereinigt nachträglich alle Fächer in der Datenbank von (U...)-Zusätzen."""
        self.logger.info("Starte Datenbank-Bereinigung für Fächer-Namen...")
        try:
            # 1. Alle Fächer laden
            cursor = self.conn.execute("SELECT fach_id, fach_kurz, fach_lang, fach_typ, wahlpflicht_gruppe FROM faecher")
            all_subjects = cursor.fetchall()
            
            changes_count = 0
            
            for subj in all_subjects:
                s_id = subj[0]
                s_kurz = subj[1]
                s_lang = subj[2]
                s_typ = subj[3]
                s_wp_gruppe = subj[4]

                # Entferne bekannte Sammel-/Artefaktfächer aus SPH-Dateien (z. B. "Al~Bio~...~WPU")
                if (s_kurz and s_kurz.count("~") >= 3) or (s_lang and s_lang.count("~") >= 3):
                    self.logger.info(f"Entferne Artefakt-Fach: ID {s_id} '{s_lang}'")
                    self.conn.execute("DELETE FROM noten WHERE fach_id = ?", (s_id,))
                    self.conn.execute("DELETE FROM faecher WHERE fach_id = ?", (s_id,))
                    changes_count += 1
                    continue

                # Check Regex
                clean_lang = re.sub(r'\s*\(\s*U\s*\d+\s*\)', '', s_lang, flags=re.IGNORECASE).strip()
                
                # Wenn Änderung nötig ist
                if clean_lang != s_lang:
                    clean_kurz = re.sub(r'\s*\(\s*U\s*\d+\s*\)', '', s_kurz, flags=re.IGNORECASE).strip()
                    
                    self.logger.info(f"Bereinige Fach: ID {s_id} '{s_lang}' -> '{clean_lang}'")
                    
                    # Prüfen: Gibt es das Clean-Fach schon?
                    # Achtung: Check muss auch Typ/Gruppe berücksichtigen, damit wir nicht Äpfel mit Birnen mergen.
                    # Aber in diesem Fall wollen wir ja gerade das "Praxistag (U1)" mit "Praxistag" mergen.
                    # Wir nehmen an, dass Typ/Gruppe identisch sein sollten oder vom Ziel übernommen werden.
                    
                    target_cursor = self.conn.execute(
                        """SELECT fach_id FROM faecher 
                           WHERE fach_lang = ? AND (fach_typ = ? OR (fach_typ IS NULL AND ? IS NULL))
                           AND (wahlpflicht_gruppe = ? OR (wahlpflicht_gruppe IS NULL AND ? IS NULL))""",
                        (clean_lang, s_typ, s_typ, s_wp_gruppe, s_wp_gruppe)
                    )
                    target_res = target_cursor.fetchone()
                    
                    if target_res:
                        # MERGE: Ziel existiert schon
                        target_id = target_res[0]
                        if target_id == s_id:
                            continue # Sollte nicht passieren, aber sicher ist sicher

                        self.logger.info(f"  -> Merge zu existierendem Fach ID {target_id}")

                        # 1. Update Noten: Setze fach_id auf target_id
                        # Konfliktlösung: IGNORE (falls Note für target_id im gleichen Halbjahr schon existiert -> behalte target_note)
                        self.conn.execute(
                            "UPDATE OR IGNORE noten SET fach_id = ? WHERE fach_id = ?",
                            (target_id, s_id)
                        )
                        
                        # 2. Lösche restliche Noten für s_id (die Duplikate wären und daher ignoriert wurden)
                        self.conn.execute("DELETE FROM noten WHERE fach_id = ?", (s_id,))
                        
                        # 3. Lösche das alte Fach
                        self.conn.execute("DELETE FROM faecher WHERE fach_id = ?", (s_id,))
                        
                    else:
                        # RENAME: Ziel existiert noch nicht -> einfach umbenennen
                        self.logger.info(f"  -> Umbenennung zu '{clean_lang}'")
                        self.conn.execute(
                            "UPDATE faecher SET fach_lang = ?, fach_kurz = ? WHERE fach_id = ?",
                            (clean_lang, clean_kurz, s_id)
                        )
                    
                    changes_count += 1
                    
            if changes_count > 0:
                self.conn.commit()
                self.logger.info(f"Bereinigung abgeschlossen. {changes_count} Fächer aktualisiert/gemerged.")
            else:
                self.logger.info("Keine bereinigungsbedürftigen Fächer gefunden.")
                
        except Exception as e:
            self.logger.error(f"Fehler bei der Datenbank-Bereinigung: {e}")
            self.conn.rollback() # Rollback safe

class OptimizedKopfnotenExporter:
    """Optimierter Exporter für horizontale 3-Zeilen-Tabellen mit korrekter erster Spalte"""
    def __init__(self, db_path: str, school_year: str = DEFAULT_SCHOOL_YEAR, term: int = DEFAULT_TERM):
        self.db_path = Path(db_path)
        self.conn = None
        self.logger = logging.getLogger("exporter")
        self.school_year = school_year
        self.term = int(term)
        if not self.db_path.exists():
            raise FileNotFoundError(f"Datenbank nicht gefunden: {self.db_path}")

    def __enter__(self):
        self.conn = sqlite3.connect(self.db_path)
        self.conn.row_factory = sqlite3.Row
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        if self.conn:
            self.conn.close()

    def _create_test_template(self, doc: Document, max_cols: int) -> None:
        """Creates a test template with dynamic columns per student
        
        Args:
            doc: Document object to add template content to
            max_cols: Maximum expected number of columns (used for display only)
        """
        # Import necessary modules for XML handling
        from docx.oxml.xmlchemy import OxmlElement
        from docx.oxml.ns import qn
        
        # Title and date (outside any control blocks)
        header_para = doc.add_paragraph()
        header_para.add_run("KOPFNOTEN - KLASSE {{ klasse }}")
        
        date_para = doc.add_paragraph()
        date_para.add_run("Export-Datum: {{ export_datum }}")
        
        doc.add_paragraph()  # Empty line for spacing

        # FOR loop opening - in its own paragraph
        for_para = doc.add_paragraph()
        for_para.add_run("{% for schueler in schueler_liste %}")

        # Student name heading
        name_para = doc.add_paragraph()
        name_para.add_run("{{ schueler.name }}")
        
        doc.add_paragraph()  # Add spacing before table
        
        # Create a placeholder table with just the first column
        # The real dynamic table will be created during template rendering
        table = doc.add_table(rows=3, cols=1)
        table.style = "Table Grid"
        table.autofit = False
        
        # Headers for first column
        headers = ["Fach", "AV", "SV"]
        for i, header in enumerate(headers):
            cell = table.cell(i, 0)
            run = cell.paragraphs[0].add_run(header)
            run.bold = True
        
        # Add Jinja code for dynamically adding the subject columns
        # This will insert the necessary columns during template rendering
        dynamic_code_para = doc.add_paragraph()
        dynamic_code_para.add_run("{{! ! ! DYNAMIC_TABLE_PLACEHOLDER ! ! !}}")
        
        # Add spacing after table
        doc.add_paragraph()
        
        # Add a paragraph for the conditional page break
        if_para = doc.add_paragraph()
        if_para.add_run("{% if not schueler.ist_letzter %}")
        
        # Add a page break using proper XML element creation
        page_break_para = doc.add_paragraph()
        run = page_break_para.add_run()
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run._r.append(br)
        
        # ENDIF statement - in its own paragraph
        endif_para = doc.add_paragraph()
        endif_para.add_run("{% endif %}")
        
        # Add spacing before endfor
        doc.add_paragraph()
        
        # ENDFOR statement - in its own paragraph
        endfor_para = doc.add_paragraph()
        endfor_para.add_run("{% endfor %}")

    def _extract_jahrgang(self, klasse: str) -> Optional[int]:
        """Extrahiert den Jahrgang aus dem Klassennamen (z.B. '7a' -> 7)"""
        match = re.search(r'(\d+)', str(klasse))
        return int(match.group(1)) if match else None

    def _format_faecher_logic(self, rows, jahrgang: Optional[int]) -> Tuple[List[str], List[str], List[str]]:
        """Zentrale Logik für Fächer-Filterung, Formatierung und Sortierung"""
        regular_subjects = []
        wp_subjects = []
        
        # Triaden-Status
        triad_grades = {} # name -> {av, sv}
        wpu_graded_list = [] # Liste der benoteten WPU-Fächer
        
        processed_rows = []
        for row in rows:
            row_dict = dict(row)
            original_fach_lang = row_dict["fach_lang"]
            fach_kurz = row_dict.get("fach_kurz")
            fach_typ = row_dict.get("fach_typ")
            
            # Canonical name lookup
            fach_lang = FAECHER_MAPPING.get(original_fach_lang, original_fach_lang)
            
            is_wp = bool(row_dict.get("ist_wahlpflicht_belegung", 0))
            wp_gruppe = row_dict.get("wahlpflicht_gruppe")
            
            av_val = row_dict["note_av"]
            sv_val = row_dict["note_sv"]
            av_special = row_dict.get("note_av_special")
            sv_special = row_dict.get("note_sv_special")
            av_note = av_special if av_special is not None else (str(av_val) if av_val is not None else "-")
            sv_note = sv_special if sv_special is not None else (str(sv_val) if sv_val is not None else "-")

            is_rel_triad = (fach_kurz == "Ethik") or (fach_kurz == "Religion" and fach_typ in ["evangelisch", "katholisch"])
            
            config_status = SUBJECT_STATUS_CONFIG.get(fach_lang, "")
            is_wpu_config = "WPU" in config_status
            is_wpu = is_wp or any(p in (wp_gruppe or "") for p in ["WPU", "WP"]) or is_wpu_config

            if is_rel_triad and (av_val is not None or sv_val is not None or av_special is not None or sv_special is not None):
                triad_grades[fach_lang] = {"av": av_note, "sv": sv_note}
            if is_wpu and (av_val is not None or sv_val is not None or av_special is not None or sv_special is not None):
                if fach_lang not in wpu_graded_list:
                    wpu_graded_list.append(fach_lang)

            processed_rows.append({
                "fach_lang": fach_lang,
                "fach_kurz": fach_kurz,
                "fach_typ": fach_typ,
                "is_wp": is_wp,
                "is_wpu": is_wpu,
                "is_rel_triad": is_rel_triad,
                "av_note": av_note,
                "sv_note": sv_note,
                "wp_gruppe": wp_gruppe
            })

        # Triaden-Sicherstellung
        triad_names = ["Ethik", "Religion (ev)", "Religion (kath)"]
        
        # WPU Limit ermitteln
        wpu_limit = 2 if jahrgang and jahrgang >= 9 else 1
        
        # Determine allowed WPUs (Top N graded ones)
        # Sort key could be added here if needed, currently implicit by encounter order or alphabetic?
        # Better sort alphabetically to be deterministic
        wpu_graded_list.sort()
        allowed_wpus = wpu_graded_list[:wpu_limit]
        
        final_rows = []
        for p_row in processed_rows:
            keep_row = True
            
            if p_row["is_rel_triad"]:
                # Logic logic for Triad (stays same: "/" if other has grade)
                has_any_triad_grade = bool(triad_grades)
                if has_any_triad_grade and p_row["fach_lang"] not in triad_grades:
                     p_row["av_note"] = "/"
                     p_row["sv_note"] = "/"
            
            elif p_row["is_wpu"]:
                # WPU-Ausschluss basierend auf Limit:
                # Wir behalten NUR die Fächer, die in `allowed_wpus` sind.
                # Alles andere (z.B. Ungradierte WPU Platzhalter oder überschüssige) FLIEGT RAUS. (nicht nur "/")
                # Außer wir haben GAR KEINE Noten (wpu_graded_list ist leer).
                
                if allowed_wpus:
                    if p_row["fach_lang"] not in allowed_wpus:
                        keep_row = False
                else:
                    # Wenn keine WPU Noten da sind -> Ungradierte placeholders rausnehmen
                    if p_row["av_note"] == "-" and p_row["sv_note"] == "-":
                        keep_row = False

            if keep_row:
                final_rows.append(p_row)
        
        # Output Lists preparation
        # Wir müssen die final_rows nun in regular_subjects und wp_subjects aufteilen für den Docx Export
        
        for p_row in final_rows:
            fach_lang = p_row["fach_lang"]
            av_note = p_row["av_note"]
            sv_note = p_row["sv_note"]
            
            # Skip if filtered (should be handled by final_rows, but check for safety)
            
            if p_row["is_wpu"]:
                # WPU-Logik nach Jahrgang
                if jahrgang in [5, 6]:
                     regular_subjects.append({"display": fach_lang, "av": av_note, "sv": sv_note})
                elif jahrgang in [7, 8]:
                    wp_subjects.append({
                        "display": f"{fach_lang} (WPU1)", # Exportiert mit Suffix (WPU1)
                        "av": av_note,
                        "sv": sv_note,
                        "sort_key": (1, fach_lang)
                    })
                elif jahrgang in [9, 10]:
                    # determine group
                    best_group = "WPU2" if (p_row["wp_gruppe"] and "2" in str(p_row["wp_gruppe"])) else "WPU1"
                    # IF we have 2, logic ensures they are both allowed.
                    # We need to ensure we don't map both to WPU1 if they are distinct.
                    # The `allowed_wpus` list maintains order.
                    # Simple heuristic: first in list is WPU1, second is WPU2.
                    
                    try:
                        idx = wpu_graded_list.index(fach_lang)
                        best_group = "WPU1" if idx == 0 else "WPU2"
                    except ValueError:
                        pass # keep default
                        
                    wp_subjects.append({
                        "display": f"{fach_lang} ({best_group})", # Suffix (WPU1) oder (WPU2)
                        "av": av_note,
                        "sv": sv_note,
                        "group": best_group, 
                        "sort_key": (1 if best_group == "WPU1" else 2, fach_lang)
                    })
                else:
                    regular_subjects.append({"display": fach_lang, "av": av_note, "sv": sv_note})
            else:
                regular_subjects.append({"display": fach_lang, "av": av_note, "sv": sv_note})

        # Sortierung
        priority = {
            "Deutsch": 1,
            "Mathematik": 2,
            "Englisch": 3,
            "Gesellschaftslehre": 4, # Korrigiert von Gesellschaftskunde
            "Ethik": 20,
            "Religion (ev)": 21,
            "Religion (kath)": 22
        }

        def regular_sort_key(item):
            name = item["display"]
            return (priority.get(name, 100), name)

        regular_subjects.sort(key=regular_sort_key)
        wp_subjects.sort(key=lambda x: x.get("sort_key", (10, x["display"])))

        all_entries = regular_subjects + wp_subjects
        
        return (
            [e["display"] for e in all_entries],
            [e["av"] for e in all_entries],
            [e["sv"] for e in all_entries]
        )


    def _process_template_with_context(self, template_path: Path, context: Dict[str, Any], output_file: Path) -> None:
        """Process a template with dynamic table creation for each student
        
        Args:
            template_path: Path to the template file
            context: Context data for template rendering
            output_file: Path where to save the output file
        """
        original_cwd = os.getcwd()
        temp_dir = None
        
        try:
            # Create temporary directory
            temp_dir = Path(tempfile.mkdtemp())
            self.logger.info(f"Created temp directory: {temp_dir}")
            
            # Copy template to temp directory
            temp_template = temp_dir / template_path.name
            shutil.copy2(template_path, temp_template)
            self.logger.info(f"Template copied to temp directory, size: {temp_template.stat().st_size} bytes")
            
            # Change to temp directory
            os.chdir(temp_dir)
            self.logger.info(f"Changed working directory to: {temp_dir}")
            
            # Modify the template to include dynamic table creation code and get buffer
            # Decoupled approach: python-docx modifies -> buffer -> docxtpl reads buffer
            template_buffer = self._prepare_dynamic_template(temp_template)
            
            # Generate template filename with timestamp to avoid conflicts
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            temp_docxtpl = f"temp_template_{timestamp}.docx"
            
            # Initialize DocxTemplate with the buffer
            template = DocxTemplate(template_buffer)
            self.logger.info(f"DocxTemplate initialized with memory buffer")
            
            # Render template with context
            template.render(context)
            
            # Change back to original directory before saving
            os.chdir(original_cwd)
            self.logger.info(f"Changed back to original directory (cleanup): {original_cwd}")
            
            # Save output
            template.save(str(output_file))
            self.logger.info(f"Output saved to: {output_file}")
            
        except Exception as e:
            self.logger.error(f"Error processing template: {e}")
            raise
            
        finally:
            # Clean up
            try:
                if original_cwd:
                    os.chdir(original_cwd)
                if temp_dir and temp_dir.exists():
                    shutil.rmtree(temp_dir)
                self.logger.info("Temporäre Dateien bereinigt")
            except Exception as cleanup_error:
                self.logger.warning(f"Warning during cleanup: {cleanup_error}")
                
    def _prepare_dynamic_template(self, template_path: Path) -> io.BytesIO:
        """Modifies the template file to include dynamic code and returns a BytesIO buffer
        
        Args:
            template_path: Path to the template file
            
        Returns:
            io.BytesIO: Buffer containing the modified docx file
        """
        try:
            from docx import Document
            from docx.shared import Inches
            import io
            
            # Read the template as a docx file using python-docx
            doc = Document(template_path)
            
            # Find paragraphs containing the placeholder
            placeholder_text = "{{! ! ! DYNAMIC_TABLE_PLACEHOLDER ! ! !}}"
            for p in doc.paragraphs:
                if placeholder_text in p.text:
                    # Replace the placeholder with Jinja2 code for dynamic table creation
                    p.text = ""
                    run = p.add_run("""
{%- set faecher_count = schueler.faecher_anzahl -%}
{%- if faecher_count > 0 -%}
    {# Create a new table with exact column count #}
    {%- set new_table = create_table(rows=3, cols=faecher_count+1) -%}
    {%- do new_table.style = 'Table Grid' -%}
    {%- do new_table.autofit = False -%}
    
    {# Set column headers in first column #}
    {%- do new_table.cell(0, 0).paragraphs[0].add_run('Fach').bold = True -%}
    {%- do new_table.cell(1, 0).paragraphs[0].add_run('AV').bold = True -%}
    {%- do new_table.cell(2, 0).paragraphs[0].add_run('SV').bold = True -%}
    
    {# Calculate column widths: fixed total width (8 inches) #}
    {%- do new_table.column(0).width = Inches(1) -%}
    {%- set data_col_width = Inches(7) / faecher_count -%}
    
    {# Create data columns and fill with content #}
    {%- for i in range(faecher_count) -%}
        {%- do new_table.column(i+1).width = data_col_width -%}
        {%- do new_table.cell(0, i+1).paragraphs[0].add_run(schueler.faecher_spalten[i]) -%}
        {%- do new_table.cell(1, i+1).paragraphs[0].add_run(schueler.av_noten[i]) -%}
        {%- do new_table.cell(2, i+1).paragraphs[0].add_run(schueler.sv_noten[i]) -%}
    {%- endfor -%}
{%- endif -%}
""")
                    break
            
            # Save the modified template to a BytesIO buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            self.logger.info(f"Template modified in-memory and saved to buffer")
            return buffer
            
        except Exception as e:
            self.logger.error(f"Error preparing dynamic template: {e}")
            raise

    def export_horizontal_tables(
        self, output_dir: Path, template_path: Path, klassen_liste: List[str], schueler_id: Optional[int] = None, export_date: Optional[str] = None
    ) -> Dict[str, Any]:
        """Exportiert horizontale 3-Zeilen-Tabellen für ausgewählte Klassen oder einen einzelnen Schüler"""
        output_dir = Path(output_dir)
        template_path = Path(template_path).resolve()
        if not template_path.exists():
            raise FileNotFoundError(f"Template nicht gefunden: {template_path}")

        if not export_date:
            export_date = datetime.now().strftime("%d.%m.%Y")

        # Make sure template exists
        if not template_path.exists():
            raise FileNotFoundError(f"Template nicht gefunden: {template_path}")

        output_dir.mkdir(parents=True, exist_ok=True)
        summary = {
            "gesamt_dateien": 0,
            "gesamt_fehler": 0,
            "klassen_details": {},
            "schueler_details": {},
            "start_time": datetime.now(),
            "export_mode": "horizontal_optimized",
        }

        try:
            # Wenn eine Schüler-ID angegeben ist, exportiere nur diesen Schüler
            if schueler_id is not None:
                self.logger.info(f"Exportiere einzelnen Schüler: ID {schueler_id}")
                # Hole Schülerdaten
                cursor = self.conn.execute(
                    "SELECT name, klasse FROM schueler WHERE schueler_id = ?",
                    (schueler_id,)
                )
                schueler_data = cursor.fetchone()
                if not schueler_data:
                    raise ValueError(f"Schüler mit ID {schueler_id} nicht gefunden")
                schueler_name = schueler_data["name"]
                klasse = schueler_data["klasse"]

                # Exportiere den einzelnen Schüler
                schueler_result = self._export_einzelschueler_horizontal(
                    schueler_id, schueler_name, klasse, output_dir, template_path, export_date
                )
                summary["schueler_details"][schueler_id] = schueler_result
                if schueler_result["datei_erstellt"]:
                    summary["gesamt_dateien"] += 1
                else:
                    summary["gesamt_fehler"] += 1
            # Sonst exportiere alle ausgewählten Klassen
            else:
                for klasse in klassen_liste:
                    self.logger.info(f"Exportiere Klasse horizontal: {klasse}")
                    klassen_result = self._export_klasse_horizontal_optimized(
                        klasse, output_dir, template_path, export_date
                    )
                    summary["klassen_details"][klasse] = klassen_result
                    if klassen_result["datei_erstellt"]:
                        summary["gesamt_dateien"] += 1
                    else:
                        summary["gesamt_fehler"] += 1
        except Exception as e:
            self.logger.error(f"Fehler beim Export: {e}")
            summary["gesamt_fehler"] += 1

        summary["end_time"] = datetime.now()
        summary["duration"] = summary["end_time"] - summary["start_time"]
        return summary

    def _export_klasse_horizontal_optimized(
        self, klasse: str, output_dir: Path, template_path: Path, export_date: str
    ) -> Dict[str, Any]:
        """Exportiert eine Klasse als horizontale Tabelle (optimiert)"""
        result = {
            "datei_erstellt": False,
            "output_file": None,
            "schueler_count": 0,
            "faecher_count": 0,
            "fehler": None,
        }

        try:
            # Get class data
            schueler_liste = self._get_schueler_horizontal_optimized(klasse)
            if not schueler_liste:
                raise ValueError(f"Keine Schüler in Klasse {klasse} gefunden")

            # Calculate maximum subjects
            max_faecher = max(s["faecher_anzahl"] for s in schueler_liste)

            # Create context
            context = {
                "klasse": klasse,
                "export_datum": export_date,
                "schueler_liste": schueler_liste,
                "max_faecher": max_faecher,
                "schueler": schueler_liste[0] if schueler_liste else None,
            }

            # Create output file path
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_file = output_dir.resolve() / f"Kopfnoten_{klasse}_horizontal_{timestamp}.docx"

            # Process template
            self._process_template_with_context(template_path, context, output_file)

            # Update result
            result.update({
                "datei_erstellt": True,
                "output_file": str(output_file),
                "schueler_count": len(schueler_liste),
                "faecher_count": max_faecher,
            })
            self.logger.info(f"Erfolgreich exportiert: {output_file.name}")
        except Exception as e:
            error_msg = f"Export error {klasse}: {str(e)}"
            self.logger.error(error_msg)
            result["fehler"] = error_msg

        return result

    def _export_einzelschueler_horizontal(
        self, schueler_id: int, schueler_name: str, klasse: str, output_dir: Path, template_path: Path, export_date: str
    ) -> Dict[str, Any]:
        """Exportiert einen einzelnen Schüler als horizontale Tabelle"""
        result = {
            "datei_erstellt": False,
            "output_file": None,
            "faecher_count": 0,
            "fehler": None,
        }

        try:
            # Get student data
            schueler_data = self._get_einzelschueler_horizontal(schueler_id, schueler_name, klasse)
            if not schueler_data:
                raise ValueError(f"No data found for student {schueler_name} (ID: {schueler_id})")

            # Create context
            context = {
                "klasse": klasse,
                "export_datum": export_date,
                "schueler_liste": [schueler_data],
                "max_faecher": schueler_data["faecher_anzahl"],
                "schueler": schueler_data,
            }

            # Create output file path
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_file = output_dir.resolve() / f"Kopfnoten_{schueler_name.replace(' ', '_')}_{timestamp}.docx"

            # Process template
            self._process_template_with_context(template_path, context, output_file)

            # Update result
            result.update({
                "datei_erstellt": True,
                "output_file": str(output_file),
                "faecher_count": schueler_data["faecher_anzahl"],
            })
            self.logger.info(f"Export successful: {output_file.name}")
        except Exception as e:
            error_msg = f"Export error for student {schueler_name}: {str(e)}"
            self.logger.error(error_msg)
            result["fehler"] = error_msg

        return result

    def _get_einzelschueler_horizontal(
        self, schueler_id: int, schueler_name: str, klasse: str
    ) -> Dict[str, Any]:
        """Sammelt Daten für einen einzelnen Schüler für horizontale Darstellung"""
        jahrgang = self._extract_jahrgang(klasse)

        # Fächer für diesen Schüler laden
        cursor = self.conn.execute(
            """
            SELECT
                f.fach_lang,
                f.fach_kurz,
                f.fach_typ,
                n.note_av,
                n.note_sv,
                n.note_av_special,
                n.note_sv_special,
                n.ist_wahlpflicht_belegung,
                f.wahlpflicht_gruppe
            FROM noten n
            JOIN faecher f ON n.fach_id = f.fach_id
            WHERE n.schueler_id = ?
              AND n.schuljahr = ?
              AND n.halbjahr = ?
            """,
            (schueler_id, self.school_year, self.term),
        )

        rows = cursor.fetchall()
        faecher_spalten, av_noten, sv_noten = self._format_faecher_logic(rows, jahrgang)

        schueler_data = {
            "name": schueler_name,
            "klasse": klasse,
            "faecher_spalten": faecher_spalten,
            "av_noten": av_noten,
            "sv_noten": sv_noten,
            "faecher_anzahl": len(faecher_spalten),
            "ist_letzter": True,
        }
        return schueler_data

    def _get_schueler_horizontal_optimized(self, klasse: str) -> List[Dict[str, Any]]:
        """Sammelt Schülerdaten für optimierte horizontale Darstellung"""
        jahrgang = self._extract_jahrgang(klasse)

        cursor = self.conn.execute(
            """
            SELECT DISTINCT s.schueler_id, s.name
            FROM schueler s
            JOIN noten n ON s.schueler_id = n.schueler_id
            WHERE s.klasse = ?
              AND COALESCE(s.is_active, 1) = 1
              AND n.schuljahr = ?
              AND n.halbjahr = ?
            ORDER BY name
            """,
            (klasse, self.school_year, self.term),
        )

        schueler_liste = []
        schueler_rows = cursor.fetchall()

        for i, schueler in enumerate(schueler_rows):
            schueler_id = schueler["schueler_id"]
            schueler_name = schueler["name"]

            # Fächer für diesen Schüler laden
            cursor = self.conn.execute(
                """
                SELECT
                    f.fach_lang,
                    f.fach_kurz,
                    f.fach_typ,
                    n.note_av,
                    n.note_sv,
                    n.note_av_special,
                    n.note_sv_special,
                    n.ist_wahlpflicht_belegung,
                    f.wahlpflicht_gruppe
                FROM noten n
                JOIN faecher f ON n.fach_id = f.fach_id
                WHERE n.schueler_id = ?
                  AND n.schuljahr = ?
                  AND n.halbjahr = ?
                """,
                (schueler_id, self.school_year, self.term),
            )

            rows = cursor.fetchall()
            faecher_spalten, av_noten, sv_noten = self._format_faecher_logic(rows, jahrgang)

            schueler_data = {
                "name": schueler_name,
                "klasse": klasse,
                "faecher_spalten": faecher_spalten,
                "av_noten": av_noten,
                "sv_noten": sv_noten,
                "faecher_anzahl": len(faecher_spalten),
                "ist_letzter": (i == len(schueler_rows) - 1),
            }
            schueler_liste.append(schueler_data)

        return schueler_liste

class SimplifiedGradeEditor:
    """Vereinfachter Noten-Editor"""
    def __init__(self, master_widget, db_path: str):
        self.master_widget = master_widget
        self.app = None # Reference to main application for callbacks
        self.db_path = db_path
        self.logger = logging.getLogger("grade_editor")

    def open_grade_editor(self, student_id: int, student_name: str, student_class: str, student_context_list: List[Dict] = None):
        """Öffnet den Editor für einen Schüler"""
        try:
            # Fenster erstellen (oder wiederverwenden falls wir Navigation machen?)
            # Wir erstellen es neu für Clean State, aber idealerweise Reuse.
            # Navigation Logic: Wenn das Fenster schon offen ist -> Reuse?
            # Einfachste Lösung: Fenster clean neu bauen.
            
            editor_window = tk.Toplevel(self.master_widget)
            editor_window.title(f"Noten bearbeiten: {student_name} ({student_class})")
            editor_window.geometry("900x700")
            
            # --- Navigation Frame (Top) ---
            if student_context_list:
                nav_frame = ttk.Frame(editor_window)
                nav_frame.pack(fill=tk.X, padx=10, pady=5)
                
                # Find current index
                current_idx = -1
                for i, s in enumerate(student_context_list):
                    if str(s["id"]) == str(student_id):
                        current_idx = i
                        break
                
                def navigate(delta):
                    new_idx = current_idx + delta
                    if 0 <= new_idx < len(student_context_list):
                        next_student = student_context_list[new_idx]
                        editor_window.destroy()
                        self.open_grade_editor(next_student["id"], next_student["name"], next_student["klasse"], student_context_list)
                
                # Buttons
                btn_prev = ttk.Button(nav_frame, text="<< Vorheriger", command=lambda: navigate(-1), state=tk.NORMAL if current_idx > 0 else tk.DISABLED)
                btn_prev.pack(side=tk.LEFT)
                
                ttk.Label(nav_frame, text=f"{current_idx+1} / {len(student_context_list)}").pack(side=tk.LEFT, padx=10)
                
                btn_next = ttk.Button(nav_frame, text="Nächster >>", command=lambda: navigate(1), state=tk.NORMAL if current_idx < len(student_context_list)-1 else tk.DISABLED)
                btn_next.pack(side=tk.RIGHT)

            
            
            # Header Frame
            header_frame = ttk.Frame(editor_window)
            header_frame.pack(fill=tk.X, padx=10, pady=5)

            tk.Label(
                header_frame,
                text=f"Noten für {student_name}",
                font=("Arial", 14, "bold"),
            ).pack(side=tk.LEFT)
            ttk.Label(
                header_frame, text=f"Klasse: {student_class}", font=("Arial", 12)
            ).pack(side=tk.RIGHT)

            # Noten-Frame
            notes_frame = ttk.LabelFrame(editor_window, text="Noten bearbeiten")
            notes_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

            # Lade aktuelle Noten
            grades_data = self._load_student_grades(student_id)
            
            # Jahrgang ermitteln
            match = re.search(r'(\d+)', student_class)
            jahrgang = int(match.group(1)) if match else 0

            # WPU Target Limit bestimmen
            wpu_target = 0
            if jahrgang in [7, 8]:
                wpu_target = 1
            elif jahrgang in [9, 10]:
                wpu_target = 2
            
            # --- BACKFILLING: Fehlende reguläre Fächer ergänzen ---
            # Ermittle alle regulären Fächer der Klasse (Präziser als Jahrgang)
            # Delegate to main app
            if self.app:
                regular_subjects_metadata = self.app._get_class_regular_subjects(student_class)
            else:
                 self.logger.warning("No app reference in Editor, cannot backfill regular subjects")
                 regular_subjects_metadata = {}
            
            # Welche Fächer hat der Schüler bereits?
            existing_subjects_set = set(g["fach_lang"] for g in grades_data)
            for subj_name, subj_id in regular_subjects_metadata.items():
                # Exclude Ethik/Religion from generic backfilling (handled by Triad logic)
                if subj_name in ["Ethik", "Religion", "Religion evangelisch", "Religion katholisch"]:
                    continue
                    
                if subj_name not in existing_subjects_set:
                    # Placeholder hinzufügen
                    # Wir brauchen eine Dummy-ID und Standardwerte
                    new_placeholder = {
                        "noten_id": f"missing_{subj_name}", # Virtuelle ID
                        "fach_id": subj_id, # Real ID from metadata!
                        "fach_lang": subj_name,
                        "fach_kurz": subj_name[:3].upper(), # Generiert
                        "note_av": None,
                        "note_sv": None,
                        "note_av_special": None,
                        "note_sv_special": None,
                        "ist_wahlpflicht_belegung": 0,
                        "ist_wahlpflicht": 0,
                        "wahlpflicht_gruppe": None,
                        "lehrer_kuerzel": "",
                        "is_placeholder": True
                    }
                    grades_data.append(new_placeholder)
                    existing_subjects_set.add(subj_name)
            # -----------------------------------------------------
            # Zähle bereits BENOTETE WPU-Fächer des Schülers
            graded_wpu_count = 0
            existing_wpu_subjects = set()
            
            for g in grades_data:
                # Prüfe ob WPU (Config oder Flag)
                status_check = SUBJECT_STATUS_CONFIG.get(g["fach_lang"], "")
                is_wpu_config = "WPU" in status_check
                is_wpu_db = bool(g["ist_wahlpflicht_belegung"]) or bool(g["ist_wahlpflicht"]) or (g["wahlpflicht_gruppe"] and "WP" in str(g["wahlpflicht_gruppe"]))
                
                is_wpu = is_wpu_config or is_wpu_db
                
                if is_wpu:
                    existing_wpu_subjects.add(g["fach_lang"])
                    # Zählt nur wenn Note vorhanden
                    if (
                        g["note_av"] is not None
                        or g["note_sv"] is not None
                        or g.get("note_av_special") is not None
                        or g.get("note_sv_special") is not None
                    ):
                        graded_wpu_count += 1
            
            # Wenn WPU-Ziel NOCH NICHT erreicht ist -> Platzhalter anzeigen
            if graded_wpu_count < wpu_target:
                # Lade mögliche WPU-Fächer der Klasse und merge sie ein
                if self.app:
                    class_wpu = self.app._get_class_wpu_subjects(student_class)
                else:
                    class_wpu = []
                
                for wpu in class_wpu:
                    # Prüfe auch via Config
                    status_check = SUBJECT_STATUS_CONFIG.get(wpu["fach_lang"], "")
                    is_wpu_config = "WPU" in status_check
                    is_wpu_db = wpu["ist_wahlpflicht"] or (wpu["wahlpflicht_gruppe"] and "WP" in wpu["wahlpflicht_gruppe"])
                    
                    # Wenn valides WPU Fach und noch nicht beim Schüler vorhanden
                    if (is_wpu_config or is_wpu_db) and wpu["fach_lang"] not in existing_wpu_subjects:
                        # Dummy Entry erstellen
                        dummy_entry = {
                            "noten_id": f"new_wpu_{wpu['fach_kurz']}", 
                            "fach_id": wpu["fach_id"], 
                            "fach_lang": wpu["fach_lang"],
                            "fach_kurz": wpu["fach_kurz"],
                            "note_av": None,
                            "note_sv": None,
                            "note_av_special": None,
                            "note_sv_special": None,
                            "ist_wahlpflicht_belegung": 0, 
                            "ist_wahlpflicht": 1,
                            "wahlpflicht_gruppe": wpu["wahlpflicht_gruppe"],
                            "lehrer_kuerzel": wpu["lehrer_kuerzel"], 
                            "is_placeholder": True 
                        }
                        grades_data.append(dummy_entry)
            else:
                # WPU-Ziel erreicht (oder überschritten):
                # Wir müssen UNBENOTETE WPU-Einträge aus der Liste entfernen, damit sie nicht angezeigt werden.
                # Das bereinigt die Ansicht von "Leichen" (importierte leere Spalten oder zuvor angelegte Platzhalter).
                
                cleaned_grades = []
                for g in grades_data:
                    # Prüfe ob WPU
                    status_check = SUBJECT_STATUS_CONFIG.get(g["fach_lang"], "")
                    is_wpu_config = "WPU" in status_check
                    is_wpu_db = bool(g["ist_wahlpflicht_belegung"]) or bool(g["ist_wahlpflicht"]) or (g["wahlpflicht_gruppe"] and "WP" in str(g["wahlpflicht_gruppe"]))
                    is_wpu = is_wpu_config or is_wpu_db
                    
                    # Behalten wenn:
                    # - KEIN WPU Fach (reguläres Fach)
                    # - ODER WPU Fach mit Note
                    # - ODER WPU Fach ist Teil der benoteten (graded_wpu_count zählt ja nur Noten, aber vielleicht will man das Fach behalten?)
                    #   Eigentlich: Wenn unbenotet und WPU -> Weg.
                    
                    has_grade = (
                        g["note_av"] is not None
                        or g["note_sv"] is not None
                        or g.get("note_av_special") is not None
                        or g.get("note_sv_special") is not None
                    )
                    
                    if not is_wpu:
                        cleaned_grades.append(g)
                    elif has_grade:
                         cleaned_grades.append(g)
                    # else: skip (hide)
                
                grades_data = cleaned_grades
            
            # Sortieren nach Fachname
            grades_data.sort(key=lambda x: x["fach_lang"])

            # Erstelle Eingabefelder
            self._create_grade_inputs(
                notes_frame, grades_data, student_id, student_class, editor_window
            )
        except Exception as e:
            self.logger.error(f"Fehler beim Öffnen des Noten-Editors: {e}")
            messagebox.showerror("Editor-Fehler", f"Fehler: {e}")

    def _load_student_grades(self, student_id: int) -> List[Dict]:
        """Lädt Noten eines Schülers (korrigierte Version)"""
        try:
            school_year, term = (DEFAULT_SCHOOL_YEAR, DEFAULT_TERM)
            if self.app and hasattr(self.app, "_get_active_period"):
                school_year, term = self.app._get_active_period()
            with sqlite3.connect(self.db_path) as conn:
                conn.row_factory = sqlite3.Row # Für Spaltenzugriff per Name
                cursor = conn.execute(
                    """
                    SELECT
                        n.noten_id,
                        f.fach_lang,
                        f.fach_kurz,
                        n.note_av,
                        n.note_sv,
                        n.note_av_special,
                        n.note_sv_special,
                        COALESCE(n.manual_av_lock, 0) AS manual_av_lock,
                        COALESCE(n.manual_sv_lock, 0) AS manual_sv_lock,
                        n.ist_wahlpflicht_belegung,
                        f.ist_wahlpflicht,
                        f.wahlpflicht_gruppe,
                        n.lehrer_kuerzel
                    FROM noten n
                    JOIN faecher f ON n.fach_id = f.fach_id
                    WHERE n.schueler_id = ?
                      AND n.schuljahr = ?
                      AND n.halbjahr = ?
                    ORDER BY f.fach_lang
                    """,
                    (student_id, school_year, term),
                )
                return [dict(row) for row in cursor.fetchall()]
        except Exception as e:
            self.logger.error(f"Fehler beim Laden der Noten: {e}")
            return []



    def _create_grade_inputs(
        self, parent_frame, grades_data: List[Dict], student_id: int, student_class: str, editor_window
    ):
        """Erstellt Eingabefelder für Noten"""
        # Scrollbares Frame Setup
        canvas = tk.Canvas(parent_frame)
        scrollbar = ttk.Scrollbar(parent_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        # Header für die Eingabefelder
        headers = ["Fach", "Lehrkraft", "AV-Note (1-6)", "SV-Note (1-6)", "Status"]
        for i, header in enumerate(headers):
            ttk.Label(scrollable_frame, text=header, font=("Arial", 10, "bold")).grid(
                row=0, column=i, padx=10, pady=5, sticky=tk.W
            )

        # Daten einfügen
        grade_vars = {}
        triad_vars = []
        wpu_vars = []
        
        
        # Helper for Subject Status
        for row_idx, grade in enumerate(grades_data, start=1):
            # Fachname bereinigen (WP-Kürzel entfernen, U1/U2 entfernen)
            fach_name = grade["fach_lang"]
            fach_name = re.sub(r"\s*\(WP\d*\)\s*$", "", fach_name)
            # Robustere Regex für (U1), (U 12) etc.
            fach_name = re.sub(r'\s*\(\s*U\s*\d+\s*\)', '', fach_name, flags=re.IGNORECASE).strip()
            fach_name = fach_name.strip()
            
            # Lehrer (nur Info)
            lehrer = grade.get("lehrer_kuerzel") or "-"
            
            # IDs handlen (New vs Existing)
            noten_id = grade["noten_id"]
            is_placeholder = grade.get("is_placeholder", False)
            fach_kurz = grade.get("fach_kurz")
            fach_typ = grade.get("fach_typ")
            
            # ... rest logic ...
            if grade["wahlpflicht_gruppe"]:
                # Auch hier WP-Bezüge in Klammern entfernen, falls vorhanden
                wp_grp = re.sub(r"WP\d*", "", str(grade['wahlpflicht_gruppe'])).strip(" ()")
                if wp_grp:
                    fach_name += f" ({wp_grp})"
            
            # WPU detection
            is_wp = bool(grade["ist_wahlpflicht_belegung"]) or bool(grade["ist_wahlpflicht"]) or any(p in (grade["wahlpflicht_gruppe"] or "") for p in ["WPU", "WP"])

            # Subject Status Text
            # Nutze FAECHER_MAPPING um den kanonischen Namen zu prüfen
            # grade["fach_lang"] ist der Name aus der DB (evtl importiert)
            # Wir nutzen den cleanen Namen für den Check
            clean_name = re.sub(r"\s*\(WP\d*\)\s*$", "", grade["fach_lang"])
            canonical_name = FAECHER_MAPPING.get(clean_name, clean_name)
            
            status_text = "Nebenfach"
            
            # Lookup in SUBJECT_STATUS_CONFIG
            if canonical_name in SUBJECT_STATUS_CONFIG:
                 status_text = SUBJECT_STATUS_CONFIG[canonical_name]
            # Fallback für Namen die nicht im Mapping sind, aber vielleicht "Hauptfach" sein sollen?
            # Für jetzt: Default "Nebenfach" ist OK.
            
            if is_wp and "WPU" not in status_text:
                status_text += " (WPU)"

            # Fach Name
            ttk.Label(scrollable_frame, text=fach_name).grid(row=row_idx, column=0, padx=10, pady=2, sticky=tk.W)
            
            # Lehrer (nur Info)
            lehrer = grade.get("lehrer_kuerzel") or "-"
            ttk.Label(scrollable_frame, text=lehrer).grid(row=row_idx, column=1, padx=10, pady=2, sticky=tk.W)
            
            grade_options = ["", "1", "2", "3", "4", "5", "6", "GB", "NF"]

            # AV Note
            av_display = grade.get("note_av_special") or (str(grade["note_av"]) if grade["note_av"] is not None else "")
            av_var = tk.StringVar(value=av_display)
            av_entry = ttk.Combobox(
                scrollable_frame, textvariable=av_var, values=grade_options, width=6, state="readonly"
            )
            av_entry.grid(row=row_idx, column=2, padx=10, pady=2)
            
            # SV Note
            sv_display = grade.get("note_sv_special") or (str(grade["note_sv"]) if grade["note_sv"] is not None else "")
            sv_var = tk.StringVar(value=sv_display)
            sv_entry = ttk.Combobox(
                scrollable_frame, textvariable=sv_var, values=grade_options, width=6, state="readonly"
            )
            sv_entry.grid(row=row_idx, column=3, padx=10, pady=2)
            
            # Status Label instead of Checkbox
            ttk.Label(scrollable_frame, text=status_text).grid(row=row_idx, column=4, padx=10, pady=2, sticky=tk.W)

            vars_data = {
                "noten_id": grade["noten_id"],
                "av_var": av_var,
                "sv_var": sv_var,
                "av_entry": av_entry,
                "sv_entry": sv_entry,
                "fach_name": fach_name,
                "manual_av_lock": bool(grade.get("manual_av_lock", 0)),
                "manual_sv_lock": bool(grade.get("manual_sv_lock", 0)),

                "fach_kurz": fach_kurz,
                "fach_typ": fach_typ,
                "fach_id": grade.get("fach_id"), # Optional, für New WPU
                "ist_wpu": is_wp,
                "lehrer_kuerzel_hint": grade.get("lehrer_kuerzel"), # Für Insert benötigt
            }
            grade_vars[grade["noten_id"]] = vars_data
            
            # Zu Gruppen hinzufügen
            if (fach_kurz == "Ethik") or (fach_kurz == "Religion" and fach_typ in ["evangelisch", "katholisch"]):
                triad_vars.append(vars_data)
            elif vars_data["ist_wpu"]:
                wpu_vars.append(vars_data)

        # Traces für gegenseitigen Ausschluss
        def update_exclusion_triad(group_vars):
            # Religion/Ethik: Max 1 Note (AV oder SV) insgesamt in der Gruppe
            any_filled = False
            filled_id = None
            for v in group_vars:
                if v["av_var"].get().strip() or v["sv_var"].get().strip():
                    any_filled = True
                    filled_id = v["noten_id"]
                    # Sobald einer befüllt ist, brauchen wir nicht weiter suchen
                    break
            
            for v in group_vars:
                # Disable if limit reached AND this one is not yet filled
                state = "disabled" if (any_filled and v["noten_id"] != filled_id) else "readonly"
                v["av_entry"].config(state=state)
                v["sv_entry"].config(state=state)

        def update_exclusion_wpu(group_vars, student_class):
            # WPU: 7-8 -> 1 Note, 9-10 -> 2 Noten
            match_year = re.search(r"(\d+)", student_class)
            jahrgang = int(match_year.group(1)) if match_year else 0
            limit = 2 if jahrgang >= 9 else 1
            
            filled_vars = [v for v in group_vars if v["av_var"].get().strip() or v["sv_var"].get().strip()]
            
            for v in group_vars:
                # Disable if limit reached AND this one is not yet filled
                is_filled = v in filled_vars
                if len(filled_vars) >= limit and not is_filled:
                    state = "disabled"
                else:
                    state = "readonly"
                v["av_entry"].config(state=state)
                v["sv_entry"].config(state=state)

        # Traces aufsetzen
        if triad_vars:
            callback_triad = lambda *args: update_exclusion_triad(triad_vars)
            for v in triad_vars:
                v["av_var"].trace_add("write", callback_triad)
                v["sv_var"].trace_add("write", callback_triad)
            update_exclusion_triad(triad_vars)

        if wpu_vars:
            callback_wpu = lambda *args: update_exclusion_wpu(wpu_vars, student_class)
            for v in wpu_vars:
                v["av_var"].trace_add("write", callback_wpu)
                v["sv_var"].trace_add("write", callback_wpu)
            update_exclusion_wpu(wpu_vars, student_class)

        # Buttons (im editor_window, nicht im scrollable_frame)
        button_frame = ttk.Frame(editor_window)
        button_frame.pack(fill=tk.X, padx=10, pady=10)

        def save_all_grades():
            """Speichert alle Noten"""
            try:
                school_year, term = (DEFAULT_SCHOOL_YEAR, DEFAULT_TERM)
                if self.app and hasattr(self.app, "_get_active_period"):
                    school_year, term = self.app._get_active_period()
                with sqlite3.connect(self.db_path) as conn:
                    saved_count = 0
                    for noten_id, vars_dict in grade_vars.items():
                        av_text = vars_dict["av_var"].get().strip()
                        sv_text = vars_dict["sv_var"].get().strip()

                        # Validiere und konvertiere Noten
                        av_value = None
                        sv_value = None
                        av_special = None
                        sv_special = None
                        if av_text:
                            if av_text.upper() in {"GB", "NF"}:
                                av_special = av_text.upper()
                            else:
                                try:
                                    av_value = int(av_text)
                                    if not (1 <= av_value <= 6):
                                        raise ValueError(
                                            f"AV-Note für {vars_dict['fach_name']} muss zwischen 1 und 6 liegen"
                                        )
                                except ValueError:
                                    raise ValueError(f"Ungültige AV-Note für {vars_dict['fach_name']}")
                                
                        if sv_text:
                            if sv_text.upper() in {"GB", "NF"}:
                                sv_special = sv_text.upper()
                            else:
                                try:
                                    sv_value = int(sv_text)
                                    if not (1 <= sv_value <= 6):
                                        raise ValueError(
                                            f"SV-Note für {vars_dict['fach_name']} muss zwischen 1 und 6 liegen"
                                        )
                                except ValueError:
                                    raise ValueError(f"Ungültige SV-Note für {vars_dict['fach_name']}")

                        # WPU-Status automatisch setzen:
                        # Wenn Fach WPU-fähig ist UND eine Note hat, ist es gewählt.
                        wp_value = vars_dict.get("ist_wpu", False) and (
                            av_value is not None or sv_value is not None or av_special is not None or sv_special is not None
                        )

                        # Speichere in Datenbank
                        if str(noten_id).startswith("new_wpu_") or str(noten_id).startswith("missing_"):
                             # INSERT Logic für neue WPU Platzhalter UND Missing Regular Placeholder
                             # ACHTUNG: Wir brauchen die fach_id !
                             # Das ist etwas tricky, da wir die hier nicht haben direkt.
                             # Aber wir haben fach_kurz und fach_lang.
                             
                             # Nur speichern, wenn auch wirklich was eingetragen wurde!
                             if av_value is None and sv_value is None and av_special is None and sv_special is None:
                                 continue
                             
                             fach_id_target = vars_dict.get("fach_id")
                             if fach_id_target:
                                 av_manual_lock = 1 if av_text else 0
                                 sv_manual_lock = 1 if sv_text else 0
                                 conn.execute(
                                    """
                                    INSERT INTO noten (
                                        schueler_id, fach_id, note_av, note_sv,
                                        note_av_special, note_sv_special,
                                        manual_av_lock, manual_sv_lock,
                                        ist_wahlpflicht_belegung, lehrer_kuerzel, schuljahr, halbjahr
                                    )
                                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                                    """,
                                    (
                                        student_id, fach_id_target, av_value, sv_value,
                                        av_special, sv_special,
                                        av_manual_lock, sv_manual_lock,
                                        wp_value, vars_dict.get("lehrer_kuerzel_hint"), school_year, term
                                    )
                                 )
                                 saved_count += 1
                             else:
                                 # Fallback falls fach_id fehlt (sollte nicht passieren)
                                 self.logger.error(f"Fehler: Keine fach_id für neuen WPU Eintrag {vars_dict['fach_name']}")
                        else:
                            # Normal UPDATE
                            av_manual_lock = 1 if av_text else 0
                            sv_manual_lock = 1 if sv_text else 0
                            conn.execute(
                                """
                                UPDATE noten
                                SET note_av = ?, note_sv = ?,
                                    note_av_special = ?, note_sv_special = ?,
                                    ist_wahlpflicht_belegung = ?,
                                    manual_av_lock = ?,
                                    manual_sv_lock = ?
                                WHERE noten_id = ?
                                  AND schuljahr = ?
                                  AND halbjahr = ?
                                """,
                                (
                                    av_value, sv_value, av_special, sv_special, wp_value,
                                    av_manual_lock, sv_manual_lock,
                                    noten_id, school_year, term
                                ),
                            )
                            saved_count += 1

                    conn.commit()
                    messagebox.showinfo(
                        "Gespeichert", f"{saved_count} Fachnoten erfolgreich aktualisiert!"
                    )
                    editor_window.destroy()

                    # Refresh parent window
                    if self.app and hasattr(self.app, "refresh_analysis_data"):
                        self.app.refresh_analysis_data()
                        # Also refresh tree selection to keep context if possible?
                        # Re-selecting might be tricky if list rebuilt.
                        # Maybe we can just stay silent.
            except ValueError as e:
                messagebox.showerror("Eingabefehler", str(e))
            except Exception as e:
                self.logger.error(f"Fehler beim Speichern der Noten: {e}")
                messagebox.showerror("Speicher-Fehler", f"Fehler: {e}")

        ttk.Button(
            button_frame, text="✅ Alle Änderungen speichern", command=save_all_grades
        ).pack(side=tk.LEFT, padx=5)

        ttk.Button(button_frame, text="Abbrechen", command=editor_window.destroy).pack(
            side=tk.RIGHT, padx=5
        )

class StatusManager:
    """Einfache Status-Verwaltung"""
    def __init__(self, parent):
        self.parent = parent
        self.status_label = None
        self.progress_bar = None

    def setup_status_bar(self, parent_frame):
        """Erstellt Statusleiste"""
        status_frame = ttk.Frame(parent_frame)
        status_frame.pack(side=tk.BOTTOM, fill=tk.X, padx=5, pady=2)
        self.status_label = ttk.Label(status_frame, text="Bereit", relief=tk.SUNKEN)
        self.status_label.pack(side=tk.LEFT, fill=tk.X, expand=True)
        self.progress_bar = ttk.Progressbar(status_frame, mode="indeterminate")
        self.progress_bar.pack(side=tk.RIGHT, padx=(5, 0))
        return status_frame

    def set_status(self, message: str, progress: bool = False):
        """Setzt Status"""
        if self.status_label:
            self.status_label.config(
                text=f"{datetime.now().strftime('%H:%M:%S')} - {message}"
            )
        if progress and self.progress_bar:
            self.progress_bar.start()
        elif not progress and self.progress_bar:
            self.progress_bar.stop()

    def clear_status(self):
        """Setzt Status zurück"""
        self.set_status("Bereit", False)

class KopfnotenGUI:
    """Hauptklasse der optimierten GUI-Anwendung"""
    def __init__(self):
        self.root = tk.Tk()
        self.paths = APP_PATHS
        # Pfade (Muss vor setup_application initialisiert sein!)
        self.db_path = self.paths.database_path
        
        self.setup_application()
        # Manager
        self.path_manager = LinuxPathManager()
        self.status_manager = StatusManager(self)
        self.template_designer = SimpleTemplateDesigner(self.root)
        # GUI-Variablen
        self.template_var = tk.StringVar()
        self.output_var = tk.StringVar(value=str(self.paths.output_word_dir))
        self.export_running = False
        self.ui_queue = queue.Queue()

        # --- LOGIN CHECK ---
        from credentials import CredentialManager
        from login_gui import LoginWindow
        
        self.credentials_manager = CredentialManager(data_dir=str(self.paths.temp_dir))
        
        # Hide main window during login
        self.root.withdraw()
        
        login_win = LoginWindow(self.root, self.credentials_manager)
        self.root.wait_window(login_win.window)
        
        if not login_win.result:
            try:
                if self.root.winfo_exists():
                    self.root.destroy()
            except:
                pass
            sys.exit(0)
            
        self.root.deiconify()
        # -------------------

        # GUI-Komponenten
        self.notebook = None
        self.import_listbox = None
        self.export_listbox = None
        self.export_log = None
        self.analysis_tree = None
        self.stats_text = None
        self.selected_schueler_var = tk.StringVar(value="")
        self.student_search_after = None # For debouncing
        self.sph_missing_overview = {}
        
        # New Filter Vars
        self.teacher_filter_var = tk.StringVar()
        self.status_filter_var = tk.StringVar(value="Alle")
        self.current_school_year_var = tk.StringVar(value=DEFAULT_SCHOOL_YEAR)
        self.current_term_var = tk.IntVar(value=DEFAULT_TERM)

        # Setup
        self.create_gui()
        self.load_sph_missing_overview()
        self.load_initial_data()
        self.setup_linux_environment()
        self.process_ui_queue()

    def process_ui_queue(self):
        """Processes UI updates from background threads"""
        try:
            while True:
                task = self.ui_queue.get_nowait()
                if callable(task):
                    task()
                self.ui_queue.task_done()
        except queue.Empty:
            pass
        self.root.after(100, self.process_ui_queue)

    def queue_ui(self, func: Callable, *args, **kwargs):
        """Queues a UI update function to be run on the main thread"""
        self.ui_queue.put(lambda: func(*args, **kwargs))

    def setup_application(self):
        """Grundlegende Anwendungseinrichtung"""
        self.root.title("Kopfnoten-Manager")
        self.root.geometry("1200x900")
        self.root.minsize(1000, 700)
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)
        
        # AUTO-CLEANUP ON START: Bereinige Datenbank von alten (U...)-Leichen
        try:
            with KopfnotenImporter(self.db_path) as importer:
                importer._clean_existing_subjects()
        except Exception as e:
            logging.error(f"Fehler beim Startup-Cleanup: {e}")

        # Style für bessere Optik
        style = ttk.Style()
        if "clam" in style.theme_names():
            style.theme_use("clam")
        style.configure("TNotebook.Tab", padding=(14, 8), font=("Segoe UI", 10))
        style.configure("Treeview", rowheight=30, font=("Segoe UI", 10))
        style.configure("Treeview.Heading", font=("Segoe UI Semibold", 10))
        style.configure("Card.TFrame", relief=tk.GROOVE, borderwidth=1)
        style.configure("CardTitle.TLabel", font=("Segoe UI", 9), foreground="#4b5563")
        style.configure("CardValue.TLabel", font=("Segoe UI Semibold", 14))

    def setup_linux_environment(self):
        """Linux-spezifische Umgebungseinrichtung"""
        directories = [
            self.paths.logs_dir,
            self.paths.templates_dir,
            self.paths.output_word_dir,
            self.paths.output_excel_dir,
            self.paths.database_path.parent,
            self.paths.temp_dir,
            self.paths.backup_dir,
        ]
        for directory in directories:
            try:
                self.path_manager.ensure_directory(directory)
                logging.info(f"Verzeichnis bereit: {directory}")
            except Exception as e:
                logging.error(f"Fehler beim Erstellen von {directory}: {e}")

    def create_gui(self):
        """Erstellt die GUI"""
        # Menü
        self.create_menu()
        # Hauptcontainer
        main_frame = ttk.Frame(self.root)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        # Status-Bar
        self.status_manager.setup_status_bar(self.root)
        # Notebook für Tabs
        self.notebook = ttk.Notebook(main_frame)
        self.notebook.pack(fill=tk.BOTH, expand=True)
        # Tabs erstellen
        self.create_import_tab()
        # self.create_sph_tab()  # SPH jetzt in Import
        self.create_analysis_tab()
        self.create_insights_tab()
        self.create_export_tab()
        if TEMPLATE_MANAGER_ENABLED:
            self.create_template_tab()
        
        # Load Config
        self.load_sph_config()
        
        # Initial Status
        self.status_manager.set_status("Anwendung gestartet")

    def _get_active_period(self) -> Tuple[str, int]:
        return (self.current_school_year_var.get().strip() or DEFAULT_SCHOOL_YEAR, int(self.current_term_var.get() or 1))

    def _get_active_period_key(self) -> str:
        school_year, term = self._get_active_period()
        return f"{school_year}|H{term}"

    def _get_sph_cache_path(self) -> Path:
        return self.paths.data_root / "sph_missing_overview.json"

    def _suggest_school_years(self) -> List[str]:
        now = datetime.now().year
        return [f"{y}/{y + 1}" for y in range(now - 2, now + 3)]

    def _refresh_period_label(self):
        if hasattr(self, "period_info_label"):
            school_year, term = self._get_active_period()
            self.period_info_label.config(text=f"Aktive Periode: {school_year} (HJ {term})")

    def _school_year_for_date(self, dt: datetime) -> str:
        start_year = dt.year if dt.month >= 8 else dt.year - 1
        return f"{start_year}/{start_year + 1}"

    def _suggest_period_from_date(self, dt: Optional[datetime] = None) -> Dict[str, Any]:
        dt = dt or datetime.now()
        month = dt.month
        school_year = self._school_year_for_date(dt)

        if month == 1:
            return {
                "school_year": school_year,
                "term": 1,
                "confident": True,
                "reason": "Januar-Fenster (Halbjahreszeugnisse)",
            }
        if month in (6, 7):
            return {
                "school_year": school_year,
                "term": 2,
                "confident": True,
                "reason": "Juni/Juli-Fenster (Ganzjahreszeugnisse)",
            }

        # Fallback außerhalb der klaren Erfassungsfenster.
        fallback_term = 1 if month <= 4 else 2
        return {
            "school_year": school_year,
            "term": fallback_term,
            "confident": False,
            "reason": "Außerhalb Januar/Juni/Juli (bitte bestätigen)",
        }

    def _apply_period(self, school_year: str, term: int):
        self.current_school_year_var.set(str(school_year).strip() or DEFAULT_SCHOOL_YEAR)
        self.current_term_var.set(1 if int(term) == 1 else 2)
        self.on_period_changed()

    def _confirm_period_before_import(self, import_source: str) -> bool:
        current_school_year, current_term = self._get_active_period()
        suggestion = self._suggest_period_from_date()
        suggested_school_year = suggestion["school_year"]
        suggested_term = suggestion["term"]
        confident = suggestion["confident"]

        current_label = f"{current_school_year} / HJ {current_term}"
        suggested_label = f"{suggested_school_year} / HJ {suggested_term}"

        if confident:
            if (current_school_year, current_term) != (suggested_school_year, suggested_term):
                use_suggestion = messagebox.askyesno(
                    "Periode prüfen",
                    f"{import_source}: Empfohlene Periode ist {suggested_label}\n"
                    f"(Grund: {suggestion['reason']}).\n\n"
                    f"Aktuell eingestellt: {current_label}\n\n"
                    "Soll auf die empfohlene Periode umgestellt werden?",
                )
                if use_suggestion:
                    self._apply_period(suggested_school_year, suggested_term)
            return True

        decision = messagebox.askyesnocancel(
            "Periode bestätigen",
            f"{import_source}: Zeitraum liegt außerhalb der Standardfenster.\n\n"
            f"Vorschlag: {suggested_label}\n"
            f"Aktuell:   {current_label}\n\n"
            "Ja = Vorschlag übernehmen\n"
            "Nein = aktuelle Periode beibehalten\n"
            "Abbrechen = Import stoppen",
        )
        if decision is None:
            return False
        if decision:
            self._apply_period(suggested_school_year, suggested_term)
        return True

    def on_period_changed(self, event=None):
        self.save_sph_config()
        self.load_sph_missing_overview()
        self._refresh_period_label()
        self.refresh_all_data()

    # Old SPH Tab removed


    # School Search Methods Removed (Handled in Login Window)
    def search_schools(self): pass
    def apply_school_selection(self, event=None): pass
    def add_recent_school(self, school_data): pass

    def load_sph_config(self):
        """Lädt SPH Konfiguration (Klassen + aktive Periode)."""
        try:
            import json
            config_path = self.paths.sph_config_path
            if config_path.exists():
                with open(config_path, "r") as f:
                    config = json.load(f)
                
                # School/User logic moved to credentials.py / LoginWindow

                period = config.get("period", {})
                school_year = str(period.get("school_year", "")).strip()
                term = period.get("term", DEFAULT_TERM)
                if school_year:
                    self.current_school_year_var.set(school_year)
                try:
                    term_int = int(term)
                except (TypeError, ValueError):
                    term_int = DEFAULT_TERM
                if term_int not in (1, 2):
                    term_int = DEFAULT_TERM
                self.current_term_var.set(term_int)
                
                if "classes" in config:
                    classes = config["classes"]
                    for year, count in classes.items():
                        if int(year) in self.spinboxes:
                            self.spinboxes[int(year)].set(count)
                self._refresh_period_label()
                            
        except Exception as e:
            logging.error(f"Fehler beim Laden der Config: {e}")

    def save_sph_config(self):
        """Speichert SPH Konfiguration (Klassen + aktive Periode)."""
        try:
            import json
            config_path = self.paths.sph_config_path
            
            # Read existing to preserve school/user if they were there (for prefill consistency with LoginWindow)
            config = {}
            if config_path.exists():
                 try: 
                     with open(config_path, "r") as f: config = json.load(f)
                 except: pass

            classes = {}
            for year, spin in self.spinboxes.items():
                classes[str(year)] = spin.get()
            
            config["classes"] = classes
            school_year, term = self._get_active_period()
            config["period"] = {
                "school_year": school_year,
                "term": term,
            }
            # Note: We rely on LoginWindow/Credentials to manage auth persistence.
            # We don't overwrite school/user here anymore because we don't have the widgets.
            
            with open(config_path, "w") as f:
                json.dump(config, f)
                
        except Exception as e:
            logging.error(f"Fehler beim Speichern der Config: {e}")

    def _load_db_transfer_meta(self) -> Dict[str, Any]:
        """Lädt optionale Metadaten zu DB-Import/Export aus der Config."""
        try:
            config_path = self.paths.sph_config_path
            if not config_path.exists():
                return {}
            with open(config_path, "r", encoding="utf-8") as f:
                config = json.load(f)
            meta = config.get("db_transfer", {})
            return meta if isinstance(meta, dict) else {}
        except Exception:
            return {}

    def _save_db_transfer_meta(self, **kwargs):
        """Speichert Metadaten zu DB-Import/Export in die Config."""
        try:
            config_path = self.paths.sph_config_path
            config = {}
            if config_path.exists():
                try:
                    with open(config_path, "r", encoding="utf-8") as f:
                        config = json.load(f)
                except Exception:
                    config = {}

            meta = config.get("db_transfer", {})
            if not isinstance(meta, dict):
                meta = {}
            meta.update(kwargs)
            config["db_transfer"] = meta

            self.path_manager.ensure_directory(config_path.parent)
            with open(config_path, "w", encoding="utf-8") as f:
                json.dump(config, f, ensure_ascii=False)
        except Exception as e:
            logging.error(f"Fehler beim Speichern der DB-Transfer-Metadaten: {e}")

    def load_sph_missing_overview(self):
        """Lädt den SPH-Abgleich für die aktive Periode."""
        try:
            self.sph_missing_overview = {}
            cache_path = self._get_sph_cache_path()
            if cache_path.exists():
                with open(cache_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                if isinstance(data, dict):
                    if "periods" in data and isinstance(data.get("periods"), dict):
                        self.sph_missing_overview = data["periods"].get(self._get_active_period_key(), {})
                    else:
                        # Backward compatibility: legacy cache without period split
                        self.sph_missing_overview = data
                    logging.info(
                        f"SPH-Abgleich aus Cache geladen ({len(self.sph_missing_overview)} Klassen)."
                    )
        except Exception as e:
            logging.error(f"Fehler beim Laden des SPH-Abgleich-Caches: {e}")

    def save_sph_missing_overview(self):
        """Speichert den aktuellen SPH-Abgleich dauerhaft auf Platte (pro Periode)."""
        try:
            cache_path = self._get_sph_cache_path()
            self.path_manager.ensure_directory(cache_path.parent)
            payload = {"periods": {}}
            if cache_path.exists():
                try:
                    with open(cache_path, "r", encoding="utf-8") as f:
                        existing = json.load(f)
                    if isinstance(existing, dict) and isinstance(existing.get("periods"), dict):
                        payload["periods"] = existing["periods"]
                except Exception:
                    payload = {"periods": {}}

            payload["periods"][self._get_active_period_key()] = self.sph_missing_overview
            with open(cache_path, "w", encoding="utf-8") as f:
                json.dump(payload, f, ensure_ascii=False)
            logging.info(
                f"SPH-Abgleich gespeichert ({len(self.sph_missing_overview)} Klassen)."
            )
        except Exception as e:
            logging.error(f"Fehler beim Speichern des SPH-Abgleich-Caches: {e}")

    def show_login_window(self):
        """Erlaubt das Ändern der Zugangsdaten"""
        from login_gui import LoginWindow
        login_win = LoginWindow(self.root, self.credentials_manager)
        self.root.wait_window(login_win.window)
        # Refresh the Import tab if it exists
        self.create_widgets_sph_section()

    def create_widgets_sph_section(self):
        """Aktualisiert nur den SPH-Anmeldebereich im Import-Tab"""
        # This is a bit tricky if we want to replace existing widgets. 
        # For simplicity, we just check if self.sph_login_container exists and refresh it.
        if hasattr(self, "sph_login_container"):
            for widget in self.sph_login_container.winfo_children():
                widget.destroy()
            
            if self.credentials_manager.credentials:
                 sch, usr, _ = self.credentials_manager.credentials
                 ttk.Label(self.sph_login_container, text=f"✅ Angemeldet als: {usr} @ {sch}", foreground="green", font=("Arial", 10, "bold")).pack(anchor=tk.W)
                 ttk.Button(self.sph_login_container, text="Anmeldung ändern", command=self.show_login_window).pack(anchor=tk.W, pady=5)
            else:
                 ttk.Label(self.sph_login_container, text="⚠️ Nicht eingeloggt", foreground="red").pack(anchor=tk.W)
                 ttk.Button(self.sph_login_container, text="Jetzt anmelden", command=self.show_login_window).pack(anchor=tk.W, pady=5)

    def run_sph_import(self):
        """Führt den SPH-Import Prozess aus"""
        if not self._confirm_period_before_import("SPH-Import"):
            self.status_manager.set_status("SPH-Import abgebrochen")
            return
        # Credentials from Manager
        if not self.credentials_manager.credentials:
             messagebox.showerror("Fehler", "Nicht eingeloggt. Bitte melden Sie sich zuerst an.")
             self.show_login_window()
             return
             
        school, user, pw = self.credentials_manager.credentials
        
        # Config aus Spinboxen
        tasks = []
        try:
            for year, spin in self.spinboxes.items():
                count_str = spin.get()
                if count_str.strip().isdigit():
                    count = int(count_str.strip())
                    if count > 0:
                        jg_formatted = f"{year:02d}"
                        tasks.append((jg_formatted, count))
        except Exception as e:
             messagebox.showerror("Konfigurationsfehler", f"Formatfehler: {e}")
             return

        # Speichern (Nur Config, nicht user/pw da jetzt secured)
        self.save_sph_config()
        self.log_to_import("SPH-Import gestartet...")
        self.status_manager.set_status("SPH-Import gestartet...", True)

        import threading
        t = threading.Thread(target=self._sph_worker, args=(school, user, pw, tasks))
        t.start()

    def _sph_worker(self, school, user, pw, tasks):
        """Hintergrund-Worker für SPH Download"""
        try:
            from sph_downloader import SPHDownloader
            downloader = SPHDownloader(logger=logging.getLogger("sph"))
            
            # Login
            self.queue_ui(self.log_to_import, "SPH: Login läuft...")
            self.queue_ui(self.status_manager.set_status, "SPH: Login...")
            downloader.login(school, user, pw)
            
            # Download Loop
            output_dir = self.paths.temp_dir
            output_dir.mkdir(parents=True, exist_ok=True)

            # 1) Primär: Autoerkennung (ohne manuelle Vorgabe)
            self.queue_ui(self.status_manager.set_status, "Autoerkenne Klassen aus SPH...")
            downloaded_files, auto_tasks = self._auto_detect_and_download_classes(
                downloader, output_dir
            )
            manual_fallback_used = False

            # 2) Backup-Fallback: manuelle Angaben nutzen, falls Autoerkennung nichts gefunden hat
            if not downloaded_files:
                if tasks:
                    manual_fallback_used = True
                    self.queue_ui(
                        self.status_manager.set_status,
                        "Autoerkennung ohne Treffer – nutze manuelle Klassenangaben (Backup)...",
                    )
                    downloaded_files = self._download_manual_tasks(
                        downloader, output_dir, tasks
                    )
                else:
                    self.queue_ui(
                        messagebox.showwarning,
                        "SPH Import",
                        "Autoerkennung hat keine Klassen gefunden und es wurden keine manuellen Klassenzahlen angegeben.",
                    )
                    self.queue_ui(self.status_manager.set_status, "SPH Import ohne Ergebnis")
                    return

            auto_summary = ", ".join([f"J{int(jg)}={cnt}" for jg, cnt in auto_tasks]) if auto_tasks else "keine Treffer"
            backup_summary = ", ".join([f"J{int(jg)}={cnt}" for jg, cnt in tasks]) if tasks else "nicht konfiguriert"
            self.queue_ui(self.log_to_import, f"Autoerkennung: {auto_summary}")
            if manual_fallback_used:
                self.queue_ui(self.log_to_import, f"Backup-Fallback aktiv: {backup_summary}")
            else:
                self.queue_ui(self.log_to_import, "Backup-Fallback nicht benötigt.")
            
            # Import Trigger
            if downloaded_files:
                self.queue_ui(self.status_manager.set_status, f"Importiere {len(downloaded_files)} Dateien...")
                run_meta = {
                    "auto_summary": auto_summary,
                    "backup_summary": backup_summary,
                    "manual_fallback_used": manual_fallback_used,
                    "downloaded_count": len(downloaded_files),
                }
                # Direkt im Worker aufrufen: _process_downloaded_files verarbeitet UI-Ausgaben selbst per queue_ui
                self._process_downloaded_files(downloaded_files, (school, user, pw), run_meta)
            else:
                self.queue_ui(messagebox.showwarning, "Ergebnis", "Keine Dateien erfolgreich geladen.")

        except Exception as e:
            self.queue_ui(messagebox.showerror, "SPH Fehler", f"{e}")
            self.queue_ui(self.status_manager.set_status, "Fehler bei SPH Import")

    def _process_downloaded_files(self, file_paths, sph_credentials=None, run_meta=None):
        """Verarbeitet heruntergeladene Dateien"""
        if threading.current_thread() == threading.main_thread():
            threading.Thread(
                target=self._process_downloaded_files,
                args=(file_paths, sph_credentials, run_meta),
                daemon=True,
            ).start()
            return
        try:
            self.path_manager.ensure_directory(self.db_path.parent)
            
            # Use KopfnotenImporter as context manager
            # Assuming KopfnotenImporter is available (it is in the same file)
            school_year, term = self._get_active_period()
            with KopfnotenImporter(str(self.db_path), school_year=school_year, term=term) as importer:
                count = 0
                failed = []
                total = len(file_paths)
                for idx, fp in enumerate(file_paths, start=1):
                    try:
                        self.queue_ui(
                            self.status_manager.set_status,
                            f"SPH-Import: {idx}/{total} - {Path(fp).name}",
                            True,
                        )
                        self.queue_ui(self.log_to_import, f"Importiere Klasse: {Path(fp).name}")
                        importer.import_excel_file(str(fp))
                        count += 1
                        self.queue_ui(self.log_to_import, f"✅ Erfolgreich: {Path(fp).name}")
                    except Exception as e:
                        failed.append((Path(fp).name, str(e)))
                        self.queue_ui(self.log_to_import, f"❌ Fehler bei {Path(fp).name}: {e}")
                        logging.getLogger("importer").error(
                            f"Fehler beim Import von {Path(fp).name}: {e}"
                        )
                # Nach Import einmal Artefakt-/Namensbereinigung ausführen
                importer._clean_existing_subjects()
                
            self.queue_ui(self.refresh_all_data)
            # SPH-Abgleich NACH abgeschlossenem Import laden
            if sph_credentials:
                school, user, pw = sph_credentials
                self.queue_ui(self.status_manager.set_status, "SPH-Abgleich wird nach Import aktualisiert...")
                threading.Thread(
                    target=self._sph_post_import_sync_worker,
                    args=(school, user, pw),
                    daemon=True
                ).start()

            summary_lines = []
            if run_meta:
                summary_lines.append(f"Autoerkennung: {run_meta.get('auto_summary', '-')}")
                summary_lines.append(
                    "Backup-Fallback: "
                    + ("verwendet" if run_meta.get("manual_fallback_used") else "nicht benötigt")
                )
                if run_meta.get("manual_fallback_used"):
                    summary_lines.append(f"Backup-Konfiguration: {run_meta.get('backup_summary', '-')}")
            summary_text = ("\n\n" + "\n".join(summary_lines)) if summary_lines else ""

            if failed:
                details = "\n".join([f"- {name}: {err}" for name, err in failed[:10]])
                more = ""
                if len(failed) > 10:
                    more = f"\n... und {len(failed) - 10} weitere Fehler."
                self.queue_ui(
                    messagebox.showwarning,
                    "SPH Import teilweise erfolgreich",
                    f"{count} Klassen erfolgreich importiert.\n"
                    f"{len(failed)} Klassen konnten nicht importiert werden.\n\n"
                    f"{details}{more}{summary_text}"
                )
            else:
                self.queue_ui(
                    messagebox.showinfo,
                    "SPH Import",
                    f"{count} Klassen erfolgreich importiert.{summary_text}"
                )

            if summary_lines:
                self.queue_ui(self.log_to_import, "Abschluss: " + " | ".join(summary_lines))
                self.queue_ui(
                    self.status_manager.set_status,
                    f"SPH-Import abgeschlossen (Auto: {run_meta.get('auto_summary', '-')}; "
                    f"Fallback: {'ja' if run_meta.get('manual_fallback_used') else 'nein'})"
                )
        except Exception as e:
            self.queue_ui(messagebox.showerror, "Import Fehler", f"{e}")

    def _download_manual_tasks(self, downloader, output_dir: Path, tasks: List[Tuple[str, int]]) -> List[Path]:
        """Lädt Klassen anhand manueller Konfiguration (Backup-Pfad)."""
        downloaded_files: List[Path] = []
        letters = CLASS_SUFFIX_LETTERS
        for jg, count in tasks:
            for i in range(min(count, MAX_CLASSES_PER_JAHRGANG)):
                suffix = letters[i]
                class_name = f"{jg}{suffix}"
                self.queue_ui(self.status_manager.set_status, f"Lade Klasse {class_name} (manuell)...")
                self.queue_ui(self.log_to_import, f"Lade Klasse {class_name} (manuell)...")
                file_path = downloader.download_class_list(class_name, jg, output_dir)
                if file_path:
                    downloaded_files.append(file_path)
                    self.queue_ui(self.log_to_import, f"✅ Download ok: {class_name}")
                else:
                    self.queue_ui(self.log_to_import, f"⚠️ Kein Download: {class_name}")
        return downloaded_files

    def _auto_detect_and_download_classes(self, downloader, output_dir: Path) -> Tuple[List[Path], List[Tuple[str, int]]]:
        """
        Erkennt Klassen pro Jahrgang automatisch durch sequenzielles Testen (05a … 05i).
        Stoppt je Jahrgang nach 2 Fehlversuchen in Folge nach erstem Treffer.
        """
        downloaded_files: List[Path] = []
        detected_tasks: List[Tuple[str, int]] = []
        letters = CLASS_SUFFIX_LETTERS
        years = sorted(self.spinboxes.keys()) if hasattr(self, "spinboxes") else [5, 6, 7, 8, 9, 10]

        for year in years:
            jg = f"{int(year):02d}"
            success_count = 0
            fail_streak = 0

            for i in range(MAX_CLASSES_PER_JAHRGANG):
                class_name = f"{jg}{letters[i]}"
                self.queue_ui(self.status_manager.set_status, f"Autocheck Klasse {class_name}...")
                self.queue_ui(self.log_to_import, f"Autocheck Klasse {class_name}...")
                file_path = downloader.download_class_list(class_name, jg, output_dir)

                if file_path:
                    downloaded_files.append(file_path)
                    success_count += 1
                    fail_streak = 0
                    self.queue_ui(self.log_to_import, f"✅ Download ok: {class_name}")
                else:
                    fail_streak += 1
                    self.queue_ui(self.log_to_import, f"— Nicht gefunden: {class_name}")
                    if success_count == 0 and fail_streak >= 1:
                        break
                    if success_count > 0 and fail_streak >= 2:
                        break

            if success_count > 0:
                detected_tasks.append((jg, success_count))

        return downloaded_files, detected_tasks

    def _sph_post_import_sync_worker(self, school, user, pw):
        """Lädt SPH-Abgleich im Anschluss an einen erfolgreichen Import."""
        try:
            from sph_downloader import SPHDownloader
            downloader = SPHDownloader(logger=logging.getLogger("sph"))
            downloader.login(school, user, pw)
            overview = downloader.fetch_missing_submissions_overview()
            self.sph_missing_overview = overview
            self.save_sph_missing_overview()
            self.queue_ui(self.refresh_analysis_data)
            self.queue_ui(self.refresh_insights_data)
            self.queue_ui(
                self.status_manager.set_status,
                f"SPH-Abgleich aktualisiert ({len(overview)} Klassen)."
            )
        except Exception as e:
            self.queue_ui(
                self.status_manager.set_status,
                f"SPH-Abgleich nach Import fehlgeschlagen: {e}"
            )

    def create_menu(self):
        """Erstellt vereinfachtes Menü"""
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)
        # Datei-Menü
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Datei", menu=file_menu)
        file_menu.add_command(label="Datenbank öffnen", command=self.open_database)
        file_menu.add_command(label="Datenbank importieren...", command=self.import_database_file)
        file_menu.add_command(label="Datenbank exportieren...", command=self.export_database_file)
        file_menu.add_command(label="Datenbank-Info", command=self.show_database_info)
        file_menu.add_command(label="Datenbank sichern", command=self.backup_database)
        file_menu.add_separator()
        file_menu.add_command(
            label="Backup-Klassenangaben...",
            command=self.show_backup_class_config,
        )
        file_menu.add_separator()
        file_menu.add_command(label="Beenden", command=self.root.quit)
        # Werkzeuge-Menü
        tools_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Werkzeuge", menu=tools_menu)
        if TEMPLATE_MANAGER_ENABLED:
            tools_menu.add_command(
                label="Template-Designer",
                command=self.template_designer.create_template_designer_window,
            )
        tools_menu.add_command(label="Logs anzeigen", command=self.show_logs)
        tools_menu.add_separator()
        tools_menu.add_command(
            label="Berechtigungen prüfen", command=self.check_permissions
        )
        # Hilfe-Menü
        help_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Hilfe", menu=help_menu)
        help_menu.add_command(label="Über", command=self.show_about)

    def create_import_tab(self):
        """Erstellt Import-Tab"""
        import_frame = ttk.Frame(self.notebook)
        self.notebook.add(import_frame, text="📥 Import")
        import_frame.columnconfigure(0, weight=1)
        import_frame.rowconfigure(0, weight=3)
        import_frame.rowconfigure(1, weight=2)
        import_frame.rowconfigure(2, weight=0)

        # --- 1. SPH (oben, präsent) ---
        sph_frame = ttk.LabelFrame(import_frame, text="1. Import aus Schulportal Hessen (SPH)")
        sph_frame.grid(row=0, column=0, sticky="nsew", padx=10, pady=(10, 5))

        sph_left = ttk.Frame(sph_frame)
        sph_left.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=8, pady=8)
        sph_right = ttk.Frame(sph_frame)
        sph_right.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=8, pady=8)

        self.sph_login_container = ttk.Frame(sph_left)
        self.sph_login_container.pack(fill=tk.X, padx=5, pady=5)
        self.create_widgets_sph_section()

        ttk.Label(
            sph_right,
            text="Automatische Klassenerkennung ist aktiv.",
            font=("Segoe UI", 10, "bold"),
            foreground="#15803d",
        ).pack(anchor=tk.W)
        ttk.Label(
            sph_right,
            text="Manuelle Klassenanzahlen nur als Backup – einstellbar unter Datei → Backup-Klassenangaben.",
            foreground="#555",
            wraplength=420,
        ).pack(anchor=tk.W, pady=(4, 12))

        self.sph_import_btn = tk.Button(
            sph_right,
            text="SPH Download & Import",
            command=self.run_sph_import,
            font=("Segoe UI", 13, "bold"),
            bg="#2563eb",
            fg="white",
            activebackground="#1d4ed8",
            activeforeground="white",
            relief=tk.FLAT,
            cursor="hand2",
            padx=16,
            pady=12,
        )
        self.sph_import_btn.pack(fill=tk.X, ipady=6)
        self.sph_status_label = ttk.Label(sph_right, text="-", foreground="#555")
        self.sph_status_label.pack(anchor=tk.W, pady=(8, 0))

        self._create_backup_class_config_dialog()

        # --- Import-Log (Mitte, für SPH und lokalen Import) ---
        log_frame = ttk.LabelFrame(import_frame, text="Import-Log")
        log_frame.grid(row=1, column=0, sticky="nsew", padx=10, pady=5)
        self.import_log = scrolledtext.ScrolledText(
            log_frame, height=10, state=tk.DISABLED, font=("Consolas", 9)
        )
        self.import_log.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # --- 2. Manueller Import (unten, kompakt) ---
        manual_frame = ttk.LabelFrame(import_frame, text="2. Excel-Dateien importieren (Lokal)")
        manual_frame.grid(row=2, column=0, sticky="ew", padx=10, pady=(5, 10))

        button_frame = ttk.Frame(manual_frame)
        button_frame.pack(fill=tk.X, padx=5, pady=4)
        ttk.Button(
            button_frame, text="Excel-Dateien auswählen", command=self.select_excel_files
        ).pack(side=tk.LEFT, padx=(0, 4))
        ttk.Button(
            button_frame, text="Alle importieren", command=self.import_all_files
        ).pack(side=tk.LEFT, padx=(0, 4))
        ttk.Button(
            button_frame, text="Auswahl löschen", command=self.clear_import_selection
        ).pack(side=tk.LEFT)
        ttk.Button(
            button_frame, text="Datenbank komplett löschen", command=self.delete_database
        ).pack(side=tk.RIGHT)

        list_frame = ttk.Frame(manual_frame)
        list_frame.pack(fill=tk.X, padx=5, pady=(0, 6))
        self.import_listbox = tk.Listbox(list_frame, selectmode=tk.EXTENDED, height=4, font=("Segoe UI", 9))
        scrollbar_import = ttk.Scrollbar(
            list_frame, orient=tk.VERTICAL, command=self.import_listbox.yview
        )
        self.import_listbox.config(yscrollcommand=scrollbar_import.set)
        self.import_listbox.pack(side=tk.LEFT, fill=tk.X, expand=True)
        scrollbar_import.pack(side=tk.RIGHT, fill=tk.Y)

    def _create_backup_class_config_dialog(self):
        """Erstellt den Backup-Klassendialog (ausgeblendet, öffnen über Datei-Menü)."""
        self.spinboxes = {}
        self._backup_cfg_dialog = tk.Toplevel(self.root)
        self._backup_cfg_dialog.withdraw()
        self._backup_cfg_dialog.title("Backup: Klassen-Konfiguration (Züge)")
        self._backup_cfg_dialog.transient(self.root)
        self._backup_cfg_dialog.resizable(False, False)

        ttk.Label(
            self._backup_cfg_dialog,
            text="Fallback, wenn die Autoerkennung keine Klassen findet.\n"
                 "Anzahl Klassen pro Jahrgang (0 = nicht verwenden).",
            wraplength=360,
            foreground="#555",
        ).pack(anchor=tk.W, padx=12, pady=(12, 8))

        config_grid = ttk.Frame(self._backup_cfg_dialog)
        config_grid.pack(fill=tk.X, padx=12, pady=(0, 8))

        years = [5, 6, 7, 8, 9, 10]
        for i, year in enumerate(years):
            r = i // 2
            c = (i % 2) * 2
            ttk.Label(config_grid, text=f"J{year}:").grid(row=r, column=c, sticky=tk.W, padx=2, pady=2)
            spin = ttk.Spinbox(config_grid, from_=0, to=MAX_CLASSES_PER_JAHRGANG, width=4)
            spin.set(3)
            spin.grid(row=r, column=c + 1, sticky=tk.W, padx=(2, 16), pady=2)
            self.spinboxes[year] = spin

        btn_frame = ttk.Frame(self._backup_cfg_dialog)
        btn_frame.pack(fill=tk.X, padx=12, pady=(0, 12))
        ttk.Button(btn_frame, text="Speichern & Schließen", command=self._close_backup_class_config).pack(
            side=tk.RIGHT
        )
        ttk.Button(btn_frame, text="Abbrechen", command=self._backup_cfg_dialog.withdraw).pack(
            side=tk.RIGHT, padx=(0, 8)
        )
        self._backup_cfg_dialog.protocol("WM_DELETE_WINDOW", self._close_backup_class_config)

    def show_backup_class_config(self):
        """Öffnet den Dialog für manuelle Backup-Klassenangaben."""
        if not hasattr(self, "_backup_cfg_dialog"):
            self._create_backup_class_config_dialog()
        self._backup_cfg_dialog.deiconify()
        self._backup_cfg_dialog.lift()
        self._backup_cfg_dialog.focus_force()

    def _close_backup_class_config(self):
        """Speichert Backup-Klassenangaben und schließt den Dialog."""
        self.save_sph_config()
        if hasattr(self, "_backup_cfg_dialog"):
            self._backup_cfg_dialog.withdraw()

    # Old create_sph_tab removed or ignored (can remove entire method if desired, but replacing logic here)
    # def create_sph_tab(self): ... -> DELETED/IGNORED content below


    def create_analysis_tab(self):
        """Erstellt Analyse-Tab"""
        analysis_frame = ttk.Frame(self.notebook)
        self.notebook.add(analysis_frame, text="🗄 Datenbank")
        period_frame = ttk.LabelFrame(analysis_frame, text="Aktive Periode")
        period_frame.pack(fill=tk.X, padx=10, pady=(8, 5))
        ttk.Label(period_frame, text="Schuljahr:").pack(side=tk.LEFT, padx=(6, 4), pady=6)
        self.school_year_combo = ttk.Combobox(
            period_frame,
            textvariable=self.current_school_year_var,
            width=12,
            state="normal",
            values=self._suggest_school_years(),
        )
        self.school_year_combo.pack(side=tk.LEFT, padx=(0, 12), pady=6)
        self.school_year_combo.bind("<<ComboboxSelected>>", self.on_period_changed)
        self.school_year_combo.bind("<FocusOut>", self.on_period_changed)
        ttk.Label(period_frame, text="Halbjahr:").pack(side=tk.LEFT, padx=(0, 4), pady=6)
        self.term_combo = ttk.Combobox(
            period_frame,
            textvariable=self.current_term_var,
            width=5,
            state="readonly",
            values=[1, 2],
        )
        self.term_combo.pack(side=tk.LEFT, padx=(0, 12), pady=6)
        self.term_combo.bind("<<ComboboxSelected>>", self.on_period_changed)
        self.period_info_label = ttk.Label(period_frame, foreground="#555")
        self.period_info_label.pack(side=tk.LEFT, padx=(6, 0), pady=6)
        self._refresh_period_label()

        # Filter
        filter_frame = ttk.LabelFrame(analysis_frame, text="Filter und Suche")
        filter_frame.pack(fill=tk.X, padx=10, pady=5)
        filter_controls = ttk.Frame(filter_frame)
        filter_controls.pack(fill=tk.X, padx=5, pady=5)
        ttk.Label(filter_controls, text="Klasse:").pack(side=tk.LEFT)
        self.class_filter = ttk.Combobox(filter_controls, width=10, state="readonly")
        self.class_filter.pack(side=tk.LEFT, padx=(5, 15))
        self.class_filter.bind("<<ComboboxSelected>>", lambda e: self.search_students())
        ttk.Label(filter_controls, text="Schüler:").pack(side=tk.LEFT)
        self.student_search = ttk.Entry(filter_controls, width=20)
        self.student_search.pack(side=tk.LEFT, padx=(5, 15))
        self.student_search.bind("<KeyRelease>", lambda e: self.search_students())
        ttk.Button(filter_controls, text="Suchen", command=self.search_students).pack(
            side=tk.LEFT, padx=(0, 5)
        )
        
        # NEW FILTERS
        ttk.Label(filter_controls, text="Lehrer:").pack(side=tk.LEFT, padx=(10, 5))
        ttk.Entry(filter_controls, textvariable=self.teacher_filter_var, width=10).pack(side=tk.LEFT, padx=(0, 5))

        ttk.Label(filter_controls, text="Status:").pack(side=tk.LEFT, padx=(10, 5))
        self.status_filter_combo = ttk.Combobox(filter_controls, textvariable=self.status_filter_var, width=12, state="readonly", values=["Alle", "Vollständig", "Unvollständig"])
        self.status_filter_combo.pack(side=tk.LEFT, padx=(0, 10))
        self.status_filter_combo.bind("<<ComboboxSelected>>", lambda e: self.search_students())
        
        ttk.Button(
            filter_controls, text="Filter zurücksetzen", command=self.reset_filters
        ).pack(side=tk.LEFT)
        ttk.Label(
            filter_controls,
            text="  SPH-Abgleich je Lernendem: rot/gelb/grün",
            foreground="#555"
        ).pack(side=tk.LEFT, padx=(10, 0))
        # Daten
        data_frame = ttk.LabelFrame(analysis_frame, text="Schüler und Noten")
        data_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        tree_frame = ttk.Frame(data_frame)
        tree_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        self.analysis_tree = ttk.Treeview(tree_frame, show="tree headings")
        v_scrollbar = ttk.Scrollbar(
            tree_frame, orient=tk.VERTICAL, command=self.analysis_tree.yview
        )
        h_scrollbar = ttk.Scrollbar(
            tree_frame, orient=tk.HORIZONTAL, command=self.analysis_tree.xview
        )
        self.analysis_tree.config(
            yscrollcommand=v_scrollbar.set, xscrollcommand=h_scrollbar.set
        )
        # Lernendenbezogene SPH-Ampelfarben
        self.analysis_tree.tag_configure("student_red", background="#f8d7da")
        self.analysis_tree.tag_configure("student_yellow", background="#fff3cd")
        self.analysis_tree.tag_configure("student_green", background="#d1e7dd")
        self.analysis_tree.grid(row=0, column=0, sticky="nsew")
        v_scrollbar.grid(row=0, column=1, sticky="ns")
        h_scrollbar.grid(row=1, column=0, sticky="ew")
        tree_frame.grid_rowconfigure(0, weight=1)
        tree_frame.grid_columnconfigure(0, weight=1)
        
        # Bindings
        # Bindings
        self.analysis_tree.bind("<Double-1>", self.on_tree_double_click)

        # Bearbeitung
        edit_frame = ttk.Frame(data_frame)
        edit_frame.pack(fill=tk.X, padx=5, pady=5)
        ttk.Button(
            edit_frame, text="Noten bearbeiten", command=self.edit_selected_grade
        ).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(
            edit_frame, text="Lernende deaktivieren", command=self.deactivate_selected_student
        ).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(
            edit_frame, text="Deaktivierte verwalten", command=self.manage_inactive_students
        ).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(
            edit_frame, text="Daten aktualisieren", command=self.refresh_analysis_data
        ).pack(side=tk.LEFT)
        # Neue Schaltfläche für Einzelschüler-Export
        ttk.Button(
            edit_frame, text="Schüler exportieren", command=self.export_selected_student
        ).pack(side=tk.RIGHT, padx=5)

        # Neue Schaltfläche für Fehllisten-Export
        ttk.Button(
            edit_frame, text="Fehlliste exportieren", command=self.export_missing_list
        ).pack(side=tk.RIGHT, padx=5)

    def create_insights_tab(self):
        """Erstellt den Analyse-Tab mit Kennzahlen, Rankings und Trends."""
        insights_frame = ttk.Frame(self.notebook)
        self.notebook.add(insights_frame, text="📊 Analyse")

        toolbar = ttk.Frame(insights_frame)
        toolbar.pack(fill=tk.X, padx=10, pady=(8, 4))
        ttk.Button(
            toolbar, text="Analyse aktualisieren", command=self.refresh_insights_data
        ).pack(side=tk.LEFT)
        self.insights_period_label = ttk.Label(toolbar, foreground="#555")
        self.insights_period_label.pack(side=tk.LEFT, padx=(12, 0))
        ttk.Label(
            toolbar,
            text="Basisperiode = Auswahl im Tab „Datenbank“ · ohne deaktivierte Lernende",
            foreground="#666",
        ).pack(side=tk.RIGHT, padx=(8, 0))

        compare_frame = ttk.LabelFrame(
            insights_frame,
            text="Vergleichsperioden (Strg/Klick für Mehrfachauswahl)",
        )
        compare_frame.pack(fill=tk.X, padx=10, pady=(0, 6))
        compare_inner = ttk.Frame(compare_frame)
        compare_inner.pack(fill=tk.X, padx=8, pady=6)
        self.insights_compare_listbox = tk.Listbox(
            compare_inner,
            selectmode=tk.EXTENDED,
            height=5,
            exportselection=False,
            font=("Segoe UI", 10),
        )
        compare_scroll = ttk.Scrollbar(
            compare_inner, orient=tk.VERTICAL, command=self.insights_compare_listbox.yview
        )
        self.insights_compare_listbox.configure(yscrollcommand=compare_scroll.set)
        self.insights_compare_listbox.pack(side=tk.LEFT, fill=tk.X, expand=True)
        compare_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self._insights_compare_period_map: List[Tuple[str, int]] = []

        kpi_frame = ttk.Frame(insights_frame)
        kpi_frame.pack(fill=tk.X, padx=10, pady=(0, 8))
        self.insights_kpi_vars = {}
        self._create_kpi_card(kpi_frame, "students", "Lernende", "0")
        self._create_kpi_card(kpi_frame, "overall_avg", "Gesamtschnitt", "-")
        self._create_kpi_card(kpi_frame, "completion", "Vollständigkeit", "-")
        self._create_kpi_card(kpi_frame, "classes", "Klassen", "0")

        self.insights_text_sections = {}
        self.insights_tables = {}
        section_notebook = ttk.Notebook(insights_frame)
        section_notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=(0, 10))
        section_defs = [
            ("overview", "Überblick", "text"),
            ("classes", "Klassen", "table"),
            ("years", "Jahrgänge", "table"),
            ("subjects", "Fächer & Ranking", "table"),
            ("top", "Top Lernende", "custom_top"),
            ("trends", "Entwicklung", "text"),
        ]
        for key, label, section_type in section_defs:
            frame = ttk.Frame(section_notebook)
            section_notebook.add(frame, text=label)
            if section_type == "text":
                text = scrolledtext.ScrolledText(
                    frame,
                    state=tk.DISABLED,
                    wrap=tk.WORD,
                    font=("Segoe UI", 10),
                    padx=12,
                    pady=10,
                )
                text.pack(fill=tk.BOTH, expand=True)
                self.insights_text_sections[key] = text
            elif key == "classes":
                self.insights_tables[key] = self._create_insights_table(
                    frame,
                    columns=["Klasse", "AV", "SV", "Gesamt", "Vollst. %", "Lernende"],
                    widths=[120, 90, 90, 90, 110, 90],
                    numeric_columns={"AV", "SV", "Gesamt", "Vollst. %", "Lernende"},
                )
            elif key == "years":
                self.insights_tables[key] = self._create_insights_table(
                    frame,
                    columns=["Jahrgang", "AV", "SV", "Gesamt", "Vollst. %", "Lernende"],
                    widths=[120, 90, 90, 90, 110, 90],
                    numeric_columns={"AV", "SV", "Gesamt", "Vollst. %", "Lernende"},
                )
            elif key == "subjects":
                self.insights_tables[key] = self._create_insights_table(
                    frame,
                    columns=["Rang", "Fach", "AV", "SV", "Gesamt", "Streuung", "Noten"],
                    widths=[70, 220, 90, 90, 90, 100, 80],
                    numeric_columns={"Rang", "AV", "SV", "Gesamt", "Streuung", "Noten"},
                )
            elif key == "top":
                self.insights_top_tables = self._create_top_section_widgets(frame)

        self.refresh_insights_data()

    def _create_kpi_card(self, parent, key: str, title: str, initial: str):
        card = ttk.Frame(parent, style="Card.TFrame", padding=(10, 8))
        card.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 8))
        ttk.Label(card, text=title, style="CardTitle.TLabel").pack(anchor=tk.W)
        value_var = tk.StringVar(value=initial)
        ttk.Label(card, textvariable=value_var, style="CardValue.TLabel").pack(anchor=tk.W, pady=(3, 0))
        self.insights_kpi_vars[key] = value_var

    def _create_insights_table(
        self,
        parent,
        columns: List[str],
        widths: List[int],
        numeric_columns: Optional[set] = None,
        stretch_columns: Optional[set] = None,
        tree_height: int = 18,
    ):
        numeric_columns = numeric_columns or set()
        stretch_columns = stretch_columns or set()
        wrapper = ttk.Frame(parent)
        wrapper.pack(fill=tk.BOTH, expand=True, padx=6, pady=6)
        tree = ttk.Treeview(wrapper, columns=columns, show="headings", height=tree_height)
        v_scroll = ttk.Scrollbar(wrapper, orient=tk.VERTICAL, command=tree.yview)
        h_scroll = ttk.Scrollbar(wrapper, orient=tk.HORIZONTAL, command=tree.xview)
        tree.configure(yscrollcommand=v_scroll.set, xscrollcommand=h_scroll.set)
        tree.tag_configure("row_even", background="#f8fafc")
        tree.tag_configure("row_odd", background="#ffffff")
        for col, width in zip(columns, widths):
            is_numeric = col in numeric_columns
            anchor = tk.E if is_numeric else tk.W
            tree.heading(col, text=col, anchor=anchor)
            tree.column(col, width=width, anchor=anchor, stretch=(col in stretch_columns))
        tree.grid(row=0, column=0, sticky="nsew")
        v_scroll.grid(row=0, column=1, sticky="ns")
        h_scroll.grid(row=1, column=0, sticky="ew")
        wrapper.grid_rowconfigure(0, weight=1)
        wrapper.grid_columnconfigure(0, weight=1)
        return tree

    def _create_top_section_widgets(self, parent):
        notebook = ttk.Notebook(parent)
        notebook.pack(fill=tk.BOTH, expand=True, padx=6, pady=6)

        top_cols = ["Rang", "Name", "Klasse", "Jg", "AV Ø", "SV Ø", "Gesamt Ø", "Fächer"]
        top_widths = [55, 260, 80, 45, 70, 70, 80, 60]
        top_numeric = {"Jg", "AV Ø", "SV Ø", "Gesamt Ø", "Fächer"}

        school_tab = ttk.Frame(notebook)
        notebook.add(school_tab, text="Top 10 – Schule")
        school_tree = self._create_insights_table(
            school_tab,
            columns=top_cols,
            widths=top_widths,
            numeric_columns=top_numeric,
            stretch_columns={"Name"},
            tree_height=12,
        )

        year_tab = ttk.Frame(notebook)
        notebook.add(year_tab, text="Top 3 – je Jahrgang")
        year_tree = self._create_insights_table(
            year_tab,
            columns=["Jahrgang", "Platz", "Name", "Klasse", "AV Ø", "SV Ø", "Gesamt Ø", "Fächer"],
            widths=[70, 55, 260, 80, 70, 70, 80, 60],
            numeric_columns={"AV Ø", "SV Ø", "Gesamt Ø", "Fächer"},
            stretch_columns={"Name"},
            tree_height=22,
        )

        class_tab = ttk.Frame(notebook)
        notebook.add(class_tab, text="Klassenbeste – je Klasse")
        class_tree = self._create_insights_table(
            class_tab,
            columns=["Klasse", "Name", "Jg", "AV Ø", "SV Ø", "Gesamt Ø", "Fächer"],
            widths=[80, 260, 45, 70, 70, 80, 60],
            numeric_columns={"Jg", "AV Ø", "SV Ø", "Gesamt Ø", "Fächer"},
            stretch_columns={"Name", "Klasse"},
            tree_height=22,
        )
        return {
            "school": school_tree,
            "year": year_tree,
            "class": class_tree,
        }

    def _rank_medal(self, rank: int) -> str:
        """Platz 1–3 als Medaille, sonst Zahl."""
        medals = {1: "🥇", 2: "🥈", 3: "🥉"}
        try:
            rank_int = int(rank)
        except (TypeError, ValueError):
            return str(rank)
        return medals.get(rank_int, str(rank_int))

    def _top_student_grade_values(self, student: Dict[str, Any]) -> Tuple[str, str, str, int]:
        """AV-, SV- und Gesamtwerte für Top-Tabellen."""
        return (
            self._fmt_avg(student.get("av_avg")),
            self._fmt_avg(student.get("sv_avg")),
            self._fmt_avg(student.get("gesamt_avg")),
            int(student.get("subjects_graded") or 0),
        )

    def _safe_avg(self, values: List[float]) -> Optional[float]:
        vals = [float(v) for v in values if v is not None]
        if not vals:
            return None
        return sum(vals) / len(vals)

    def _fmt_avg(self, value: Optional[float]) -> str:
        return "-" if value is None else f"{value:.2f}"

    def _extract_jahrgang_from_klasse(self, klasse: str) -> int:
        match = re.search(r"(\d+)", str(klasse or ""))
        return int(match.group(1)) if match else 0

    def _class_sort_key(self, klasse: Any) -> Tuple[int, str]:
        raw = str(klasse or "").strip()
        m = re.match(r"^(\d+)\s*([A-Za-z].*)?$", raw)
        if m:
            num = int(m.group(1))
            suffix = (m.group(2) or "").lower()
            return (num, suffix)
        m2 = re.match(r"^(\d+)", raw)
        if m2:
            return (int(m2.group(1)), raw.lower())
        return (999, raw.lower())

    def _period_sort_key(self, school_year: str, term: int) -> Tuple[int, int]:
        try:
            start_year = int(str(school_year).split("/")[0])
        except Exception:
            start_year = 0
        return (start_year, int(term))

    def _period_label(self, school_year: str, term: int) -> str:
        return f"{school_year} · HJ {term}"

    def _get_available_periods(self) -> List[Tuple[str, int]]:
        if not self.db_path.exists():
            return []
        with sqlite3.connect(self.db_path) as conn:
            rows = conn.execute(
                """
                SELECT DISTINCT n.schuljahr, n.halbjahr
                FROM noten n
                JOIN schueler s ON n.schueler_id = s.schueler_id
                WHERE n.schuljahr IS NOT NULL
                  AND n.halbjahr IS NOT NULL
                  AND COALESCE(s.is_active, 1) = 1
                """
            ).fetchall()
        periods = []
        for sy, term in rows:
            try:
                periods.append((str(sy), int(term)))
            except Exception:
                continue
        periods.sort(key=lambda p: self._period_sort_key(p[0], p[1]))
        return periods

    def _sync_insights_compare_periods(self, current_school_year: str, current_term: int):
        """Befüllt die Vergleichsperioden-Liste (ohne aktive Basisperiode)."""
        if not hasattr(self, "insights_compare_listbox"):
            return
        previous_selection = set(self._get_selected_compare_periods())
        available = self._get_available_periods()
        self.insights_compare_listbox.delete(0, tk.END)
        self._insights_compare_period_map = []

        for sy, term in available:
            if sy == current_school_year and int(term) == int(current_term):
                continue
            self._insights_compare_period_map.append((sy, term))
            self.insights_compare_listbox.insert(tk.END, self._period_label(sy, term))

        if not self._insights_compare_period_map:
            return

        selected_indices = []
        if previous_selection:
            for idx, period in enumerate(self._insights_compare_period_map):
                if period in previous_selection:
                    selected_indices.append(idx)
        if not selected_indices:
            prev_period = self._find_previous_period(current_school_year, current_term)
            if prev_period and prev_period in self._insights_compare_period_map:
                selected_indices = [self._insights_compare_period_map.index(prev_period)]

        for idx in selected_indices:
            self.insights_compare_listbox.selection_set(idx)

    def _get_selected_compare_periods(self) -> List[Tuple[str, int]]:
        if not hasattr(self, "insights_compare_listbox"):
            return []
        return [
            self._insights_compare_period_map[idx]
            for idx in self.insights_compare_listbox.curselection()
            if 0 <= idx < len(self._insights_compare_period_map)
        ]

    def _find_previous_period(self, school_year: str, term: int) -> Optional[Tuple[str, int]]:
        periods = self._get_available_periods()
        current_key = self._period_sort_key(school_year, term)
        previous = None
        for sy, t in periods:
            if self._period_sort_key(sy, t) < current_key:
                previous = (sy, t)
        if previous:
            return previous
        # Fallback: anderes Halbjahr im selben Schuljahr
        alt_term = 1 if int(term) == 2 else 2
        if (school_year, alt_term) in periods:
            return (school_year, alt_term)
        return None

    def _collect_analysis_dataset(self, school_year: str, term: int) -> Dict[str, Any]:
        dataset = {
            "students": [],
            "class_stats": {},
            "year_stats": {},
            "subject_stats": [],
            "school": {},
            "top_by_class": {},
            "top_by_year": {},
            "top_school": [],
            "period": (school_year, term),
        }
        if not self.db_path.exists():
            return dataset

        with sqlite3.connect(self.db_path) as conn:
            rows = conn.execute(
                """
                SELECT
                    s.schueler_id,
                    s.name,
                    s.klasse,
                    s.target_subjects,
                    n.fach_id,
                    COALESCE(f.fach_lang, f.fach_kurz, '') AS fach_name,
                    n.note_av,
                    n.note_sv,
                    n.note_av_special,
                    n.note_sv_special
                FROM schueler s
                LEFT JOIN noten n ON s.schueler_id = n.schueler_id
                    AND n.schuljahr = ?
                    AND n.halbjahr = ?
                LEFT JOIN faecher f ON n.fach_id = f.fach_id
                WHERE COALESCE(s.is_active, 1) = 1
                ORDER BY s.klasse, s.name
                """,
                (school_year, term),
            ).fetchall()

        student_map = {}
        subject_map = {}
        for s_id, name, klasse, target_subjects, fach_id, fach_name, av, sv, av_special, sv_special in rows:
            if s_id not in student_map:
                student_map[s_id] = {
                    "id": s_id,
                    "name": name,
                    "klasse": klasse,
                    "jahrgang": self._extract_jahrgang_from_klasse(klasse),
                    "target_subjects": target_subjects,
                    "av_notes": [],
                    "sv_notes": [],
                    "combined_notes": [],
                    "graded_subjects": set(),
                    "filled_entries": 0,
                }
            sm = student_map[s_id]
            if av is not None:
                sm["av_notes"].append(float(av))
                sm["combined_notes"].append(float(av))
            if av is not None or av_special is not None:
                sm["filled_entries"] += 1
            if sv is not None:
                sm["sv_notes"].append(float(sv))
                sm["combined_notes"].append(float(sv))
            if sv is not None or sv_special is not None:
                sm["filled_entries"] += 1
            if fach_id and (av is not None or sv is not None or av_special is not None or sv_special is not None):
                sm["graded_subjects"].add(int(fach_id))

            subject_name = (fach_name or "").strip()
            if subject_name:
                if subject_name not in subject_map:
                    subject_map[subject_name] = {"av_notes": [], "sv_notes": [], "combined_notes": []}
                if av is not None:
                    subject_map[subject_name]["av_notes"].append(float(av))
                    subject_map[subject_name]["combined_notes"].append(float(av))
                if sv is not None:
                    subject_map[subject_name]["sv_notes"].append(float(sv))
                    subject_map[subject_name]["combined_notes"].append(float(sv))

        for sm in student_map.values():
            av_avg = self._safe_avg(sm["av_notes"])
            sv_avg = self._safe_avg(sm["sv_notes"])
            ges_avg = self._safe_avg(sm["combined_notes"])
            notes_total = sm["filled_entries"]
            target = sm["target_subjects"] or self.get_default_target_for_grade(sm["jahrgang"])
            completion_pct = None
            if target:
                completion_pct = min(100.0, (notes_total / max(1, target * 2)) * 100.0)
            dataset["students"].append(
                {
                    "id": sm["id"],
                    "name": sm["name"],
                    "klasse": sm["klasse"],
                    "jahrgang": sm["jahrgang"],
                    "av_avg": av_avg,
                    "sv_avg": sv_avg,
                    "gesamt_avg": ges_avg,
                    "notes_total": notes_total,
                    "subjects_graded": len(sm["graded_subjects"]),
                    "completion_pct": completion_pct,
                }
            )

        def summarize_group(students: List[Dict[str, Any]]) -> Dict[str, Any]:
            av_values = [s["av_avg"] for s in students if s["av_avg"] is not None]
            sv_values = [s["sv_avg"] for s in students if s["sv_avg"] is not None]
            total_values = [s["gesamt_avg"] for s in students if s["gesamt_avg"] is not None]
            completion_vals = [s["completion_pct"] for s in students if s["completion_pct"] is not None]
            return {
                "count": len(students),
                "av_avg": self._safe_avg(av_values),
                "sv_avg": self._safe_avg(sv_values),
                "gesamt_avg": self._safe_avg(total_values),
                "completion_pct": self._safe_avg(completion_vals),
            }

        class_groups = {}
        year_groups = {}
        for s in dataset["students"]:
            class_groups.setdefault(s["klasse"], []).append(s)
            year_groups.setdefault(s["jahrgang"], []).append(s)

        dataset["class_stats"] = {
            klasse: summarize_group(students)
            for klasse, students in sorted(class_groups.items(), key=lambda item: self._class_sort_key(item[0]))
        }
        dataset["year_stats"] = {
            jahrgang: summarize_group(students)
            for jahrgang, students in sorted(year_groups.items(), key=lambda item: item[0])
        }
        dataset["school"] = summarize_group(dataset["students"])

        subject_stats = []
        for subject_name, vals in subject_map.items():
            combined = vals["combined_notes"]
            stddev = statistics.pstdev(combined) if len(combined) >= 2 else 0.0
            subject_stats.append(
                {
                    "subject": subject_name,
                    "av_avg": self._safe_avg(vals["av_notes"]),
                    "sv_avg": self._safe_avg(vals["sv_notes"]),
                    "gesamt_avg": self._safe_avg(combined),
                    "count": len(combined),
                    "stddev": stddev,
                }
            )
        subject_stats.sort(
            key=lambda x: (
                float("inf") if x["gesamt_avg"] is None else x["gesamt_avg"],
                x["subject"].lower(),
            )
        )
        dataset["subject_stats"] = subject_stats

        def top_students(students: List[Dict[str, Any]], limit: int = 5, min_subjects: int = 4):
            valid = [
                s for s in students
                if s["gesamt_avg"] is not None and s["subjects_graded"] >= min_subjects
            ]
            valid.sort(key=lambda s: (s["gesamt_avg"], -s["notes_total"], s["name"].lower()))
            return valid[:limit]

        dataset["top_school"] = top_students(dataset["students"], limit=10)
        dataset["top_by_class"] = {
            klasse: top_students(students, limit=1, min_subjects=3)
            for klasse, students in sorted(class_groups.items(), key=lambda item: self._class_sort_key(item[0]))
        }
        dataset["top_by_year"] = {
            jahrgang: top_students(students, limit=3, min_subjects=3)
            for jahrgang, students in sorted(year_groups.items(), key=lambda item: item[0])
        }

        return dataset

    def _render_overview_section(
        self, current: Dict[str, Any], compare_datasets: Optional[List[Dict[str, Any]]] = None
    ) -> str:
        school_year, term = current["period"]
        school = current["school"]
        students = current["students"]
        classes_count = len(current["class_stats"])
        years_count = len(current["year_stats"])
        complete_count = sum(1 for s in students if (s["completion_pct"] or 0) >= 99.9)
        incomplete_count = len(students) - complete_count
        lines = [
            f"Periode: {school_year} (HJ {term})",
            "Hinweis: Deaktivierte Lernende sind in allen Kennzahlen ausgeschlossen.",
            "",
            "Gesamtkennzahlen:",
            f"- Lernende (aktiv): {len(students)}",
            f"- Klassen: {classes_count}",
            f"- Jahrgänge: {years_count}",
            f"- AV-Durchschnitt: {self._fmt_avg(school.get('av_avg'))}",
            f"- SV-Durchschnitt: {self._fmt_avg(school.get('sv_avg'))}",
            f"- Gesamtdurchschnitt: {self._fmt_avg(school.get('gesamt_avg'))}",
            f"- Durchschnittliche Vollständigkeit: {self._fmt_avg(school.get('completion_pct'))}%",
            f"- Vollständig (>=99.9%): {complete_count}",
            f"- Unvollständig: {incomplete_count}",
        ]
        compare_datasets = compare_datasets or []
        if compare_datasets:
            lines.extend(["", "Vergleich zu ausgewählten Perioden:"])
            for previous in compare_datasets:
                prev_school = previous["school"]
                prev_year, prev_term = previous["period"]
                delta = None
                if school.get("gesamt_avg") is not None and prev_school.get("gesamt_avg") is not None:
                    delta = school["gesamt_avg"] - prev_school["gesamt_avg"]
                student_delta = len(students) - len(previous.get("students", []))
                lines.extend([
                    "",
                    f"→ {prev_year} (HJ {prev_term}):",
                    f"  Gesamtdurchschnitt: {self._fmt_avg(prev_school.get('gesamt_avg'))} "
                    f"(Delta {self._fmt_avg(delta)})",
                    f"  Lernende aktiv: {len(previous.get('students', []))} (Delta {student_delta:+d})",
                    f"  Vollständigkeit: {self._fmt_avg(prev_school.get('completion_pct'))}%",
                ])
        return "\n".join(lines)

    def _render_class_section(self, current: Dict[str, Any]) -> str:
        lines = ["Klassendurchschnitte (aktive Periode):", ""]
        ranked = sorted(
            current["class_stats"].items(),
            key=lambda item: (float("inf") if item[1]["gesamt_avg"] is None else item[1]["gesamt_avg"], str(item[0])),
        )
        for klasse, stats in ranked:
            lines.append(
                f"- {klasse}: AV {self._fmt_avg(stats['av_avg'])} | "
                f"SV {self._fmt_avg(stats['sv_avg'])} | Gesamt {self._fmt_avg(stats['gesamt_avg'])} | "
                f"Vollständigkeit {self._fmt_avg(stats['completion_pct'])}% | Lernende {stats['count']}"
            )
        if len(lines) == 2:
            lines.append("- Keine Klassendaten")
        return "\n".join(lines)

    def _render_year_section(self, current: Dict[str, Any]) -> str:
        lines = ["Jahrgangsdurchschnitte (aktive Periode):", ""]
        for jahrgang, stats in current["year_stats"].items():
            lines.append(
                f"- Jg {jahrgang}: AV {self._fmt_avg(stats['av_avg'])} | "
                f"SV {self._fmt_avg(stats['sv_avg'])} | Gesamt {self._fmt_avg(stats['gesamt_avg'])} | "
                f"Vollständigkeit {self._fmt_avg(stats['completion_pct'])}% | Lernende {stats['count']}"
            )
        if len(lines) == 2:
            lines.append("- Keine Jahrgangsdaten")
        return "\n".join(lines)

    def _render_subject_section(self, current: Dict[str, Any]) -> str:
        lines = ["Fächerdurchschnitte und Ranking:", ""]
        ranked = [s for s in current["subject_stats"] if s["gesamt_avg"] is not None]
        if not ranked:
            return "\n".join(lines + ["- Keine Fachdaten"])

        lines.append("Beste Fächer (niedrigster Schnitt):")
        for idx, row in enumerate(ranked[:10], start=1):
            lines.append(
                f"- {idx}. {row['subject']}: AV {self._fmt_avg(row['av_avg'])}, "
                f"SV {self._fmt_avg(row['sv_avg'])}, Gesamt {self._fmt_avg(row['gesamt_avg'])}, "
                f"Streuung {row['stddev']:.2f}, Noten {row['count']}"
            )

        lines.append("")
        lines.append("Schwächste Fächer (höchster Schnitt):")
        for idx, row in enumerate(list(reversed(ranked[-10:])), start=1):
            lines.append(
                f"- {idx}. {row['subject']}: AV {self._fmt_avg(row['av_avg'])}, "
                f"SV {self._fmt_avg(row['sv_avg'])}, Gesamt {self._fmt_avg(row['gesamt_avg'])}, "
                f"Streuung {row['stddev']:.2f}, Noten {row['count']}"
            )
        return "\n".join(lines)

    def _render_top_section(self, current: Dict[str, Any]) -> str:
        lines = [
            "Top Lernende (Mindestkriterium: >=4 benotete Fächer):",
            "",
            "Schulweit:",
        ]
        if current["top_school"]:
            for idx, s in enumerate(current["top_school"], start=1):
                lines.append(
                    f"- {idx}. {s['name']} ({s['klasse']}): Schnitt {self._fmt_avg(s['gesamt_avg'])}, "
                    f"AV {self._fmt_avg(s['av_avg'])}, SV {self._fmt_avg(s['sv_avg'])}, Noten {s['notes_total']}"
                )
        else:
            lines.append("- Keine ausreichend bewerteten Lernenden")

        lines.append("")
        lines.append("Top pro Jahrgang:")
        for jahrgang, students in current["top_by_year"].items():
            if students:
                top = students[0]
                lines.append(f"- Jg {jahrgang}: {top['name']} ({top['klasse']}) mit {self._fmt_avg(top['gesamt_avg'])}")
            else:
                lines.append(f"- Jg {jahrgang}: keine ausreichenden Daten")

        lines.append("")
        lines.append("Top pro Klasse:")
        for klasse, students in current["top_by_class"].items():
            if students:
                top = students[0]
                lines.append(f"- {klasse}: {top['name']} mit {self._fmt_avg(top['gesamt_avg'])}")
            else:
                lines.append(f"- {klasse}: keine ausreichenden Daten")
        return "\n".join(lines)

    def _render_trends_section(
        self, current: Dict[str, Any], compare_datasets: Optional[List[Dict[str, Any]]] = None
    ) -> str:
        compare_datasets = compare_datasets or []
        if not compare_datasets:
            return (
                "Keine Vergleichsperioden ausgewählt.\n\n"
                "Wählen Sie im Feld „Vergleichsperioden“ eine oder mehrere Perioden "
                "(Strg+Klick für Mehrfachauswahl) und klicken Sie auf „Analyse aktualisieren“."
            )

        current_year, current_term = current["period"]
        lines = [
            f"Entwicklung der Basisperiode {current_year} (HJ {current_term}) "
            f"gegen {len(compare_datasets)} ausgewählte Periode(n):",
            "Hinweis: Deaktivierte Lernende sind ausgeschlossen.",
            "",
        ]

        for previous in compare_datasets:
            prev_year, prev_term = previous["period"]
            lines.append(f"=== Vergleich mit {prev_year} (HJ {prev_term}) ===")

            current_school = current["school"]
            prev_school = previous["school"]
            school_delta = None
            if current_school.get("gesamt_avg") is not None and prev_school.get("gesamt_avg") is not None:
                school_delta = current_school["gesamt_avg"] - prev_school["gesamt_avg"]
            lines.append(
                f"- Schule gesamt: {self._fmt_avg(prev_school.get('gesamt_avg'))} -> "
                f"{self._fmt_avg(current_school.get('gesamt_avg'))} (Delta {self._fmt_avg(school_delta)})"
            )
            lines.append(
                f"- Lernende aktiv: {len(previous.get('students', []))} -> "
                f"{len(current.get('students', []))}"
            )

            lines.append("")
            lines.append("Klassenentwicklung:")
            class_lines = 0
            for klasse, curr_stats in current["class_stats"].items():
                prev_stats = previous["class_stats"].get(klasse)
                if not prev_stats:
                    continue
                delta = None
                if curr_stats.get("gesamt_avg") is not None and prev_stats.get("gesamt_avg") is not None:
                    delta = curr_stats["gesamt_avg"] - prev_stats["gesamt_avg"]
                lines.append(
                    f"  - {klasse}: {self._fmt_avg(prev_stats.get('gesamt_avg'))} -> "
                    f"{self._fmt_avg(curr_stats.get('gesamt_avg'))} (Delta {self._fmt_avg(delta)})"
                )
                class_lines += 1
            if class_lines == 0:
                lines.append("  - Keine Klassen mit Daten in beiden Perioden")

            lines.append("")
            lines.append("Jahrgangsentwicklung:")
            year_lines = 0
            for jahrgang, curr_stats in current["year_stats"].items():
                prev_stats = previous["year_stats"].get(jahrgang)
                if not prev_stats:
                    continue
                delta = None
                if curr_stats.get("gesamt_avg") is not None and prev_stats.get("gesamt_avg") is not None:
                    delta = curr_stats["gesamt_avg"] - prev_stats["gesamt_avg"]
                lines.append(
                    f"  - Jg {jahrgang}: {self._fmt_avg(prev_stats.get('gesamt_avg'))} -> "
                    f"{self._fmt_avg(curr_stats.get('gesamt_avg'))} (Delta {self._fmt_avg(delta)})"
                )
                year_lines += 1
            if year_lines == 0:
                lines.append("  - Keine Jahrgänge mit Daten in beiden Perioden")

            lines.append("")
            lines.append("Fachentwicklung (Top 10 Deltas):")
            prev_subject_map = {row["subject"]: row for row in previous["subject_stats"]}
            deltas = []
            for row in current["subject_stats"]:
                prev_row = prev_subject_map.get(row["subject"])
                if not prev_row:
                    continue
                if row["gesamt_avg"] is None or prev_row["gesamt_avg"] is None:
                    continue
                deltas.append(
                    (row["subject"], row["gesamt_avg"] - prev_row["gesamt_avg"], prev_row["gesamt_avg"], row["gesamt_avg"])
                )
            if deltas:
                deltas.sort(key=lambda x: abs(x[1]), reverse=True)
                for subject, delta, old_val, new_val in deltas[:10]:
                    lines.append(
                        f"  - {subject}: {old_val:.2f} -> {new_val:.2f} (Delta {delta:+.2f})"
                    )
            else:
                lines.append("  - Keine vergleichbaren Fachdaten vorhanden")
            lines.append("")

        return "\n".join(lines).strip()

    def _set_insights_text_section(self, section_key: str, text: str):
        widget = self.insights_text_sections.get(section_key)
        if not widget:
            return
        widget.config(state=tk.NORMAL)
        widget.delete("1.0", tk.END)
        widget.insert(tk.END, text)
        widget.config(state=tk.DISABLED)

    def _set_insights_table_rows(self, section_key: str, rows: List[Tuple[Any, ...]]):
        tree = self.insights_tables.get(section_key)
        if not tree:
            return
        tree.delete(*tree.get_children())
        for idx, row in enumerate(rows):
            tag = "row_even" if idx % 2 == 0 else "row_odd"
            tree.insert("", tk.END, values=row, tags=(tag,))

    def _set_tree_rows(self, tree, rows: List[Tuple[Any, ...]]):
        if not tree:
            return
        tree.delete(*tree.get_children())
        for idx, row in enumerate(rows):
            tag = "row_even" if idx % 2 == 0 else "row_odd"
            tree.insert("", tk.END, values=row, tags=(tag,))

    def refresh_insights_data(self):
        """Aktualisiert aggregierte Analysen für die aktive Periode."""
        if not hasattr(self, "insights_text_sections"):
            return
        school_year, term = self._get_active_period()
        if hasattr(self, "insights_period_label"):
            self.insights_period_label.config(text=f"Basisperiode: {school_year} (HJ {term})")
        self._sync_insights_compare_periods(school_year, term)

        if not self.db_path.exists():
            if hasattr(self, "insights_kpi_vars"):
                self.insights_kpi_vars["students"].set("0")
                self.insights_kpi_vars["overall_avg"].set("-")
                self.insights_kpi_vars["completion"].set("-")
                self.insights_kpi_vars["classes"].set("0")
            for key in self.insights_text_sections:
                self._set_insights_text_section(key, "Keine Datenbank gefunden.")
            for key in self.insights_tables:
                self._set_insights_table_rows(key, [])
            if hasattr(self, "insights_top_tables"):
                self._set_tree_rows(self.insights_top_tables.get("school"), [])
                self._set_tree_rows(self.insights_top_tables.get("year"), [])
                self._set_tree_rows(self.insights_top_tables.get("class"), [])
            return

        try:
            current = self._collect_analysis_dataset(school_year, term)
            compare_datasets = []
            for cmp_sy, cmp_term in self._get_selected_compare_periods():
                compare_datasets.append(self._collect_analysis_dataset(cmp_sy, cmp_term))

            if hasattr(self, "insights_kpi_vars"):
                self.insights_kpi_vars["students"].set(str(len(current.get("students", []))))
                self.insights_kpi_vars["overall_avg"].set(self._fmt_avg(current.get("school", {}).get("gesamt_avg")))
                completion = current.get("school", {}).get("completion_pct")
                completion_text = self._fmt_avg(completion)
                self.insights_kpi_vars["completion"].set("-" if completion_text == "-" else f"{completion_text}%")
                self.insights_kpi_vars["classes"].set(str(len(current.get("class_stats", {}))))

            self._set_insights_text_section("overview", self._render_overview_section(current, compare_datasets))
            self._set_insights_text_section("trends", self._render_trends_section(current, compare_datasets))

            class_rows = []
            for klasse, stats in sorted(current["class_stats"].items(), key=lambda item: self._class_sort_key(item[0])):
                class_rows.append(
                    (
                        klasse,
                        self._fmt_avg(stats["av_avg"]),
                        self._fmt_avg(stats["sv_avg"]),
                        self._fmt_avg(stats["gesamt_avg"]),
                        self._fmt_avg(stats["completion_pct"]),
                        stats["count"],
                    )
                )
            self._set_insights_table_rows("classes", class_rows)

            year_rows = []
            for jahrgang, stats in sorted(current["year_stats"].items(), key=lambda item: item[0]):
                year_rows.append(
                    (
                        f"Jg {jahrgang}",
                        self._fmt_avg(stats["av_avg"]),
                        self._fmt_avg(stats["sv_avg"]),
                        self._fmt_avg(stats["gesamt_avg"]),
                        self._fmt_avg(stats["completion_pct"]),
                        stats["count"],
                    )
                )
            self._set_insights_table_rows("years", year_rows)

            subject_rows = []
            ranked_subjects = [s for s in current["subject_stats"] if s["gesamt_avg"] is not None]
            for idx, row in enumerate(ranked_subjects, start=1):
                subject_rows.append(
                    (
                        idx,
                        row["subject"],
                        self._fmt_avg(row["av_avg"]),
                        self._fmt_avg(row["sv_avg"]),
                        self._fmt_avg(row["gesamt_avg"]),
                        f"{row['stddev']:.2f}",
                        row["count"],
                    )
                )
            self._set_insights_table_rows("subjects", subject_rows)

            school_rows = []
            for idx, s in enumerate(current["top_school"][:10], start=1):
                av, sv, gesamt, faecher = self._top_student_grade_values(s)
                school_rows.append(
                    (self._rank_medal(idx), s["name"], s["klasse"], s["jahrgang"], av, sv, gesamt, faecher)
                )
            year_rows_top = []
            for jahrgang, students in sorted(current["top_by_year"].items(), key=lambda item: item[0]):
                for idx, s in enumerate(students[:3], start=1):
                    av, sv, gesamt, faecher = self._top_student_grade_values(s)
                    year_rows_top.append(
                        (f"Jg {jahrgang}", self._rank_medal(idx), s["name"], s["klasse"], av, sv, gesamt, faecher)
                    )
            class_rows_top = []
            for klasse in sorted(current["class_stats"].keys(), key=self._class_sort_key):
                students = current["top_by_class"].get(klasse, [])
                if students:
                    s = students[0]
                    av, sv, gesamt, faecher = self._top_student_grade_values(s)
                    class_rows_top.append(
                        (klasse, s["name"], s["jahrgang"], av, sv, gesamt, faecher)
                    )
                else:
                    class_rows_top.append(
                        (klasse, "— (keine ausreichenden Daten)", "-", "-", "-", "-", "-")
                    )
            if hasattr(self, "insights_top_tables"):
                self._set_tree_rows(self.insights_top_tables.get("school"), school_rows)
                self._set_tree_rows(self.insights_top_tables.get("year"), year_rows_top)
                self._set_tree_rows(self.insights_top_tables.get("class"), class_rows_top)
        except Exception as e:
            logging.error(f"Fehler beim Aktualisieren der Analysekennzahlen: {e}")
            if hasattr(self, "insights_kpi_vars"):
                self.insights_kpi_vars["students"].set("-")
                self.insights_kpi_vars["overall_avg"].set("-")
                self.insights_kpi_vars["completion"].set("-")
                self.insights_kpi_vars["classes"].set("-")
            for key in self.insights_text_sections:
                self._set_insights_text_section(key, f"Analyse konnte nicht berechnet werden:\n{e}")
            for key in self.insights_tables:
                self._set_insights_table_rows(key, [])
            if hasattr(self, "insights_top_tables"):
                self._set_tree_rows(self.insights_top_tables.get("school"), [])
                self._set_tree_rows(self.insights_top_tables.get("year"), [])
                self._set_tree_rows(self.insights_top_tables.get("class"), [])

    def create_export_tab(self):
        """Erstellt vereinfachten Export-Tab"""
        export_frame = ttk.Frame(self.notebook)
        self.notebook.add(export_frame, text="📤 Export")
        self.export_tab = export_frame
        # Export-Optionen
        options_frame = ttk.LabelFrame(export_frame, text="Export-Optionen")
        options_frame.pack(fill=tk.X, padx=10, pady=5)
        # Template-Auswahl
        template_frame = ttk.Frame(options_frame)
        template_frame.pack(fill=tk.X, padx=5, pady=5)
        ttk.Label(template_frame, text="Template:").pack(side=tk.LEFT)
        template_entry = ttk.Entry(
            template_frame, textvariable=self.template_var, width=50
        )
        template_entry.pack(side=tk.LEFT, padx=(5, 5), fill=tk.X, expand=True)
        ttk.Button(
            template_frame, text="Durchsuchen", command=self.select_template
        ).pack(side=tk.RIGHT, padx=(5, 0))
        # Ausgabe-Verzeichnis
        output_frame = ttk.Frame(options_frame)
        output_frame.pack(fill=tk.X, padx=5, pady=5)
        ttk.Label(output_frame, text="Ausgabe:").pack(side=tk.LEFT)
        output_entry = ttk.Entry(output_frame, textvariable=self.output_var, width=50)
        output_entry.pack(side=tk.LEFT, padx=(5, 5), fill=tk.X, expand=True)
        ttk.Button(
            output_frame, text="Durchsuchen", command=self.select_output_dir
        ).pack(side=tk.RIGHT)
        # Export-Datum
        date_frame = ttk.Frame(options_frame)
        date_frame.pack(fill=tk.X, padx=5, pady=5)
        ttk.Label(date_frame, text="Export-Datum:").pack(side=tk.LEFT)
        self.export_date_var = tk.StringVar(value=datetime.now().strftime("%d.%m.%Y"))
        date_entry = ttk.Entry(date_frame, textvariable=self.export_date_var, width=15)
        date_entry.pack(side=tk.LEFT, padx=(5, 5))
        ttk.Button(date_frame, text="Heute", command=lambda: self.export_date_var.set(datetime.now().strftime("%d.%m.%Y"))).pack(side=tk.LEFT)
        # Klassenauswahl
        class_frame = ttk.LabelFrame(export_frame, text="Klassenauswahl")
        class_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        class_controls = ttk.Frame(class_frame)
        class_controls.pack(fill=tk.X, padx=5, pady=5)
        ttk.Button(
            class_controls, text="Alle auswählen", command=self.select_all_classes
        ).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(
            class_controls, text="Auswahl umkehren", command=self.invert_class_selection
        ).pack(side=tk.LEFT, padx=(0, 5))
        # Export-Button
        self.export_btn = ttk.Button(
            class_controls,
            text="🚀 Export Serienbrief",
            command=self.start_optimized_export,
        )
        self.export_btn.pack(side=tk.RIGHT, padx=(5, 0))
        # Content
        content_frame = ttk.Frame(class_frame)
        content_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        # Klassen-Liste
        list_frame = ttk.Frame(content_frame)
        list_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        ttk.Label(list_frame, text="Verfügbare Klassen:").pack(anchor=tk.W)
        self.export_listbox = tk.Listbox(list_frame, selectmode=tk.EXTENDED)
        list_scrollbar = ttk.Scrollbar(
            list_frame, orient=tk.VERTICAL, command=self.export_listbox.yview
        )
        self.export_listbox.config(yscrollcommand=list_scrollbar.set)
        self.export_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        list_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        # Export-Log
        log_frame = ttk.Frame(content_frame)
        log_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(10, 0))
        ttk.Label(log_frame, text="Export-Log:").pack(anchor=tk.W)
        self.export_log = scrolledtext.ScrolledText(
            log_frame, width=50, state=tk.DISABLED
        )
        self.export_log.pack(fill=tk.BOTH, expand=True)
        # Progress-Bar
        self.export_progress = ttk.Progressbar(export_frame, mode="indeterminate")
        self.export_progress.pack(fill=tk.X, padx=10, pady=5)

    def create_template_tab(self):
        """Erstellt Template-Tab (nur wenn TEMPLATE_MANAGER_ENABLED)."""
        if not TEMPLATE_MANAGER_ENABLED:
            return
        template_frame = ttk.Frame(self.notebook)
        self.notebook.add(template_frame, text="📝 Templates")
        # Template-Designer
        designer_frame = ttk.LabelFrame(template_frame, text="Template-Designer")
        designer_frame.pack(fill=tk.X, padx=10, pady=5)
        designer_text = ttk.Label(
            designer_frame,
            text="Erstellen Sie einfache Templates für den Export.\n"
            "Der Designer hilft Ihnen bei der Erstellung von Templates ohne komplexe Validierung.",
        )
        designer_text.pack(padx=10, pady=10)
        ttk.Button(
            designer_frame,
            text="Template-Designer öffnen",
            command=self.template_designer.create_template_designer_window,
        ).pack(pady=5)
        # Template-Übersicht
        overview_frame = ttk.LabelFrame(template_frame, text="Verfügbare Templates")
        overview_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        self.template_list = tk.Listbox(overview_frame)
        template_scroll = ttk.Scrollbar(
            overview_frame, orient=tk.VERTICAL, command=self.template_list.yview
        )
        self.template_list.config(yscrollcommand=template_scroll.set)
        self.template_list.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5, pady=5)
        template_scroll.pack(side=tk.RIGHT, fill=tk.Y, pady=5)
        # Template-Buttons
        template_buttons = ttk.Frame(overview_frame)
        template_buttons.pack(fill=tk.X, padx=5, pady=5)
        ttk.Button(
            template_buttons, text="Aktualisieren", command=self.refresh_template_list
        ).pack(side=tk.LEFT, padx=5)
        ttk.Button(
            template_buttons,
            text="Template verwenden",
            command=self.use_selected_template,
        ).pack(side=tk.LEFT, padx=5)

    # ===================== NEUE FUNKTIONEN =====================

    def export_selected_student(self):
        """Exportiert ausgewählten Schüler"""
        selection = self.analysis_tree.selection()
        if not selection:
            messagebox.showwarning(
                "Keine Auswahl", "Bitte wählen Sie einen Schüler aus."
            )
            return
        item = self.analysis_tree.item(selection[0])
        values = item["values"]
        student_id = values[0]
        student_name = values[1]
        student_class = values[2]

        # Template-Prüfung
        template_path = Path(self.template_var.get().strip())
        # Make sure we're using an absolute path
        if not template_path.is_absolute():
            template_path = Path.cwd() / template_path
        if not template_path or not template_path.exists():
            messagebox.showerror(
                "Template fehlt", "Bitte wählen Sie eine gültige Template-Datei aus."
            )
            # Zur Export-Registerkarte wechseln für Template-Auswahl
            self.notebook.select(getattr(self, "export_tab", 3))
            return

        # Ausgabeverzeichnis prüfen
        output_dir = Path(self.output_var.get().strip())
        if not output_dir:
            output_dir = self.paths.output_word_dir
            self.output_var.set(str(output_dir))
        try:
            self.path_manager.ensure_directory(output_dir)
        except Exception as e:
            messagebox.showerror(
                "Ausgabe-Fehler",
                f"Ausgabeverzeichnis konnte nicht erstellt werden:\n{e}",
            )
            return
        # Export-UI vorbereiten
        self.export_running = True
        self.status_manager.set_status(f"Exportiere Schüler: {student_name}...", True)

        # Export in separatem Thread
        export_date = self.export_date_var.get().strip()
        export_thread = threading.Thread(
            target=self.run_student_export,
            args=(student_id, student_name, student_class, template_path, output_dir, export_date),
            daemon=True,
        )
        export_thread.start()

    def run_student_export(
        self,
        student_id: int,
        student_name: str,
        student_class: str,
        template_path: Path,
        output_dir: Path,
        export_date: str
    ):
        """Führt Schüler-Export in separatem Thread aus"""
        try:
            school_year, term = self._get_active_period()
            with OptimizedKopfnotenExporter(self.db_path, school_year=school_year, term=term) as exporter:
                start_time = datetime.now()
                # Export durchführen für einzelnen Schüler
                summary = exporter.export_horizontal_tables(
                    output_dir, template_path, [student_class], student_id, export_date=export_date
                )
                end_time = datetime.now()
                duration = end_time - start_time

                # Erfolgs-Meldung
                if summary["gesamt_dateien"] > 0:
                    schueler_result = list(summary["schueler_details"].values())[0]
                    success_msg = (
                        f"✅ Export für {student_name} erfolgreich!\n\n"
                        f"Dauer: {duration.total_seconds():.1f} Sekunden\n"
                        f"Datei: {Path(schueler_result['output_file']).name}\n"
                        f"Ausgabe: {output_dir}\n\n"
                        f"Format: Horizontale Tabelle"
                    )
                    # GUI-Thread für MessageBox
                    self.root.after(
                        100, lambda: messagebox.showinfo(f"Export für {student_name}", success_msg)
                    )
                else:
                    error_msg = f"❌ Export für {student_name} fehlgeschlagen"
                    self.root.after(
                        100, lambda: messagebox.showerror("Export-Fehler", error_msg)
                    )
        except Exception as e:
            error_msg = f"❌ Export für {student_name} fehlgeschlagen: {str(e)}"
            logging.error(f"Student-Export-Fehler: {e}")
            self.root.after(
                100, lambda: messagebox.showerror("Export-Fehler", error_msg)
            )
        finally:
            self.root.after(100, lambda: self.status_manager.clear_status())
            self.root.after(100, lambda: setattr(self, "export_running", False))

    # ===================== OPTIMIERTE FUNKTIONEN =====================

    def start_optimized_export(self):
        """Startet optimierten horizontalen Export"""
        if self.export_running:
            messagebox.showwarning("Export läuft", "Es läuft bereits ein Export!")
            return

        # Validierungen
        selected_indices = self.export_listbox.curselection()
        if not selected_indices:
            messagebox.showwarning(
                "Keine Auswahl", "Bitte wählen Sie mindestens eine Klasse aus."
            )
            return

        template_path = Path(self.template_var.get().strip())
        # Make sure we're using an absolute path
        if not template_path.is_absolute():
            template_path = Path.cwd() / template_path
        if not template_path or not template_path.exists():
            messagebox.showerror(
                "Template fehlt", "Bitte wählen Sie eine gültige Template-Datei aus."
            )
            return

        output_dir = Path(self.output_var.get().strip())
        if not output_dir:
            output_dir = self.paths.output_word_dir
            self.output_var.set(str(output_dir))
        try:
            self.path_manager.ensure_directory(output_dir)
        except Exception as e:
            messagebox.showerror(
                "Ausgabe-Fehler",
                f"Ausgabeverzeichnis konnte nicht erstellt werden:\n{e}",
            )
            return

        selected_classes = [self.export_listbox.get(i) for i in selected_indices]

        # Export-UI vorbereiten
        self.export_running = True
        self.export_btn.config(state=tk.DISABLED, text="Export läuft...")
        self.export_progress.start()
        self.clear_export_log()
        self.log_to_export(
            f"Starte optimierten horizontalen Export für {len(selected_classes)} Klassen"
        )
        self.status_manager.set_status("Optimierter Export läuft...", True)

        # Export in separatem Thread
        export_date = self.export_date_var.get().strip()
        export_thread = threading.Thread(
            target=self.run_optimized_export,
            args=(selected_classes, template_path, output_dir, export_date),
            daemon=True,
        )
        export_thread.start()

    def run_optimized_export(
        self, klassen: List[str], template_path: Path, output_dir: Path, export_date: str
    ):
        """Führt optimierten Export aus"""
        try:
            school_year, term = self._get_active_period()
            with OptimizedKopfnotenExporter(self.db_path, school_year=school_year, term=term) as exporter:
                start_time = datetime.now()
                self.log_to_export(f"Start: {start_time.strftime('%H:%M:%S')}")

                summary = exporter.export_horizontal_tables(
                    output_dir, template_path, klassen, export_date=export_date
                )
                end_time = datetime.now()
                duration = end_time - start_time
                # Detaillierte Log-Ausgabe
                self.log_to_export(f"\nOptimierter Export Details:")
                for klasse, details in summary["klassen_details"].items():
                    if details["datei_erstellt"]:
                        self.log_to_export(
                            f"✅ {klasse}: {details['schueler_count']} Schüler, "
                            f"max. {details['faecher_count']} Fächer"
                        )
                        self.log_to_export(
                            f" Datei: {Path(details['output_file']).name}"
                        )
                    else:
                        self.log_to_export(
                            f"❌ {klasse}: {details.get('fehler', 'Unbekannter Fehler')}"
                        )

                # Erfolg-Meldung
                success_msg = (
                    f"✅ Optimierter horizontaler Export erfolgreich!\n\n"
                    f"Dateien erstellt: {summary['gesamt_dateien']}\n"
                    f"Fehler: {summary['gesamt_fehler']}\n"
                    f"Dauer: {duration.total_seconds():.1f} Sekunden\n"
                    f"Ausgabe: {output_dir}\n\n"
                    f"Format: Horizontale Tabellen mit Beschriftungsspalte"
                )
                self.log_to_export(success_msg)

                # GUI-Thread für MessageBox
                self.root.after(
                    100, lambda: messagebox.showinfo("Export erfolgreich", success_msg)
                )
        except Exception as e:
            error_msg = f"❌ Optimierter Export fehlgeschlagen: {str(e)}"
            self.log_to_export(error_msg)
            logging.error(f"Export-Fehler: {e}")
            self.root.after(
                100, lambda: messagebox.showerror("Export-Fehler", error_msg)
            )
        finally:
            self.root.after(100, self.reset_export_ui)

    def edit_selected_grade(self):
        """Öffnet vereinfachten Noten-Editor"""
        selection = self.analysis_tree.selection()
        if not selection:
            messagebox.showwarning(
                "Keine Auswahl", "Bitte wählen Sie einen Schüler aus."
            )
            return
        item = self.analysis_tree.item(selection[0])
        values = item["values"]
        student_id = values[0]
        student_name = values[1]
        student_class = values[2]

        # Gather all student IDs in current order for navigation
        all_items = self.analysis_tree.get_children()
        student_context_list = []
        for it in all_items:
             v = self.analysis_tree.item(it)["values"]
             if v:
                 student_context_list.append({
                     "id": v[0],
                     "name": v[1],
                     "klasse": v[2]
                 })

        # Verwende vereinfachten Editor
        grade_editor = SimplifiedGradeEditor(self.root, str(self.db_path))
        # Pass parent (self) to allow refresh callback from Editor
        grade_editor.app = self 
        grade_editor.open_grade_editor(student_id, student_name, student_class, student_context_list)

    def on_tree_double_click(self, event):
        """Handler für Doppelklick auf Treeview"""
        region = self.analysis_tree.identify("region", event.x, event.y)
        if region != "cell":
            return
            
        col = self.analysis_tree.identify_column(event.x)
        # col ist string wie '#1', '#2'...
        # ID=0, Name=1 (#1), Klasse=2 (#2), Fächer=3 (#3)...
        # Spalten-Definition in refresh_analysis_data:
        # columns = ["ID", "Name", "Klasse", "Fächer", "AV-Noten", "SV-Noten", "Status"]
        # Treeview columns start indices: #1 -> Name (ID hidden is #0 text?), No wait.
        # Treeview columns: identifier list. 
        # identify_column returns '#n' where n is 1-based index of display columns?
        # Let's check columns definition:
        # self.analysis_tree["columns"] = ["ID", "Name", "Klasse", "Fächer", ...]
        # #0 is tree column (hidden). #1 is ID. #2 is Name. #3 is Klasse. #4 is Fächer.
        # Wait, let's verify column definition in refresh_analysis_data.
        # columns = ["ID", "Name", "Klasse", "Fächer", ...]
        # Treeview identifier '#1' corresponds to the first column in the values list?
        # Actually usually:
        # #0 = Label column (hidden).
        # #1 = First data column (ID).
        # #4 = Fächer.
        
        if col == "#4": # Spalte "Fächer"
            selection = self.analysis_tree.selection()
            if not selection: return
            item_id = selection[0]
            values = self.analysis_tree.item(item_id, "values")
            if not values: return
            
            s_id = values[0]
            s_name = values[1]
            current_target_str = values[3]
            
            try:
                current_target = int(current_target_str)
            except:
                current_target = 0
                
            new_target = simpledialog.askinteger(
                "Soll-Fächer bearbeiten",
                f"Soll-Wert für Fächeranzahl für {s_name}:",
                initialvalue=current_target,
                minvalue=1,
                maxvalue=30,
                parent=self.root
            )
            
            if new_target is not None:
                try:
                    with sqlite3.connect(self.db_path) as conn:
                        conn.execute("UPDATE schueler SET target_subjects = ? WHERE schueler_id = ?", (new_target, s_id))
                        conn.commit()
                    self.refresh_analysis_data()
                except Exception as e:
                    messagebox.showerror("Fehler", f"Konnte Wert nicht speichern: {e}")
            return
            
        # Default behavior for other columns: Edit Grades
        self.edit_selected_grade()

    def refresh_template_list(self):
        """Aktualisiert Template-Liste"""
        if not TEMPLATE_MANAGER_ENABLED or not hasattr(self, "template_list"):
            return
        self.template_list.delete(0, tk.END)
        template_dir = self.paths.templates_dir
        if template_dir.exists():
            for template_file in template_dir.glob("*.docx"):
                self.template_list.insert(tk.END, template_file.name)

    def use_selected_template(self):
        """Verwendet ausgewähltes Template"""
        selection = self.template_list.curselection()
        if not selection:
            messagebox.showwarning(
                "Keine Auswahl", "Bitte wählen Sie ein Template aus."
            )
            return
        template_name = self.template_list.get(selection[0])
        template_path = self.paths.templates_dir / template_name
        if template_path.exists():
            # Make sure we're using an absolute path
            absolute_path = template_path.resolve()
            self.template_var.set(str(absolute_path))
            messagebox.showinfo(
                "Template gewählt", f"Template ausgewählt: {template_name}"
            )
            # Wechsle zum Export-Tab
            self.notebook.select(getattr(self, "export_tab", 3))

    # ===================== UTILITY-FUNKTIONEN =====================

    def reset_export_ui(self):
        """Setzt Export-UI zurück"""
        self.export_running = False
        self.export_btn.config(state=tk.NORMAL, text="🚀 Export Serienbrief")
        self.export_progress.stop()
        self.status_manager.clear_status()

    def select_template(self):
        """Template-Datei auswählen"""
        filename = filedialog.askopenfilename(
            title="Template-Datei auswählen",
            filetypes=[("Word-Dokumente", "*.docx"), ("Alle Dateien", "*.*")],
            initialdir=str(self.paths.templates_dir.resolve()),
        )
        if filename:
            # Make sure we're using an absolute path
            absolute_path = Path(filename).resolve()
            self.template_var.set(str(absolute_path))
            self.log_to_export(f"Template ausgewählt: {absolute_path.name}")

    def select_output_dir(self):
        """Ausgabeverzeichnis auswählen"""
        dirname = filedialog.askdirectory(
            title="Ausgabeverzeichnis auswählen",
            initialdir=str(self.paths.output_word_dir.resolve()),
        )
        if dirname:
            self.output_var.set(dirname)
            self.log_to_export(f"Ausgabeverzeichnis: {dirname}")

    def select_excel_files(self):
        """Excel-Dateien für Import auswählen"""
        filenames = filedialog.askopenfilenames(
            title="Excel-Dateien auswählen",
            filetypes=[("Excel-Dateien", "*.xlsx *.xls"), ("Alle Dateien", "*.*")],
            initialdir=str(self.paths.import_dir.resolve()),
        )
        if filenames:
            for filename in filenames:
                if filename not in self.import_listbox.get(0, tk.END):
                    self.import_listbox.insert(tk.END, filename)
            self.log_to_import(f"{len(filenames)} Datei(en) hinzugefügt")

    def import_all_files(self):
        """Importiert alle ausgewählten Excel-Dateien"""
        if not self._confirm_period_before_import("Excel-Import"):
            self.status_manager.set_status("Excel-Import abgebrochen")
            return
        if self.import_listbox.size() == 0:
            messagebox.showwarning(
                "Keine Dateien", "Bitte wählen Sie zuerst Excel-Dateien aus."
            )
            return
        files = list(self.import_listbox.get(0, tk.END))
        self.status_manager.set_status(f"Importiere {len(files)} Dateien...", True)
        import_thread = threading.Thread(
            target=self.run_import, args=(files,), daemon=True
        )
        import_thread.start()

    def run_import(self, files: List[str]):
        """Führt Import in separatem Thread aus"""
        try:
            self.path_manager.ensure_directory(self.db_path.parent)
            school_year, term = self._get_active_period()
            with KopfnotenImporter(str(self.db_path), school_year=school_year, term=term) as importer:
                successful = 0
                total = len(files)
                for idx, file_path in enumerate(files, start=1):
                    try:
                        self.status_manager.set_status(
                            f"Import läuft: {idx}/{total} - {Path(file_path).name}", True
                        )
                        self.log_to_import(f"Importiere: {Path(file_path).name}")
                        importer.import_excel_file(file_path)
                        successful += 1
                        self.log_to_import(f"✅ Erfolgreich: {Path(file_path).name}")
                    except Exception as e:
                        self.log_to_import(f"❌ Fehler bei {Path(file_path).name}: {e}")
                        logging.error(f"Import-Fehler für {file_path}: {e}")
                self.log_to_import(
                    f"\nImport abgeschlossen: {successful}/{len(files)} erfolgreich"
                )
                self.root.after(100, self.refresh_all_data)
        except Exception as e:
            self.log_to_import(f"❌ Import-Fehler: {e}")
            logging.error(f"Allgemeiner Import-Fehler: {e}")
        finally:
            self.root.after(100, lambda: self.status_manager.clear_status())

    def clear_import_selection(self):
        """Löscht Auswahl der Import-Dateien"""
        self.import_listbox.delete(0, tk.END)
        self.log_to_import("Auswahl gelöscht")

    def select_all_classes(self):
        """Wählt alle Klassen aus"""
        self.export_listbox.selection_set(0, tk.END)

    def invert_class_selection(self):
        """Kehrt Klassenauswahl um"""
        for i in range(self.export_listbox.size()):
            if i in self.export_listbox.curselection():
                self.export_listbox.selection_clear(i)
            else:
                self.export_listbox.selection_set(i)

    def log_to_import(self, message: str):
        """Fügt Nachricht zum Import-Log hinzu"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.update_log_widget(self.import_log, f"[{timestamp}] {message}")

    def log_to_export(self, message: str):
        """Fügt Nachricht zum Export-Log hinzu"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.update_log_widget(self.export_log, f"[{timestamp}] {message}")

    def clear_export_log(self):
        """Leert das Export-Log"""
        self.export_log.config(state=tk.NORMAL)
        self.export_log.delete(1.0, tk.END)
        self.export_log.config(state=tk.DISABLED)

    def update_log_widget(self, widget, message: str):
        """Thread-sichere Log-Widget-Aktualisierung"""
        def update():
            widget.config(state=tk.NORMAL)
            widget.insert(tk.END, message + "\n")
            widget.see(tk.END)
            widget.config(state=tk.DISABLED)

        if threading.current_thread() == threading.main_thread():
            update()
        else:
            self.root.after(10, update)

    def export_missing_list(self):
        """Exportiert detaillierte Liste der Schüler mit Status Unvollständig nach Excel"""
        try:
            school_year, term = self._get_active_period()
            # Output Directory sicherstellen
            output_dir = self.paths.output_excel_dir
            self.path_manager.ensure_directory(output_dir)

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = output_dir / f"Kontrollliste_Luecken_{timestamp}.xlsx"

            import pandas as pd
            from collections import defaultdict
            from copy import copy

            with sqlite3.connect(self.db_path) as conn:
                summary_rows = conn.execute(
                    """
                    SELECT
                        s.schueler_id,
                        s.name,
                        s.klasse,
                        s.target_subjects,
                        COUNT(CASE WHEN n.note_av IS NOT NULL OR n.note_av_special IS NOT NULL THEN 1 END) AS av_count,
                        COUNT(CASE WHEN n.note_sv IS NOT NULL OR n.note_sv_special IS NOT NULL THEN 1 END) AS sv_count,
                        COUNT(DISTINCT n.fach_id) AS faecher_count
                    FROM schueler s
                    LEFT JOIN noten n ON s.schueler_id = n.schueler_id
                        AND n.schuljahr = ?
                        AND n.halbjahr = ?
                    WHERE COALESCE(s.is_active, 1) = 1
                    GROUP BY s.schueler_id, s.name, s.klasse, s.target_subjects
                    ORDER BY s.klasse, s.name
                    """
                , (school_year, term)).fetchall()

                detail_rows = conn.execute(
                    """
                    SELECT
                        s.schueler_id,
                        s.klasse,
                        COALESCE(f.fach_lang, f.fach_kurz, '') AS fach_name,
                        n.lehrer_kuerzel,
                        n.note_av,
                        n.note_sv,
                        n.note_av_special,
                        n.note_sv_special
                    FROM schueler s
                    LEFT JOIN noten n ON s.schueler_id = n.schueler_id
                        AND n.schuljahr = ?
                        AND n.halbjahr = ?
                    LEFT JOIN faecher f ON n.fach_id = f.fach_id
                    WHERE COALESCE(s.is_active, 1) = 1
                    ORDER BY s.klasse, s.schueler_id
                    """
                , (school_year, term)).fetchall()

            subjects_local_map = defaultdict(dict)
            class_subject_teachers = defaultdict(lambda: defaultdict(set))
            for s_id, klasse, fach_name, lehrer, av, sv, av_special, sv_special in detail_rows:
                fach = (fach_name or "").strip()
                if not fach:
                    continue
                if lehrer:
                    class_subject_teachers[str(klasse)][fach].add(str(lehrer).strip())
                if fach not in subjects_local_map[s_id]:
                    subjects_local_map[s_id][fach] = {"av": False, "sv": False}
                if av is not None or av_special is not None:
                    subjects_local_map[s_id][fach]["av"] = True
                if sv is not None or sv_special is not None:
                    subjects_local_map[s_id][fach]["sv"] = True

            class_subject_grade_counts = defaultdict(lambda: defaultdict(int))
            for s_id, name, klasse, target_db, av_count, sv_count, faecher_count in summary_rows:
                for subj, vals in subjects_local_map.get(s_id, {}).items():
                    if vals.get("av") or vals.get("sv"):
                        subj_key = self._normalize_subject_for_sph(subj)
                        class_subject_grade_counts[str(klasse)][subj_key] += 1

            export_rows = []
            sheets_data = defaultdict(list)

            for s_id, name, klasse, target_db, av_count, sv_count, faecher_count in summary_rows:
                match_year = re.match(r"(\d+)", str(klasse))
                jahrgang = int(match_year.group(1)) if match_year else 0
                target = target_db if target_db else self.get_default_target_for_grade(jahrgang)
                current_notes = (av_count or 0) + (sv_count or 0)
                local_status = self._calculate_status(jahrgang, current_notes, faecher_count or 0, target)

                sph_text, _tag, sph_status_override = self._get_sph_alignment_for_student(
                    klasse,
                    subjects_local_map.get(s_id, {}),
                    local_status,
                    dict(class_subject_grade_counts.get(str(klasse), {})),
                )
                final_status = sph_status_override if sph_status_override else local_status

                # Für Fehlliste nur unvollständige Einträge
                if final_status != "Unvollständig":
                    continue

                missing_details = []
                student_subjects = subjects_local_map.get(s_id, {})
                for fach in sorted(student_subjects.keys()):
                    vals = student_subjects[fach]
                    av_ok = bool(vals.get("av"))
                    sv_ok = bool(vals.get("sv"))
                    if av_ok and sv_ok:
                        continue

                    teachers = sorted(class_subject_teachers.get(str(klasse), {}).get(fach, set()))
                    teacher_text = ", ".join(teachers) if teachers else "-"
                    missing_details.append(f"{fach} [{teacher_text}]")

                # Fallback falls Status unvollständig ist, aber keine fachscharfen Lücken ermittelt wurden
                if not missing_details:
                    expected_notes = (target or 0) * 2 if target else (faecher_count or 0) * 2
                    diff = max(0, expected_notes - current_notes)
                    if diff > 0:
                        missing_details.append(f"Nicht zuordenbare Lücken: {diff} Note(n)")
                    else:
                        missing_details.append("Unvollständig laut Statuslogik")

                row = {
                    "Klasse": klasse,
                    "Name": name,
                    "Soll-Fächer": target,
                    "AV-Noten": av_count or 0,
                    "SV-Noten": sv_count or 0,
                    "Status": final_status,
                    "SPH-Abgleich": sph_text,
                    "Fehlende Fächer / Lehrkräfte": " | ".join(missing_details),
                }
                export_rows.append(row)
                sheets_data[str(klasse)].append(row)

            if not export_rows:
                messagebox.showinfo("Info", "Keine Schüler mit Status 'Unvollständig' für den Export gefunden.")
                return

            with pd.ExcelWriter(filename, engine="openpyxl") as writer:
                df_gesamt = pd.DataFrame(export_rows)
                df_gesamt.to_excel(writer, sheet_name="Gesamtübersicht", index=False)

                for klasse, rows in sheets_data.items():
                    df_klasse = pd.DataFrame(rows)
                    df_klasse.to_excel(writer, sheet_name=str(klasse)[:31], index=False)

                workbook = writer.book
                for sheet_name in workbook.sheetnames:
                    ws = workbook[sheet_name]
                    for col in ws.columns:
                        column = col[0].column_letter
                        header_cell = ws[f"{column}1"]
                        if header_cell.value in ["Name", "SPH-Abgleich"]:
                            ws.column_dimensions[column].width = 28
                        elif header_cell.value == "Fehlende Fächer / Lehrkräfte":
                            ws.column_dimensions[column].width = 80
                        elif header_cell.value in ["Status", "Klasse"]:
                            ws.column_dimensions[column].width = 16
                        else:
                            ws.column_dimensions[column].width = 14
                        for cell in col:
                            new_align = copy(cell.alignment)
                            new_align.wrap_text = True
                            cell.alignment = new_align
                    ws.page_setup.orientation = "landscape"
                    ws.page_setup.fitToWidth = 1

            self.log_to_export(f"Fehlliste exportiert: {filename}")
            messagebox.showinfo(
                "Export erfolgreich",
                f"Datei erstellt:\n{filename}\n({len(export_rows)} Einträge)"
            )
            try:
                if os.name == "nt":
                    os.startfile(str(output_dir))
            except Exception:
                pass

        except Exception as e:
            logging.error(f"Fehler beim Listen-Export: {e}")
            messagebox.showerror("Export-Fehler", f"Fehler: {e}")

    def delete_database(self):
        """Löscht die gesamte Datenbank"""
        if messagebox.askyesno(
            "Datenbank löschen",
            "Sind Sie sicher? Dies löscht ALLE importierten Daten unwiderruflich!",
            icon="warning"
        ):
            try:
                if self.db_path.exists():
                    # Schließe Verbindungen? SQLite macht das automatisch bei Context Manager Exit
                    # Aber sicherheitshalber: Garbage Collection oder einfach remove probieren
                    try:
                        self.db_path.unlink()
                    except PermissionError:
                        messagebox.showerror("Fehler", "Datenbank ist noch geöffnet. Bitte Neustart versuchen.")
                        return

                # UI leeren
                self.analysis_tree.delete(*self.analysis_tree.get_children())
                self.status_manager.set_status("Datenbank gelöscht")
                
                # Neu initialisieren (erstellt leere DB)
                self.load_initial_data()
                messagebox.showinfo("Erfolg", "Datenbank wurde zurückgesetzt.")
                
            except Exception as e:
                logging.error(f"Fehler beim Löschen der DB: {e}")
                messagebox.showerror("Fehler", f"Konnte Datenbank nicht löschen: {e}")

    def load_initial_data(self):
        """Lädt initiale Daten"""
        if self.db_path.exists():
            # Bereinigung: TuT entfernen & Schema Update
            try:
                with sqlite3.connect(self.db_path) as conn:
                    # 1. TuT Bereinigung
                    conn.execute("DELETE FROM noten WHERE fach_id IN (SELECT fach_id FROM faecher WHERE fach_kurz LIKE '%TuT%' OR fach_lang LIKE '%TuT%')")
                    conn.execute("DELETE FROM faecher WHERE fach_kurz LIKE '%TuT%' OR fach_lang LIKE '%TuT%'")
                    
                    # 2. Schema Erweiterung (Migration)
                    try:
                        conn.execute("ALTER TABLE noten ADD COLUMN lehrer_kuerzel TEXT")
                    except sqlite3.OperationalError:
                        # Spalte existiert bereits
                        pass
                        
            except Exception as e:
                logging.error(f"Fehler bei DB-Wartung: {e}")
            
            self.refresh_all_data()
        else:
            # Erstellt DB neu
            school_year, term = self._get_active_period()
            with KopfnotenImporter(str(self.db_path), school_year=school_year, term=term):
                pass 
            self.log_to_import(
                "Keine Datenbank gefunden (Neu erstellt). Bitte importieren Sie Daten."
            )

    def refresh_all_data(self):
        """Aktualisiert alle Daten"""
        try:
            self.load_classes_for_export()
            self.load_classes_for_analysis()
            self.refresh_analysis_data()
            self.refresh_insights_data()
            if TEMPLATE_MANAGER_ENABLED:
                self.refresh_template_list()
            self.status_manager.set_status("Daten aktualisiert")
        except Exception as e:
            logging.error(f"Fehler beim Aktualisieren: {e}")
            self.status_manager.set_status(f"Fehler: {e}")

    def load_classes_for_export(self):
        """Lädt Klassen für Export"""
        try:
            if not self.db_path.exists():
                return
            school_year, term = self._get_active_period()
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.execute(
                    """
                    SELECT DISTINCT s.klasse
                    FROM schueler s
                    JOIN noten n ON s.schueler_id = n.schueler_id
                    WHERE n.schuljahr = ? AND n.halbjahr = ?
                      AND COALESCE(s.is_active, 1) = 1
                    ORDER BY s.klasse
                    """,
                    (school_year, term),
                )
                classes = sorted([row[0] for row in cursor.fetchall()], key=self._class_sort_key)
                self.export_listbox.delete(0, tk.END)
                for class_name in classes:
                    self.export_listbox.insert(tk.END, class_name)
                self.log_to_export(f"{len(classes)} Klassen gefunden")
        except Exception as e:
            logging.error(f"Fehler beim Laden der Klassen: {e}")

    def load_classes_for_analysis(self):
        """Lädt Klassen für Analyse"""
        try:
            if not self.db_path.exists():
                return
            school_year, term = self._get_active_period()
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.execute(
                    """
                    SELECT DISTINCT s.klasse
                    FROM schueler s
                    JOIN noten n ON s.schueler_id = n.schueler_id
                    WHERE n.schuljahr = ? AND n.halbjahr = ?
                      AND COALESCE(s.is_active, 1) = 1
                    ORDER BY s.klasse
                    """,
                    (school_year, term),
                )
                class_values = sorted([row[0] for row in cursor.fetchall()], key=self._class_sort_key)
                classes = ["Alle"] + class_values
                self.class_filter["values"] = classes
                if classes:
                    self.class_filter.set(classes[0])
        except Exception as e:
            logging.error(f"Fehler beim Laden der Klassen für Analyse: {e}")

    def refresh_analysis_data(self, class_filter=None, student_filter=None):
        """Aktualisiert Analyse-Daten, optional gefiltert"""
        try:
            if not self.db_path.exists():
                return
            school_year, term = self._get_active_period()
            self.analysis_tree.delete(*self.analysis_tree.get_children())
            columns = [
                "ID",
                "Name",
                "Klasse",
                "Fächer",
                "AV-Noten",
                "SV-Noten",
                "Status",
                "SPH-Abgleich",
            ]
            self.analysis_tree["columns"] = columns
            self.analysis_tree.column("#0", width=0, stretch=False)
            self.analysis_tree.column("ID", width=0, stretch=False)
            for col in columns:
                self.analysis_tree.heading(col, text=col)
                if col == "Name":
                     self.analysis_tree.column(col, width=150)
                elif col in ["Status", "SPH-Abgleich"]:
                     self.analysis_tree.column(col, width=120)
                else:
                     self.analysis_tree.column(col, width=80)

            with sqlite3.connect(self.db_path) as conn:
                # 1. Rohdaten abrufen (Detailliert für Deduplizierung)
                query = """
                    SELECT
                        s.schueler_id, s.name, s.klasse,
                        s.target_subjects,
                        f.fach_kurz, f.fach_typ, f.fach_lang,
                        n.note_av, n.note_sv,
                        n.note_av_special, n.note_sv_special,
                        f.ist_wahlpflicht, n.ist_wahlpflicht_belegung,
                        f.wahlpflicht_gruppe,
                        n.lehrer_kuerzel
                    FROM schueler s
                    LEFT JOIN noten n ON s.schueler_id = n.schueler_id
                        AND n.schuljahr = ?
                        AND n.halbjahr = ?
                    LEFT JOIN faecher f ON n.fach_id = f.fach_id
                    WHERE COALESCE(s.is_active, 1) = 1
                    ORDER BY s.klasse, s.name
                """
                cursor = conn.execute(query, (school_year, term))
                rows = cursor.fetchall()

                # 2. Daten nach Schüler gruppieren und deduplizieren
                from collections import defaultdict
                student_map = {} # schueler_id -> data
                
                # Mapper: Klasse -> Fach (Canonical) -> Lehrer (Set, da mehrere möglich?)
                # Wir nehmen den letzten bekannten Lehrer oder sammeln alle.
                # Besser: Klasse -> Fach -> Teacher Set
                class_teacher_map = defaultdict(lambda: defaultdict(set))
                
                for row in rows:
                    s_id, s_name, s_klasse, s_target, f_kurz, f_typ, f_lang, av, sv, av_special, sv_special, f_wp, n_wp, wp_grp, lehrer_kuerzel = row
                    
                    if s_id not in student_map:
                        student_map[s_id] = {
                            "id": s_id, "name": s_name, "klasse": s_klasse,
                            "target_subjects_db": s_target,
                            "av_count": 0, "sv_count": 0,
                            "dedup_subjects": set(),
                            "teachers": set()
                        }
                    
                    sm = student_map[s_id]
                    if f_kurz or f_lang:
                        # STANDARDIZED CANONICAL NAME
                        fach_canonical = self._get_canonical_name(f_kurz, f_lang)
                        
                        # 2. Check Config (using CLEANED name)
                        config_status = SUBJECT_STATUS_CONFIG.get(fach_canonical, "")
                        is_wpu_config = "WPU" in config_status
                        
                        # DB Flags prüfen
                        is_wpu = bool(f_wp) or bool(n_wp) or any(p in (wp_grp or "") for p in ["WPU", "WP"]) or is_wpu_config
                        
                        # OVERRIDE: Prioritize explicit config (User Request for Praxistag Fix)
                        if config_status in ["Nebenfach", "Hauptfach"]:
                             is_wpu = False
                        
                        # Store raw data for later processing (we need to count first across all rows)
                        # Wir sammeln alle Fächer des Schülers erst in einer Liste
                        if "raw_subjects" not in sm:
                            sm["raw_subjects"] = []
                        
                        sm["raw_subjects"].append({
                            "f_kurz": f_kurz,
                            "f_lang": f_lang,
                            "f_typ": f_typ,
                            "av": av,
                            "sv": sv,
                            "av_special": av_special,
                            "sv_special": sv_special,
                            "is_wpu": is_wpu,
                            "f_canonical": fach_canonical
                        })
                        
                        if lehrer_kuerzel:
                             sm["teachers"].add(lehrer_kuerzel)
                             class_teacher_map[s_klasse][fach_canonical].add(lehrer_kuerzel)


                # 3. Post-Process per Student: Deduplizierung & WPU-Filterung
                for s_id, sm in student_map.items():
                    if "raw_subjects" not in sm: continue
                    
                    raw_list = sm["raw_subjects"]
                    sm["subjects_local"] = {}
                    
                    # Jahrgang ermitteln (aus Klasse)
                    klasse = sm["klasse"]
                    match = re.search(r'(\d+)', str(klasse))
                    jahrgang = int(match.group(1)) if match else 0
                    wpu_limit = 2 if jahrgang >= 9 else 1
                    
                    # A. Zähle WPU Noten und sammle benotete WPU Fächer
                    graded_wpu_subjects = []
                    for r in raw_list:
                         if r["is_wpu"] and (r["av"] is not None or r["sv"] is not None):
                             if r["f_canonical"] not in graded_wpu_subjects:
                                 graded_wpu_subjects.append(r["f_canonical"])
                    
                    # Limit anwenden
                    allowed_wpus = graded_wpu_subjects[:wpu_limit]
                    
                    for r in raw_list:
                         f_kurz = r["f_kurz"]
                         f_lang = r["f_lang"]
                         f_typ = r["f_typ"]
                         av = r["av"]
                         sv = r["sv"]
                         av_special = r.get("av_special")
                         sv_special = r.get("sv_special")
                         is_wpu = r["is_wpu"]
                         f_canonical = r["f_canonical"]
                         
                         is_rel_triad = (f_kurz == "Ethik") or (f_kurz == "Religion" and f_typ in ["evangelisch", "katholisch"])
                         
                         count_subject = True
                         
                         if is_wpu:
                             # Exclude logic:
                             if f_canonical not in allowed_wpus:
                                 count_subject = False
                             if av is None and sv is None and av_special is None and sv_special is None:
                                 count_subject = False
                         
                         # Name Cleaning für Display (User Request: "Praxistag WU" -> "Praxistag", "Chemie (U1)" -> "Chemie")
                         # Wir entfernen "WU", "WP" am Ende oder als Wort.
                         # Auch (U1), (U 2), (U1) entfernen
                         clean_name = f_canonical
                         clean_name = re.sub(r'\s+WU$', '', clean_name)
                         clean_name = re.sub(r'\s+WP$', '', clean_name)
                         clean_name = re.sub(r'\s*\(U\s*\d+\)', '', clean_name)
                         clean_name = clean_name.strip()
                         
                         if count_subject:
                             if is_rel_triad:
                                 sm["dedup_subjects"].add("REL_TRIAD")
                             elif is_wpu:
                                 # Use cleaned canonical name + WPU prefix for safety unique keying
                                 # This ensures distinct subjects are distinct, but merges identicals.
                                 sm["dedup_subjects"].add(f"WPU_{clean_name}")
                             else:
                                 # Using cleaned canonical name instead of f_kurz/f_typ tuple 
                                 # to avoid splitting same subject with diff abbreviations
                                 # But we keep f_typ if needed (Rel/Eth)? 
                                 # Rel/Eth is handled by REL_TRIAD.
                                 # For others: "Deutsch" is "Deutsch".
                                 sm["dedup_subjects"].add(clean_name)

                         # Für SPH-Abgleich: nur gezählte/relevante Fächer (keine ausgeschlossenen WPU-Leichen)
                         if count_subject:
                             local_subject_name = (f_lang or f_kurz or "").strip()
                             if local_subject_name:
                                 if local_subject_name not in sm["subjects_local"]:
                                     sm["subjects_local"][local_subject_name] = {"av": False, "sv": False}
                                 if av is not None or av_special is not None:
                                     sm["subjects_local"][local_subject_name]["av"] = True
                                 if sv is not None or sv_special is not None:
                                     sm["subjects_local"][local_subject_name]["sv"] = True
                         
                         if av is not None or av_special is not None: sm["av_count"] += 1
                         if sv is not None or sv_special is not None: sm["sv_count"] += 1
                
                # Pro Klasse: wie viele Lernende je Fach bereits Noten haben (für SPH-Abgleich)
                class_subject_grade_counts = defaultdict(lambda: defaultdict(int))
                for sm in student_map.values():
                    for subj, vals in sm.get("subjects_local", {}).items():
                        if vals.get("av") or vals.get("sv"):
                            subj_key = self._normalize_subject_for_sph(subj)
                            class_subject_grade_counts[sm["klasse"]][subj_key] += 1

                # In Klassen gruppieren für Referenzwert-Berechnung
                class_students = defaultdict(list)
                for sm in student_map.values():
                    # Religion/Ethik Gruppe immer zusammenfassen
                    # Falls REL_TRIAD existiert, ist es bereits ein Eintrag
                    # Wir stellen sicher, dass alle Rel/Eth Einträge (auch falls fälschlich mehrfach vorhanden) als einer zählen.
                    sm["faecher_count"] = len(sm["dedup_subjects"])
                    class_students[sm["klasse"]].append(sm)

                # 3. Status berechnen und in Treeview einfügen (nach Klasse sortiert)
                for klasse in sorted(class_students.keys(), key=self._class_sort_key):
                    students = class_students[klasse]
                    
                    # Schwellenwert: Maximalanzahl der Noten in dieser Klasse (AV + SV)
                    max_notes_in_class = 0
                    for s in students:
                        notes_total = s["av_count"] + s["sv_count"]
                        if notes_total > max_notes_in_class:
                            max_notes_in_class = notes_total
                    
                    # Jahrgang aus Klasse extrahieren
                    match_year = re.match(r"(\d+)", klasse)
                    jahrgang = int(match_year.group(1)) if match_year else 0

                    for s in students:
                        # FILTER LOGIK
                        show_student = True
                        if class_filter and class_filter != "Alle":
                             if s["klasse"] != class_filter:
                                 show_student = False
                        
                        if student_filter:
                             if student_filter.lower() not in s["name"].lower():
                                 show_student = False
                        
                        if not show_student:
                            continue

                        current_notes = s["av_count"] + s["sv_count"]
                        
                        # ZIELWERT ERMITTELN
                        # Entweder DB-Wert oder Default für Jahrgang
                        final_target = s["target_subjects_db"] if s["target_subjects_db"] else self.get_default_target_for_grade(jahrgang)
                        
                        status = self._calculate_status(jahrgang, current_notes, s["faecher_count"], final_target)
                        sph_alignment, row_tag, sph_status_override = self._get_sph_alignment_for_student(
                            s["klasse"],
                            s.get("subjects_local", {}),
                            status,
                            dict(class_subject_grade_counts.get(s["klasse"], {})),
                        )
                        if sph_status_override:
                            status = sph_status_override

                        # STATUS FILTER
                        status_filter_val = self.status_filter_var.get()
                        if status_filter_val != "Alle":
                             if status != status_filter_val:
                                 show_student = False

                        # TEACHER FILTER
                        teacher_filter_val = self.teacher_filter_var.get().strip().lower()
                        if teacher_filter_val:
                             # Logic Refinement:
                             # If Status == Unvollständig (detected by status var) AND Teacher Filter Set:
                             # Show ONLY if missing grades are from THIS teacher.
                             
                             if status == "Unvollständig":
                                 # 1. Welche Fächer fehlen?
                                 # Wir müssen wissen, welche 'regular subjects' der Jahrgang hat und was der Schüler hat.
                                 # Das ist etwas teuer hier, wir machen es on demand.
                                 
                                 # Wir haben dedup_subjects (IST-Fächer).
                                 # Wir rufen _get_class_regular_subjects ab (Per Class Logic)
                                 # Class subjects
                                 jg_subjects_meta = self._get_class_regular_subjects(s["klasse"]) # Dict[name, id]
                                 jg_subject_names = set(jg_subjects_meta.keys())
                                 

                                 
                                 
                                 # SIMPLIFIED LOGIC (Compromise):
                                 # Show if student is incomplete AND teacher has ANY relevance to this student (Active or Potential).
                                 
                                 is_relevant_teacher = False
                                 
                                 # 1. Active: Teacher has already given a grade
                                 for t in s["teachers"]:
                                     if teacher_filter_val in t.lower():
                                         is_relevant_teacher = True
                                         break
                                         
                                 # 2. Potential: Teacher teaches a subject this class takes
                                 if not is_relevant_teacher:
                                     # Check class map
                                     class_map = class_teacher_map.get(s["klasse"], {})
                                     for subj_name, teachers in class_map.items():
                                         # Is this subject relevant for the student?
                                         # Check against regular subjects or WPU
                                         is_regular = subj_name in jg_subject_names
                                         
                                         # Simplified WPU check: simple existence in map is strong hint
                                         # But strictly we should check if it's a "class subject". 
                                         # jg_subject_names (now loaded via _get_class_regular_subjects) covers all NON-WPU subjects found in the class.
                                         # For WPU, we might need to be broader or check class_wpus.
                                         
                                         if is_regular:
                                              for t in teachers:
                                                  if teacher_filter_val in t.lower():
                                                      is_relevant_teacher = True
                                                      break
                                         if is_relevant_teacher: break
                                         
                                         # Check if WPU (if not regular)
                                         if not is_regular:
                                              # If teacher matches a WPU subject in this class, we assume potential relevance
                                              # (Showing too many is better than too few here)
                                              for t in teachers:
                                                  if teacher_filter_val in t.lower():
                                                       is_relevant_teacher = True
                                                       break
                                         if is_relevant_teacher: break

                                 if not is_relevant_teacher:
                                     show_student = False
                             
                             else:
                                 # Standard Logic (Show if ANY grade from teacher)
                                 has_match = False
                                 for t in s["teachers"]:
                                     if teacher_filter_val in t.lower():
                                         has_match = True
                                         break
                                 if not has_match:
                                     show_student = False
                        
                        if not show_student:
                            continue

                        # Zeile einfügen
                        self.analysis_tree.insert(
                            "",
                            tk.END,
                            values=(
                                s["id"],
                                s["name"],
                                s["klasse"],
                                final_target, # Zeige SOLL anstatt IST (User Request)
                                s["av_count"],
                                s["sv_count"],
                                status,
                                sph_alignment
                            ),
                            tags=(row_tag,) if row_tag else ()
                        )



        except Exception as e:
            logging.error(f"Fehler beim Aktualisieren der Analyse-Daten: {e}")
            import traceback
            traceback.print_exc()

    def _normalize_class_for_sph(self, klasse: str) -> str:
        """Normalisiert Klassenkennung für SPH-Map (z. B. 05a -> 05A, daz2 -> DAZ2)."""
        if not klasse:
            return ""
        k = str(klasse).strip()
        m = re.match(r"^(\d{1,2})([a-zA-Z])$", k)
        if m:
            return f"{int(m.group(1)):02d}{m.group(2).upper()}"
        if k.lower().startswith("daz"):
            return k.upper()
        return k.upper()

    def _normalize_subject_for_sph(self, text: str) -> str:
        """Normalisiert Fachbezeichnungen für SPH-Abgleich."""
        raw = (text or "").strip()
        if raw in FAECHER_MAPPING:
            raw = FAECHER_MAPPING[raw]
        raw = raw.lower()
        raw = re.sub(r"\b\d{1,2}[a-z]\b", " ", raw)
        raw = re.sub(r"\bdaz\d+\b", " ", raw)
        raw = re.sub(r"\(.*?\)", " ", raw)
        raw = re.sub(r"[^a-z0-9äöüß]+", " ", raw)
        return re.sub(r"\s+", " ", raw).strip()

    def _subjects_match_for_sph(self, local_subject: str, sph_subject: str) -> bool:
        a = self._normalize_subject_for_sph(local_subject)
        b = self._normalize_subject_for_sph(sph_subject)
        if not a or not b:
            return False
        return a == b

    def _find_local_subject_for_sph(
        self, sph_subject: str, subjects_local: Dict[str, Dict[str, bool]]
    ):
        """Findet das lokale Fach, das zu einer SPH-Lerngruppe passt."""
        for local_subject, vals in subjects_local.items():
            if self._subjects_match_for_sph(local_subject, sph_subject):
                return local_subject, vals
        return None, None

    def _sph_row_is_multi_class(self, row: Dict[str, Any]) -> bool:
        """True, wenn die SPH-Lerngruppe mehrere Klassen kombiniert (z. B. 05c/05d)."""
        lerngruppe = row.get("lerngruppe", "") or ""
        matches = re.findall(r"\b(\d{2}[a-z]|daz\d+)\b", lerngruppe, re.IGNORECASE)
        return len(matches) > 1

    def _get_sph_alignment_for_student(
        self,
        klasse: str,
        subjects_local: Dict[str, Dict[str, bool]],
        local_status: str,
        class_subject_graded: Optional[Dict[str, int]] = None,
    ):
        """
        SPH-Abgleich je Lernendem.
        Rückgabe: (anzeigetext, row_tag, status_override)

        „Kurs fehlt“ nur, wenn das Fach beim Lernenden gar nicht vorkommt,
        SPH rot meldet und andere Lernende der Klasse dafür Noten haben.
        Fehlende Noten bei vorhandenem Fach → „Einzelnoten fehlen“.
        """
        if local_status == "Vollständig":
            return "SPH: vollständig", "student_green", "Vollständig"

        cls_key = self._normalize_class_for_sph(klasse)
        data = self.sph_missing_overview.get(cls_key)
        if not data:
            return "Kein SPH-Abgleich", None, None

        rows = data.get("rows", [])
        if not rows:
            return "Kein SPH-Abgleich", None, None

        student_has_any_grade = any(
            bool(vals.get("av")) or bool(vals.get("sv")) for vals in subjects_local.values()
        )
        if not student_has_any_grade:
            return "SPH: Einzelnoten fehlen", "student_yellow", "Unvollständig"

        has_red = False
        has_yellow = False
        has_any_complete = False
        class_counts = class_subject_graded or {}
        matched_any = False

        for row in rows:
            sph_subject = row.get("fach_raw", "")
            color = row.get("farbe")
            local_subject, vals = self._find_local_subject_for_sph(sph_subject, subjects_local)

            if local_subject:
                matched_any = True
                av_ok = bool(vals.get("av"))
                sv_ok = bool(vals.get("sv"))
                complete = av_ok and sv_ok
                partial_missing = av_ok != sv_ok
                full_missing = not av_ok and not sv_ok

                if complete:
                    has_any_complete = True
                elif partial_missing or full_missing or color in ("rot", "gelb"):
                    # Fach ist importiert → fehlende Noten, kein fehlender Kurs.
                    has_yellow = True
                continue

            # Kein lokales Fach: nur dann „Kurs fehlt“ prüfen.
            if color != "rot" or self._sph_row_is_multi_class(row):
                continue

            subject_key = self._normalize_subject_for_sph(sph_subject)
            if class_counts.get(subject_key, 0) > 0:
                has_red = True
                matched_any = True

        if not matched_any:
            return "SPH: Einzelnoten fehlen", "student_yellow", "Unvollständig"

        if has_red:
            return "SPH: Kurs fehlt", "student_red", "Unvollständig"
        if has_yellow or local_status == "Unvollständig":
            return "SPH: Einzelnoten fehlen", "student_yellow", "Unvollständig"
        if has_any_complete:
            return "SPH: vollständig", "student_green", "Vollständig"

        return "SPH: vollständig", "student_green", "Vollständig"

    def _calculate_status(self, jahrgang: int, current_notes: int, faecher_count: int, target_subjects: int = None) -> str:
        """
        Berechnet Status
        target_subjects: Expliziter Sollwert (aus DB oder Default). Falls gesetzt, wird er genutzt.
        """
        # Default Logik (falls target_subjects nicht übergeben - sollte aber jetzt immer sein)
        if target_subjects and target_subjects > 0:
            target_subj = target_subjects
        else:
            # Fallback
            target_subj = self.get_default_target_for_grade(jahrgang)
        
        is_complete = False
        
        if target_subj > 0:
            # Wir prüfen, ob GENUG Noten da sind (Anzahl Noten >= Ziel * 2)
            # ACHTUNG: Der User will Sollwert.
            # Wenn Soll = 13, dann erwarte ich 26 Noten.
            # Aber wir haben auch faecher_count (IST).
            # Wenn IST < SOLL => Unvollständig? Ja.
            # Wenn IST >= SOLL und Noten >= SOLL*2 => Vollständig
            
            # Simple Check: Habe ich genug Noten für das Ziel?
            if current_notes >= (target_subj * 2):
                 is_complete = True
        else:
             # Fallback logic if Jahrgang is missing or other
             if faecher_count > 0 and current_notes >= (faecher_count * 2):
                 is_complete = True # Simple check
        
        return "Vollständig" if is_complete else "Unvollständig"

    def get_default_target_for_grade(self, jahrgang: int) -> int:
        """Liefert den Standard-Zielwert für die Fächeranzahl pro Jahrgang"""
        if jahrgang in [5, 6]: return 9
        elif jahrgang == 7: return 10
        elif jahrgang == 8: return 11
        elif jahrgang == 9: return 13
        elif jahrgang == 10: return 13
        return 0 if is_complete else "Unvollständig"


    def search_students(self, event=None):
        """Sucht Schüler mit Debounce"""
        if self.student_search_after:
            self.root.after_cancel(self.student_search_after)
            
        def run_search():
            class_filter = self.class_filter.get()
            student_filter = self.student_search.get().strip()
            self.refresh_analysis_data(class_filter, student_filter)
            self.status_manager.set_status(f"Filter angewendet: {class_filter}, Suche: {student_filter}")

        self.student_search_after = self.root.after(300, run_search)

    def reset_filters(self):
        """Setzt Filter zurück"""
        self.class_filter.set("Alle")
        self.student_search.delete(0, tk.END)
        self.teacher_filter_var.set("")
        self.status_filter_var.set("Alle")
        self.refresh_analysis_data()

    def deactivate_selected_student(self):
        """Deaktiviert ausgewählten Schüler (persistiert über Re-Importe)."""
        selection = self.analysis_tree.selection()
        if not selection:
            messagebox.showwarning(
                "Keine Auswahl", "Bitte wählen Sie einen Schüler aus."
            )
            return
        item = self.analysis_tree.item(selection[0])
        values = item["values"]
        student_id = values[0]
        student_name = values[1]

        if messagebox.askyesno(
            "Deaktivierung bestätigen",
            f"Lernende/r '{student_name}' wird deaktiviert und in Listen/Exporten ausgeblendet.\n\nFortfahren?",
        ):
            try:
                with sqlite3.connect(self.db_path) as conn:
                    conn.execute(
                        "UPDATE schueler SET is_active = 0 WHERE schueler_id = ?",
                        (student_id,),
                    )
                    conn.commit()
                    messagebox.showinfo(
                        "Deaktiviert",
                        f"Lernende/r '{student_name}' wurde deaktiviert.\n"
                        "Die Deaktivierung bleibt auch nach Neuimporten erhalten.",
                    )
                    self.refresh_all_data()
            except Exception as e:
                logging.error(f"Fehler beim Deaktivieren: {e}")
                messagebox.showerror("Deaktivierungs-Fehler", f"Fehler: {e}")

    def manage_inactive_students(self):
        """Zeigt deaktivierte Lernende und erlaubt Reaktivierung."""
        try:
            with sqlite3.connect(self.db_path) as conn:
                rows = conn.execute(
                    """
                    SELECT schueler_id, name, klasse
                    FROM schueler
                    WHERE COALESCE(is_active, 1) = 0
                    ORDER BY klasse, name
                    """
                ).fetchall()
        except Exception as e:
            messagebox.showerror("Fehler", f"Deaktivierte Lernende konnten nicht geladen werden:\n{e}")
            return

        if not rows:
            messagebox.showinfo("Deaktivierte Lernende", "Es sind aktuell keine Lernenden deaktiviert.")
            return

        dlg = tk.Toplevel(self.root)
        dlg.title("Deaktivierte Lernende verwalten")
        dlg.geometry("520x420")
        dlg.transient(self.root)
        dlg.grab_set()

        ttk.Label(
            dlg,
            text="Ausgewählte Lernende reaktivieren",
            font=("Arial", 11, "bold"),
        ).pack(anchor=tk.W, padx=10, pady=(10, 6))

        frame = ttk.Frame(dlg)
        frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=(0, 10))
        tree = ttk.Treeview(frame, columns=("id", "name", "klasse"), show="headings")
        tree.heading("id", text="ID")
        tree.heading("name", text="Name")
        tree.heading("klasse", text="Klasse")
        tree.column("id", width=60, anchor=tk.CENTER)
        tree.column("name", width=300)
        tree.column("klasse", width=100, anchor=tk.CENTER)
        tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        sb = ttk.Scrollbar(frame, orient=tk.VERTICAL, command=tree.yview)
        tree.configure(yscrollcommand=sb.set)
        sb.pack(side=tk.RIGHT, fill=tk.Y)

        for sid, name, klasse in rows:
            tree.insert("", tk.END, values=(sid, name, klasse))

        def reactivate_selected():
            selected = tree.selection()
            if not selected:
                messagebox.showwarning("Keine Auswahl", "Bitte mindestens eine Person auswählen.")
                return
            ids = [tree.item(i)["values"][0] for i in selected]
            try:
                with sqlite3.connect(self.db_path) as conn:
                    conn.executemany(
                        "UPDATE schueler SET is_active = 1 WHERE schueler_id = ?",
                        [(sid,) for sid in ids],
                    )
                    conn.commit()
                dlg.destroy()
                self.refresh_all_data()
                messagebox.showinfo("Reaktiviert", f"{len(ids)} Lernende wurden reaktiviert.")
            except Exception as e:
                messagebox.showerror("Fehler", f"Reaktivierung fehlgeschlagen:\n{e}")

        btns = ttk.Frame(dlg)
        btns.pack(fill=tk.X, padx=10, pady=(0, 10))
        ttk.Button(btns, text="Auswahl reaktivieren", command=reactivate_selected).pack(side=tk.LEFT)
        ttk.Button(btns, text="Schließen", command=dlg.destroy).pack(side=tk.RIGHT)

    # ===================== HILFS-FUNKTIONEN =====================

    def open_database(self):
        """Öffnet Datenbank"""
        filename = filedialog.askopenfilename(
            title="Datenbank öffnen",
            filetypes=[("SQLite-Datenbank", "*.db"), ("Alle Dateien", "*.*")],
            initialdir=str(self.paths.database_path.parent.resolve()),
        )
        if filename:
            self.db_path = Path(filename)
            self.refresh_all_data()
            messagebox.showinfo(
                "Datenbank geöffnet", f"Datenbank geöffnet: {self.db_path.name}"
            )

    def _validate_database_schema(self, db_file: Path) -> Tuple[bool, str]:
        """Prüft, ob eine Datei eine kompatible SQLite-Datenbank ist."""
        try:
            if not db_file.exists() or not db_file.is_file():
                return False, "Datei wurde nicht gefunden."
            if db_file.stat().st_size <= 0:
                return False, "Datei ist leer."

            with sqlite3.connect(str(db_file)) as conn:
                conn.execute("PRAGMA quick_check")
                cursor = conn.execute(
                    "SELECT name FROM sqlite_master WHERE type = 'table'"
                )
                tables = {row[0] for row in cursor.fetchall()}
            required = {"schueler", "faecher", "noten"}
            missing = sorted(required - tables)
            if missing:
                return (
                    False,
                    "Inkompatible Datenbank. Fehlende Tabellen: " + ", ".join(missing),
                )
            return True, ""
        except sqlite3.Error as e:
            return False, f"Ungültige SQLite-Datei: {e}"
        except Exception as e:
            return False, f"Validierung fehlgeschlagen: {e}"

    def export_database_file(self):
        """Exportiert die aktive Datenbank als .db-Datei."""
        if not self.db_path.exists():
            messagebox.showwarning("Keine Datenbank", "Es gibt keine aktive Datenbank zum Exportieren.")
            return

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        default_name = f"{self.db_path.stem}_export_{ts}.db"
        target = filedialog.asksaveasfilename(
            title="Datenbank exportieren",
            defaultextension=".db",
            initialfile=default_name,
            filetypes=[("SQLite-Datenbank", "*.db"), ("Alle Dateien", "*.*")],
            initialdir=str(self.paths.database_path.parent.resolve()),
        )
        if not target:
            return

        target_path = Path(target)
        if target_path.suffix.lower() != ".db":
            target_path = target_path.with_suffix(".db")
        try:
            self.path_manager.ensure_directory(target_path.parent)
            shutil.copy2(self.db_path, target_path)
            self._save_db_transfer_meta(
                last_export_path=str(target_path),
                last_export_time=datetime.now().isoformat(timespec="seconds"),
            )
            self.status_manager.set_status("Datenbank-Export erfolgreich")
            messagebox.showinfo(
                "Export erfolgreich",
                "Die Datenbank wurde erfolgreich exportiert.\n\n"
                "Hinweis: Diese Datei kann zentral abgelegt und auf anderen Rechnern importiert werden.\n\n"
                f"Export-Datei:\n{target_path}",
            )
        except Exception as e:
            messagebox.showerror("Exportfehler", f"Datenbank-Export fehlgeschlagen:\n{e}")

    def import_database_file(self):
        """Importiert eine .db-Datei und ersetzt die lokale aktive Datenbank."""
        source = filedialog.askopenfilename(
            title="Datenbank importieren",
            filetypes=[("SQLite-Datenbank", "*.db"), ("Alle Dateien", "*.*")],
            initialdir=str(self.paths.import_dir.resolve()),
        )
        if not source:
            return

        source_path = Path(source)
        is_valid, reason = self._validate_database_schema(source_path)
        if not is_valid:
            messagebox.showerror("Importfehler", reason)
            return

        proceed = messagebox.askyesno(
            "Import bestätigen",
            "Die ausgewählte Datenbank ersetzt die lokale aktive Datenbank vollständig.\n"
            "Diese Aktion ist für zentrale Ablage und Nutzung auf mehreren Rechnern gedacht.\n\n"
            "Möchten Sie fortfahren?",
        )
        if not proceed:
            self.status_manager.set_status("Datenbank-Import abgebrochen")
            return

        target_db = Path(self.paths.database_path)
        backup_path = None
        temp_path = target_db.with_suffix(".import_tmp.db")

        try:
            self.path_manager.ensure_directory(target_db.parent)
            self.path_manager.ensure_directory(self.paths.backup_dir)

            if target_db.exists():
                ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                backup_path = self.paths.backup_dir / f"{target_db.stem}_preimport_{ts}.db"
                shutil.copy2(target_db, backup_path)

            shutil.copy2(source_path, temp_path)
            temp_valid, temp_reason = self._validate_database_schema(temp_path)
            if not temp_valid:
                raise ValueError(temp_reason)

            os.replace(str(temp_path), str(target_db))
            self.db_path = target_db
            self._save_db_transfer_meta(
                last_import_path=str(source_path),
                last_import_time=datetime.now().isoformat(timespec="seconds"),
            )
            self.load_initial_data()
            self.refresh_all_data()
            self.status_manager.set_status("Datenbank-Import erfolgreich")
            messagebox.showinfo(
                "Import erfolgreich",
                f"Datenbank wurde erfolgreich importiert.\n\n"
                f"Import-Quelle:\n{source_path}\n\n"
                f"Aktive Datenbank:\n{target_db}\n\n"
                + (f"Sicherungsdatei: {backup_path}" if backup_path else "Es war keine vorherige Datenbank vorhanden."),
            )
        except Exception as e:
            # Rollback wenn möglich
            try:
                if backup_path and backup_path.exists():
                    shutil.copy2(backup_path, target_db)
                    self.db_path = target_db
                    self.load_initial_data()
            except Exception as rollback_error:
                logging.error(f"Rollback nach Importfehler fehlgeschlagen: {rollback_error}")
            messagebox.showerror("Importfehler", f"Datenbank-Import fehlgeschlagen:\n{e}")
        finally:
            try:
                if temp_path.exists():
                    temp_path.unlink()
            except Exception:
                pass

    def show_database_info(self):
        """Zeigt Datenbank-Informationen"""
        if not self.db_path.exists():
            messagebox.showwarning("Keine Datenbank", "Keine Datenbank gefunden.")
            return
        try:
            school_year, term = self._get_active_period()
            with sqlite3.connect(self.db_path) as conn:
                schueler_count = conn.execute(
                    "SELECT COUNT(*) FROM schueler WHERE COALESCE(is_active, 1) = 1"
                ).fetchone()[0]
                schueler_inactive_count = conn.execute(
                    "SELECT COUNT(*) FROM schueler WHERE COALESCE(is_active, 1) = 0"
                ).fetchone()[0]
                klassen_count = conn.execute(
                    "SELECT COUNT(DISTINCT klasse) FROM schueler WHERE COALESCE(is_active, 1) = 1"
                ).fetchone()[0]
                faecher_count = conn.execute("SELECT COUNT(*) FROM faecher").fetchone()[
                    0
                ]
                noten_count = conn.execute(
                    "SELECT COUNT(*) FROM noten WHERE schuljahr = ? AND halbjahr = ?",
                    (school_year, term),
                ).fetchone()[0]
                db_size = self.db_path.stat().st_size / (1024 * 1024)
                transfer_meta = self._load_db_transfer_meta()
                last_import_path = transfer_meta.get("last_import_path", "-")
                last_import_time = transfer_meta.get("last_import_time", "-")
                last_export_path = transfer_meta.get("last_export_path", "-")
                last_export_time = transfer_meta.get("last_export_time", "-")
                info_text = f"""Datenbank-Informationen:

Datei: {self.db_path.name}
Größe: {db_size:.2f} MB
Inhalt:
• Schüler: {schueler_count}
• Deaktiviert: {schueler_inactive_count}
• Klassen: {klassen_count}
• Fächer: {faecher_count}
• Noten: {noten_count}

Letzte Übertragung:
• Letzter Import: {last_import_time}
• Import-Quelle: {last_import_path}
• Letzter Export: {last_export_time}
• Export-Ziel: {last_export_path}"""
                messagebox.showinfo("Datenbank-Info", info_text)
        except Exception as e:
            messagebox.showerror("Datenbankfehler", f"Fehler: {e}")

    def backup_database(self):
        """Erstellt eine Zeitstempel-Sicherung der aktiven Datenbank."""
        if not self.db_path.exists():
            messagebox.showwarning("Keine Datenbank", "Es gibt keine Datenbank zum Sichern.")
            return

        try:
            self.path_manager.ensure_directory(self.paths.backup_dir)
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup_file = self.paths.backup_dir / f"{self.db_path.stem}_{ts}.db"
            shutil.copy2(self.db_path, backup_file)
            messagebox.showinfo(
                "Sicherung erstellt",
                f"Datenbank wurde gesichert:\n{backup_file}",
            )
        except Exception as e:
            messagebox.showerror("Sicherungsfehler", f"Datenbank-Sicherung fehlgeschlagen:\n{e}")

    def show_logs(self):
        """Zeigt Logs"""
        log_dir = self.paths.logs_dir
        if not log_dir.exists():
            messagebox.showinfo("Keine Logs", "Keine Log-Dateien gefunden.")
            return
        log_files = list(log_dir.glob("*.log"))
        if not log_files:
            messagebox.showinfo("Keine Logs", "Keine Log-Dateien gefunden.")
            return
        latest_log = max(log_files, key=lambda f: f.stat().st_mtime)
        try:
            with open(latest_log, "r", encoding="utf-8") as f:
                content = f.read()
                log_window = tk.Toplevel(self.root)
                log_window.title(f"Logs - {latest_log.name}")
                log_window.geometry("800x600")
                log_text = scrolledtext.ScrolledText(
                    log_window, wrap=tk.WORD, font=("Courier", 9)
                )
                log_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
                log_text.insert(tk.END, content)
                log_text.config(state=tk.DISABLED)
                log_text.see(tk.END)
                ttk.Button(log_window, text="Schließen", command=log_window.destroy).pack(
                    pady=5
                )
        except Exception as e:
            messagebox.showerror("Log-Fehler", f"Fehler: {e}")

    def check_permissions(self):
        """Prüft Berechtigungen"""
        try:
            dirs = [
                self.paths.templates_dir,
                self.paths.output_word_dir,
                self.paths.database_path.parent,
                self.paths.logs_dir,
            ]
            report = []
            for directory in dirs:
                if directory.exists():
                    readable = os.access(directory, os.R_OK)
                    writable = os.access(directory, os.W_OK)
                    executable = os.access(directory, os.X_OK)
                    status = "✅" if (readable and writable and executable) else "❌"
                    report.append(
                        f"{status} {directory.name}: R:{readable} W:{writable} X:{executable}"
                    )
                else:
                    report.append(f"❓ {directory.name}: Existiert nicht")
            messagebox.showinfo(
                "Berechtigungen", "Berechtigungsprüfung:\n\n" + "\n".join(report)
            )
        except Exception as e:
            messagebox.showerror("Prüfungsfehler", f"Fehler: {e}")

    def show_about(self):
        """Zeigt Über-Dialog"""
        about_text = f"""Kopfnoten-Manager

Entwickelt für IGS in Hessen

GitHub: {GITHUB_REPO_URL}

© Jörg Pospischil 2026"""
        messagebox.showinfo("Über", about_text)
        
    def _get_canonical_name(self, f_kurz: str, f_lang: str) -> str:
        """Standardisierte Logik für Fach-Namen"""
        # 1. Canonical name lookup (Prefer Kurz mapping, fallback to Lang)
        # Check if f_kurz is in mapping
        if f_kurz in FAECHER_MAPPING:
            base_name = FAECHER_MAPPING[f_kurz]
        else:
            base_name = f_lang if f_lang else f_kurz
            
        # CLEANING: Suffixes entfernen
        clean_name = base_name
        if clean_name:
            clean_name = re.sub(r'\s+WU$', '', clean_name)
            clean_name = re.sub(r'\s+WP$', '', clean_name)
            clean_name = re.sub(r'\s*\(U\s*\d+\)', '', clean_name)
            clean_name = clean_name.strip()
            
        return clean_name

    def _get_class_wpu_subjects(self, student_class: str) -> List[Dict]:
        """Holt ALLE WPU-Fächer, die in dieser Klasse existieren"""
        try:
            school_year, term = self._get_active_period()
            with sqlite3.connect(self.db_path) as conn:
                conn.row_factory = sqlite3.Row
                # Finde alle Fächer, die eine 'ist_wahlpflicht' Flag oder WPU Gruppe haben
                # UND die von Schülern dieser Klasse belegt sind
                cursor = conn.execute(
                    """
                    SELECT DISTINCT
                        f.fach_id,
                        f.fach_lang,
                        f.fach_kurz,
                        f.ist_wahlpflicht,
                        f.wahlpflicht_gruppe,
                         -- Wir holen ein beispielhaftes Lehrerkürzel (wenn vorhanden) für dieses Fach in dieser Klasse
                        (SELECT lehrer_kuerzel FROM noten n2 
                         JOIN schueler s2 ON n2.schueler_id = s2.schueler_id 
                         WHERE n2.fach_id = f.fach_id AND s2.klasse = ? 
                           AND n2.schuljahr = ? AND n2.halbjahr = ?
                          AND COALESCE(s2.is_active, 1) = 1
                         ORDER BY n2.lehrer_kuerzel DESC LIMIT 1) as lehrer_kuerzel
                    FROM faecher f
                    JOIN noten n ON f.fach_id = n.fach_id
                    JOIN schueler s ON n.schueler_id = s.schueler_id
                    WHERE s.klasse = ?
                    AND n.schuljahr = ? AND n.halbjahr = ?
                    AND COALESCE(s.is_active, 1) = 1
                    AND (f.ist_wahlpflicht = 1 OR f.wahlpflicht_gruppe LIKE '%WPU%' OR f.wahlpflicht_gruppe LIKE '%WP%')
                    """,
                    (student_class, school_year, term, student_class, school_year, term),
                )
                return [dict(row) for row in cursor.fetchall()]
        except Exception as e:
            logging.error(f"Fehler beim Laden der Klassen-WPU-Fächer: {e}")
            return []

    def _get_class_regular_subjects(self, student_class: str) -> Dict[str, int]:
        """Ermittelt alle regulären Fächer (kein WPU) für eine spezifische Klasse.
           Returns: Dict[Fachname (Canonical), FachID]"""
        regular_subjects = {}
        if not student_class:
            return {}
            
        try:
            school_year, term = self._get_active_period()
            with sqlite3.connect(self.db_path) as conn:
                conn.row_factory = sqlite3.Row
                
                # Query: Finde alle Fächer, die in DIESER Klasse vorkommen
                query = """
                    SELECT DISTINCT f.fach_id, f.fach_lang, f.fach_kurz, f.ist_wahlpflicht, f.wahlpflicht_gruppe
                    FROM faecher f
                    JOIN noten n ON f.fach_id = n.fach_id
                    JOIN schueler s ON n.schueler_id = s.schueler_id
                    WHERE s.klasse = ?
                      AND n.schuljahr = ?
                      AND n.halbjahr = ?
                      AND COALESCE(s.is_active, 1) = 1
                """
                cursor = conn.execute(query, (student_class, school_year, term))
                rows = cursor.fetchall()
                
                for row in rows:
                    fach_lang = row["fach_lang"]
                    fach_kurz = row["fach_kurz"]
                    fach_id = row["fach_id"]
                    
                    fach_canonical = self._get_canonical_name(fach_kurz, fach_lang)
                    
                    # Check Config first
                    config_status = SUBJECT_STATUS_CONFIG.get(fach_canonical, "")
                    is_wpu_config = "WPU" in config_status
                    
                    # Check DB flags
                    is_wpu_db = bool(row["ist_wahlpflicht"]) or (row["wahlpflicht_gruppe"] and ("WPU" in row["wahlpflicht_gruppe"] or "WP" in row["wahlpflicht_gruppe"]))
                    
                    # Entscheidung: Ist es regulär?
                    is_regular = False
                    if config_status in ["Hauptfach", "Nebenfach"]:
                        is_regular = True # Override: Config sagt regulär (z.B. Praxistag)
                    elif is_wpu_config:
                        is_regular = False # Override: Config sagt WPU
                    else:
                        is_regular = not is_wpu_db # Fallback auf DB
                    
                    if is_regular:
                        regular_subjects[fach_canonical] = fach_id # Store ID keyed by CANONICAL NAME
                        
        except Exception as e:
            logging.error(f"Fehler beim Laden der Klassenfächer: {e}")
            
        return regular_subjects

    def _get_year_regular_subjects(self, jahrgang: int) -> Dict[str, int]:
        """Ermittelt alle regulären Fächer (kein WPU) für einen kompletten Jahrgang.
           Returns: Dict[Fachname, FachID]"""
        regular_subjects = {}
        if not jahrgang:
            return {}
            
        try:
            school_year, term = self._get_active_period()
            with sqlite3.connect(self.db_path) as conn:
                conn.row_factory = sqlite3.Row
                # Patterns für Jahrgangssuche (z.B. 9% für 9a, 09% für 09a)
                patterns = [f"{jahrgang}%", f"{jahrgang:02d}%"]
                
                # Query: Finde alle Fächer, die in diesem Jahrgang vorkommen
                # Wir holen fach_id dazu
                query = """
                    SELECT DISTINCT f.fach_id, f.fach_kurz, f.fach_lang, f.ist_wahlpflicht, f.wahlpflicht_gruppe
                    FROM faecher f
                    JOIN noten n ON f.fach_id = n.fach_id
                    JOIN schueler s ON n.schueler_id = s.schueler_id
                    WHERE (""" + " OR ".join(["s.klasse LIKE ?"] * len(patterns)) + """)
                      AND n.schuljahr = ?
                      AND n.halbjahr = ?
                      AND COALESCE(s.is_active, 1) = 1
                """
                cursor = conn.execute(query, [*patterns, school_year, term])
                rows = cursor.fetchall()
                
                for row in rows:
                    fach_lang = row["fach_lang"]
                    fach_kurz = row["fach_kurz"]
                    fach_id = row["fach_id"]
                    
                    fach_canonical = self._get_canonical_name(fach_kurz, fach_lang)
                    
                    # Check Config first
                    config_status = SUBJECT_STATUS_CONFIG.get(fach_canonical, "")
                    is_wpu_config = "WPU" in config_status
                    
                    # Check DB flags
                    is_wpu_db = bool(row["ist_wahlpflicht"]) or (row["wahlpflicht_gruppe"] and ("WPU" in row["wahlpflicht_gruppe"] or "WP" in row["wahlpflicht_gruppe"]))
                    
                    # Entscheidung: Ist es regulär?
                    is_regular = False
                    if config_status in ["Hauptfach", "Nebenfach"]:
                        is_regular = True # Override: Config sagt regulär (z.B. Praxistag)
                    elif is_wpu_config:
                        is_regular = False # Override: Config sagt WPU
                    else:
                        is_regular = not is_wpu_db # Fallback auf DB
                    
                    if is_regular:
                        regular_subjects[fach_canonical] = fach_id # Store ID keyed by CANONICAL NAME
                        
        except Exception as e:
            logging.error(f"Fehler beim Laden der Jahrgangsfächer: {e}")
            
        return regular_subjects

    def show_linux_help(self):
        """Zeigt Linux-Hilfe"""
        help_text = """Linux-Hilfe - Optimierte Version

SCHNELLSTART:

1. chmod 755 templates/ output_word/ output_database/ logs/
2. python kopfnotenapp.py

FEATURES:

• Template-Designer: Einfache Template-Erstellung
• Optimierter Export: Horizontale 3-Zeilen-Tabellen mit Beschriftungsspalte
• Einzelschüler-Export: Export für einzelne ausgewählte Schüler
• Vereinfachte Noten-Bearbeitung

PROBLEMLÖSUNG:

• Bei Berechtigungsfehlern: sudo chown $USER:$USER -R ./
• Templates in LibreOffice Writer bearbeiten
• Logs prüfen: tail -f logs/kopfnoten_gui.log"""
        messagebox.showinfo("Linux-Hilfe", help_text)

    def on_closing(self):
        """Behandelt Schließen"""
        try:
            self.save_sph_config()
            self.save_sph_missing_overview()
        except Exception:
            pass
        self.root.destroy()

def main():
    """Hauptfunktion"""
    try:
        # Verzeichnisse erstellen
        APP_PATHS.ensure_runtime_dirs()
        # Anwendung starten
        app = KopfnotenGUI()
        app.root.mainloop()
    except Exception as e:
        logging.error(f"Kritischer Anwendungsfehler: {e}")
        print(f"Fehler beim Starten: {e}")

if __name__ == "__main__":
    main()
